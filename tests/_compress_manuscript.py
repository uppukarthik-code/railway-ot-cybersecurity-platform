"""
_compress_manuscript.py

Editorial compression of the IRSE manuscript narrative prose to the
4,000-5,000-word band. Loads outputs/IRSE_OT_SOC_Final_Submission.docx (which
already carries the Author Note, Appendices A-E and the terminology fix) and
replaces selected verbose body paragraphs with tighter wording.

GUARANTEES (compression is editorial only):
  - No table, figure, caption or reference paragraph is touched.
  - Every standard ([n]), table ref, figure ref, UC-id, FR-id, finding term
    and the conclusion's recommendations are preserved in the rewrites.
  - No finding, count or conclusion is altered.
Target metric: body narrative prose words in [4000, 5000].
"""

import re
from pathlib import Path
import docx

ROOT = Path(__file__).resolve().parent.parent
OUT = ROOT / "outputs"
SRC = OUT / "IRSE_OT_SOC_Final_Submission.docx"
DST = OUT / "IRSE_OT_SOC_Final_Submission_5000W.docx"


def wc(s):
    return len(re.findall(r"[A-Za-z0-9][A-Za-z0-9'\-]*", s))


# Compressed replacements keyed by the LEADING TEXT of the target paragraph
# (robust to index drift). Each value is the tightened paragraph.
REPLACEMENTS = {
"The nationwide deployment of the Train Collision Avoidance System (Kavach), electronic interlocking (EI), centralised":
 "The nationwide deployment of the Train Collision Avoidance System (Kavach), electronic interlocking (EI), centralised Traffic Management Systems (TMS) and converged Multi-Protocol Label Switching (MPLS) telecom backbones has transformed Indian Railways signalling from electrically isolated vital circuits into a networked, safety-critical cyber-physical system, exposing Operational Technology (OT) to a threat surface that enterprise Information Technology (IT) Security Operations Centres (SOCs) are neither designed nor competent to monitor. This paper presents an architecture-driven methodology for planning and establishing dedicated OT Security Operations Centres (OT-SOCs) for Indian Railways. Its central thesis is that a machine-validated, standards-traceable model of the signalling architecture - its IEC 62443 zones and conduits, EN 50159 transmission categories and Safety Integrity Level (SIL) allocations - must precede and continuously inform SOC operating-model selection, sensor placement, detection engineering, incident response and governance. The evidence base is an independent cyber security assessment of a reference railway OT architecture MODEL with formally separated and validated semantic, governance and enforcement authorities; the findings are therefore properties of that model and are presented as illustrative monitoring requirements rather than measured deficiencies of any deployed installation, with field validation on representative Kavach and EI installations identified as essential future work. With that scope understood, the assessment surfaced durable engineering findings (absence of cryptographic authentication on vital Category-1 buses, absence of multi-factor authentication on privileged engineering access, absence of timeliness monitoring on open radio bearers, incompletely instrumented trust boundaries and an unprotected telecom backhaul) and standards-compliance findings (the EN 50159 / IEC 62443 authentication-crediting tension and the need for TS 50701 lifecycle evidence). This evidence is integrated with field-proven OT-SOC practice spanning operating-model selection, a five-layer visibility architecture, SL-T derivation, the Monitoring Paradox, a four-gate safety-first incident-response doctrine, a SEV1-SEV4 severity model, threat hunting, a three-audience KPI framework, staffing and governance, residual-risk management for SIL-4 systems, a maturity roadmap and a 90-day pilot programme.",

"That premise no longer holds.":
 "That premise no longer holds. The contemporary Indian Railways signalling estate is a converged, IP-enabled, geographically distributed cyber-physical system. Electronic interlockings exchange vital state with object controllers over digital buses; the Kavach Automatic Train Protection (ATP) system extends vital communication across an open radio-frequency (RF) air interface between the Stationary Kavach (S-Kavach) at the wayside and the Loco Kavach (L-Kavach) onboard; TMS and Network Management Systems (NMS) supervise wide geographies over MPLS backbones; and engineering workstations (EWS) able to download vital logic connect, via VPN, jump hosts and the industrial de-militarised zone (IDMZ), into the heart of the safety domain. Each capability is a deliberate engineering advance; each also opens a conduit through which an adversary may attempt to influence a safety function.",

"The instinctive institutional response":
 "The instinctive response - to extend the enterprise IT SOC to cover signalling - is inadequate and, in places, dangerous. Cross-sector OT-SOC experience [17] and the wider OT-security literature [14] are unambiguous: extending an IT SOC into OT consistently fails because the tools, telemetry, threat models, playbooks and expertise required are fundamentally different, producing a false sense of security. An IT SOC optimises for confidentiality; an OT-SOC must optimise for safety and for the availability of a physical process. An IT SOC may quarantine an endpoint on detection; an OT-SOC that interrupts a vital communication path may itself precipitate the hazard it was deployed to prevent.",

"This paper argues that the establishment of OT-SOC capability":
 "This paper argues that OT-SOC capability for Indian Railways must be a discipline in its own right, whose single most important precondition is an architecture validated against standards before it is monitored. One cannot monitor what one has not modelled, prioritise detection on a flow whose criticality, transmission category and trust boundary are unknown, or defend a zone whose conduits have not been enumerated. The argument is developed from concrete evidence and fused with field-proven OT-SOC practice, so that Indian Railways need not re-derive at cost what other critical-infrastructure sectors have already learned.",

"Four phases characterise the trajectory":
 "Four phases characterise railway signalling cyber security in India. Phase I - electromechanical isolation - comprised relay interlockings and token working; security was physical and procedural. Phase II - first-generation electronics - introduced solid-state and electronic interlockings, data loggers and digital axle counters; digital logic entered the vital domain but communication remained local, proprietary and serial, and data loggers became the first forensic record of security-relevant events. Phase III - networked OT - brought MPLS backbones, centralised TMS, NMS and remote diagnostics, the first genuine IT/OT convergence; engineering access became remotely brokered and the threat surface expanded from the fence to the network. Phase IV - safety over open bearers - is the present, defined by the nationwide rollout of Kavach, which carries vital movement-authority information across an open RF bearer and integrates with the EI at the station. For the first time a vital, SIL-4 function depends on a transmission system that is open in EN 50159 terms. The RDSO Kavach specification [9] and the Kavach Cyber Security SOP [10] recognise that cryptographic protection of the Kavach interface is necessary but not sufficient: protection without monitoring gives no assurance it is functioning and no evidence for incident response. The OT-SOC is the institutional mechanism through which protection is continuously verified.",

"Railway cyber security has an established literature":
 "Railway cyber security has an established literature on which this paper builds. The CENELEC lifecycle standards EN 50126 [7] and EN 50129 [8] and the dedicated cyber standard CLC/TS 50701 [5] provide the safety-security integration framework; IEC 62443 [1-4] supplies the zone/conduit and security-level model adopted here; and EN 50159 [6] supplies the communication-threat taxonomy. EU research programmes (CYRAIL, X2Rail, SAFETY4RAILS) [21] have addressed railway threat assessment and resilience, and ENISA has published transport and railway threat-landscape analyses [20]. In the operational domain, NIST SP 800-82r3 [14] and the ISA/IEC 62443 series define monitoring and management-system practice, MITRE ATT&CK for ICS [19] provides the adversary-technique reference, and cross-sector OT-SOC practice [17] supplies the operating-model, visibility, incident-response, KPI and maturity frameworks adapted in Sections 8-21. The distinctive contribution here is twofold: the argument that a machine-validated, authority-separated architecture model must precede and configure the OT-SOC, and the end-to-end railway contextualisation - from operating-model selection to a 90-day programme - anchored to Kavach and Indian Railways' institutional structure. The structural frameworks adapted here are drawn from cross-sector practice and not claimed as original.",

"The worst-case outcome is not data loss":
 "The worst-case outcome is not data loss but a wrong-side failure: a signal cleared against a conflicting movement, a point moved under a train, a movement authority issued where none is safe, or a suppressed emergency brake command. The SIL of the affected asset is therefore the primary axis along which monitoring effort is prioritised. The adversary spectrum runs from the opportunistic (commodity malware via an EWS or USB device) through the capable insider (a maintainer with privileged access and protocol knowledge) to the resourced nation-state actor. The capable insider is of particular concern because the vital buses rely, by standard, on closed-transmission assumptions and safety coding rather than cryptographic source authentication - robust against the outsider but weaker against an actor inside the closed segment. Sectoral data indicate that most OT incidents requiring SOC response begin with human error or supply-chain compromise rather than sophisticated attack [14,17,20]; a railway OT-SOC must therefore detect configuration drift, accidental connectivity changes and unauthorised removable media as reliably as it hunts nation-state implants.",

"EN 50159 [6] classifies safety-related transmission systems":
 "EN 50159 [6] classifies safety-related transmission systems into three categories by the controllability of the medium and the consequent need for cryptographic protection. In Category 1 (closed) the participants and transmission characteristics are fixed and unauthorised access is excluded by design, so defences are non-cryptographic (safety code, sequence number, time-out). In Category 2 (controlled) the medium is not fully under the designer's control but the residual risk of unauthorised access is assessed as negligible, so cryptographic protection is not mandated. In Category 3 (open) unauthorised access must be assumed, so cryptographic procedures are required in addition to the non-cryptographic defences. The Kavach radio interface and the balise air-gap are Category 3; the managed MPLS/IP backhaul is Category 2; the vital station buses are Category 1. Because the category fixes the threat model, it fixes the detector set the OT-SOC must field; Table 2 makes the threat-defence-detection relationship explicit.",

"Two consequences follow and resolve the central evidence finding.":
 "Two consequences follow and resolve the central evidence finding. First, for Category-1 vital buses EN 50159 does not mandate cryptographic source authentication - the closed-medium assumption and the safety code carry the assurance - so the reference model's vital protocols provide integrity and replay but not authentication; this is standards-consistent PROVIDED the closed-medium assumption holds, and the OT-SOC's role is to police that assumption continuously. Second, for Category-3 open bearers cryptographic protection is mandatory but not sufficient: timeliness, replay and RF-anomaly monitoring remain obligatory because delay and masquerade are first-class EN 50159 threats that cryptography alone does not detect in operation.",

"The central methodological proposition is that architecture validation":
 "The central methodological proposition is that architecture validation must precede OT-SOC deployment. A SOC detects deviation from expected behaviour, and the definition of expected is the architecture; if the architecture is not modelled, the baseline is whatever traffic existed when monitoring began - including any pre-existing compromise. In the reference assessment [18] the architecture was expressed as three separated authorities: a semantic authority (the ontology of asset types, protocols, conduit classes, Purdue levels and protocol capabilities), a governance authority (zoning, trust boundaries, flow rules, conduit security profiles and EN 50159 transmission categories) and an enforcement authority (the validators that consume the first two and emit findings). This separation lets a deviation be classified unambiguously as a real security gap, a standards-compliance gap, a modelling defect or an enforcement defect.",

"The assessment demonstrated this concretely.":
 "The assessment demonstrated this concretely. Before validation, several legitimate flows - operator workstation to TMS, the network-management plane, the telecom backhaul and the Loco-to-Stationary Kavach safety association - were unmodelled, and a vital-to-open-RF flow that should never exist (S-Kavach directly cabled to the radio base station, Figure 3) was present but indistinguishable from the legitimate ones. Only after the governance model was completed could the forbidden flow be classified as forbidden and prevented, and the legitimate flows classified as authorised-but-monitored. The validated architecture is thus not documentation; it is the configuration source for every sensor, detector and boundary control the OT-SOC deploys, and the residual findings become its standing monitoring requirements (consolidated in the residual-risk register of Section 19).",

"The findings cited here were obtained by assessing":
 "The findings cited here were obtained by assessing a reference architecture MODEL, not by instrumenting deployed Kavach or electronic-interlocking equipment. They are properties of the model and of the governance rules it encodes; their value is to illustrate, traceably, the classes of monitoring requirement an OT-SOC must satisfy. Whether any specific finding holds for a given fielded installation is a question of fact for a field assessment under the TS 50701 lifecycle. Such field validation - on a representative Kavach corridor and electronic interlocking - is the immediate next step and is recorded as future work (Section 23).",

"TS 50701 [5] and IEC 62443-3-2 [1] require that target Security Levels":
 "TS 50701 [5] and IEC 62443-3-2 [1] require that target Security Levels be derived for each zone and conduit from a documented risk assessment, and the OT-SOC uses SL-T to calibrate monitoring intensity and response targets. This paper recommends a consequence-by-likelihood method. Consequence is scored primarily by the safety significance of the affected function - its SIL and its potential to cause a wrong-side failure - and secondarily by operational and reputational impact; a SIL-4 vital function takes the highest consequence band by default. Likelihood is scored from exposure (enterprise/Internet reachability, open RF, vendor/engineering access), the EN 50159 transmission category and the strength of existing controls. The SL-T for a zone is the level required to reduce the residual risk of its highest-consequence credible scenario to acceptable; conduits inherit the higher SL-T of the zones they join. Table 3 is a worked extract for a Kavach corridor; the full register is an artefact of the cyber security case and the Phase-0 prerequisite for any sensor deployment.",

"Every piece of monitoring infrastructure deployed inside the OT network":
 "Every piece of monitoring infrastructure inside the OT network is itself an attack surface [17]; a compromised sensor on a vital network is a bridge into the most sensitive environment a railway operates. In the railway context this Monitoring Paradox has a sharper edge than in any other sector, because the monitored assets are SIL-4 vital systems whose disturbance can directly cause a wrong-side failure. The passivity requirement has a precise meaning: on vital (Category-1) and onboard segments the collection device must be a passive optical or copper TAP incapable of transmitting onto the monitored link; SPAN mirroring, a switch function that can silently drop frames under load, is excluded from vital segments and reserved for non-vital ones. The Monitoring Paradox is then enforced architecturally (Figure 2): telemetry leaves the Station echelon strictly upward through a data diode, there is no online downward control path into the vital tier, and sensor configuration is applied through a separate, scheduled, out-of-band management channel under change control.",

"Terminology. To avoid ambiguity this paper distinguishes":
 "Terminology. To avoid ambiguity this paper distinguishes three uses of the word formerly rendered as 'tier': the geographic SOC structure in ECHELONS - E1 Station/Wayside, E2 Zonal/Divisional, E3 National; the analyst workforce in GRADES - G1 Operator, G2 Investigator, G3 Senior/Hunter, G4 Lead; and the KPI framework by reporting AUDIENCE - Board, SOC-Lead, Analyst. The federated hybrid of Section 8 is realised as a three-echelon architecture (Figure 2). E1 (Station/Wayside) is co-located with the relay room - passive TAP sensors on field and interlocking conduits, data-logger ingestion and a data-diode log export; it performs local detection and buffers evidence, never hosting an in-line control on a vital path. E2 (Zonal) hosts the SIEM, correlation, IT/OT-boundary sensors in detection mode, vulnerability management for non-vital assets and case management, correlating events across stations. E3 (National) is custodian of the validated architecture and detection-content repository, the SL-T and cyber security case, cross-zone correlation, national threat intelligence, the Railway-sector ISAC and KPI consolidation; when the architecture changes it re-validates and pushes updated baselines downward. Log flow is strictly upward and diode-enforced across the IT/OT boundary, preserving FR5 within the SOC's own architecture.",

"Detection engineering converts the validated architecture":
 "Detection engineering converts the validated architecture and its findings into deployable content: every conduit security requirement and every assessment finding becomes a detection use case. The vital Category-1 buses - the EI to S-Kavach interface, EI to object controller and speed-sensor to L-Kavach over CAN, and L-Kavach to driver-machine interface and brake-interface unit over MVB (identifiers per the reference model and [9]) - provide integrity and replay through safety coding but not cryptographic source authentication; the SOC cannot add authentication but must detect the symptoms its absence would allow. Privileged engineering access is not, in the modelled flows, MFA-gated, so every privileged session is a high-value target. The open RF conduits require timeliness, jamming and RF-anomaly monitoring; several trust boundaries are not fully instrumented; and the telecom backhaul carries traffic without an evidenced transport-security overlay. Table 6 maps threats to detections; Table 7 is the use-case library; and Table 8 maps the use cases to MITRE ATT&CK for ICS [19] so coverage can be assessed against a recognised adversary model. Each use case also records, in the content repository, its data source, detection logic, validated playbook reference and documented false-positive baseline.",

"The scarcest resource is the individual fluent":
 "The scarcest resource is the individual fluent in both signalling safety and cyber security; cross-sector experience confirms a 12-24-month development pipeline and warns that neither IT-only analysts nor signalling engineers without security training produce effective practitioners [17]. The model builds cyber capability atop IRSSE signalling expertise and embeds dedicated railway expertise with no analogue in a generic OT-SOC: the IRSSE signalling SME, the Kavach SME and the Functional Safety specialist (Table 13). Analyst grades G1-G4 are distinct from the geographic echelons E1-E3; the SMEs are G2/G3 roles serving across echelons. A pilot corridor requires of the order of 14-18 full-time equivalents for 24x7 cover (a figure adopted from [17] for validation during the pilot); RDSO is custodian of the competency curriculum and certification.",

"Governance spans the full institutional hierarchy":
 "Governance spans the full institutional hierarchy and the national cyber security apparatus. Signalling and train-protection systems are Critical Information Infrastructure; their protection engages NCIIPC for designation and oversight under section 70 of the Information Technology Act 2000 [12], and CERT-In for statutory incident reporting under the CERT-In Directions of 28 April 2022, which require reporting within six hours of noticing a reportable incident [13], alongside the internal railway hierarchy. The SOC charter must define the interplay: the OT-SOC notifies the internal RDSO/Railway-Board chain and, for incidents meeting statutory thresholds, CERT-In within the mandated window, with NCIIPC engaged for CII-level events. Table 14 defines the governance model.",

"Residual risk must be quantified, documented and owned.":
 "Residual risk must be quantified, documented and owned. The open findings of the reference assessment are entered into a register with named ownership and a treatment decision (Table 15) and remain visible as standing monitoring requirements. For SIL-4 systems three constraints apply. First, residual risk affecting a SIL-4 vital function cannot be accepted at SOC or divisional level; it is escalated to the Railway Board with a supporting safety and cyber security argument. Second, where a control mandated by IEC 62443 cannot be implemented for sound engineering reasons - the EN 50159 / IEC 62443 authentication tension on vital Category-1 buses is the archetype - the residual risk is reduced As Low As Reasonably Practicable through compensating controls, of which the OT-SOC's continuous monitoring is one, and is argued in the cyber security case. Third, no risk treatment may weaken a SIL-4 boundary, an EN 50159 protection, an authentication or MFA requirement, an integrity or replay-protection requirement, or a trust-boundary control.",

"At the technical level the assessment's findings are durable":
 "The assessment's technical findings are consolidated as the standing monitoring requirements developed in Sections 6, 12 and 19 and as the residual-risk items of Section 19, and are not repeated here. At the programme level the lessons are distinct: an OT-SOC is an operational capability, not a technology purchase; the mundane and the sophisticated both matter; talent is the binding constraint; the SOC must remain outside the vital network; and SIL-4 residual risk must be owned at Board level.",

"The convergence that has modernised Indian Railways signalling":
 "The convergence that has modernised Indian Railways signalling has dissolved the isolation on which signalling security historically depended. Protecting this estate requires dedicated OT-SOCs that are passive-first, safety-aware, SIL-prioritised, protocol-literate and bound to a validated architecture. The central contribution of this paper is the argument, evidenced by an independent standards-traceable model assessment and reinforced by field-proven practice, that architecture validation must precede OT-SOC deployment, together with a demonstration of how a validated architecture mechanically yields the SOC's operating model, visibility plan, SL-T calibration, detection specification, response and severity models, KPIs, staffing, governance, residual-risk register and maturity roadmap. The recommendation to Indian Railways, RDSO and the Railway Board is unambiguous: invest first in the validated architecture and the TS 50701 evidence base; adopt a federated hybrid SOC model with explicit OT decision rights; instrument the vital conduits and trust boundaries before broadening coverage; build the dual-competency workforce under RDSO certification; govern across the Railway Board, RDSO, NCIIPC, CERT-In, the Zonal Railways and the Divisions; and constitute the OT-SOC as a permanent, standards-anchored assurance organ. Isolation can no longer be assumed; it must now be observed.",

"IEC 62443-3-2 (risk assessment for system design) [1]":
 "IEC 62443-3-2 (risk assessment for system design) [1] and IEC 62443-3-3 (system security requirements and security levels) [2] decompose a system into zones connected by conduits and define seven Foundational Requirements: FR1 Identification and Authentication, FR2 Use Control, FR3 System Integrity, FR4 Data Confidentiality, FR5 Restricted Data Flow, FR6 Timely Response to Events and FR7 Resource Availability. FR6 (and SR 6.2 continuous monitoring) is itself a monitoring mandate. The reference architecture realises the model as a zone hierarchy mapped to the Purdue model - Enterprise (L5), IDMZ and Security-Management (L3.5), Operations (L3), Telecom Core and Radio Access (L2/L1 telecom), Interlocking (L2/L1 signalling), Field (L0) and Onboard. The dual use of Levels 1-2 for the telecom and signalling sub-domains is a railway-specific extension of the Purdue model, annotated in Figure 1.",

"The RDSO Kavach specification [9] and the Kavach Cyber Security SOP [10] impose":
 "The RDSO Kavach specification [9] and the Kavach Cyber Security SOP [10] impose railway-specific obligations: cryptographic protection of the Kavach radio interface, integrity and replay protection of vital telegrams, controlled key management, and continuous-monitoring and incident-reporting duties the OT-SOC discharges. The interlock is this paper's conceptual core: IEC 62443 tells the OT-SOC where the boundaries are; EN 50159 what to look for on each conduit; TS 50701 why and to what assurance target; and the Kavach SOP the railway-specific obligations it must satisfy and report against.",

"Recommendation - a federated hybrid.":
 "Recommendation - a federated hybrid. A dedicated national OT-SOC (Model A) at echelon E3, under RDSO and Railway Board governance, holds the validated architecture, the SL-T allocation, the cyber security case, national threat intelligence and a Railway-sector ISAC function (Model D). Zonal OT tiers (Model B), fused with existing IT-SOC infrastructure but with explicit authority to override IT actions affecting signalling, provide regional correlation. A distributed sensing layer (Model E) instruments the dispersed wayside, interlocking and Kavach sites. Model C is admissible only as a time-boxed transitional measure. Across all echelons the single non-negotiable control is a charter defining decision rights: the OT tier must be able to block any action - including from IT-SOC staff - that would affect a vital system.",

"Two planning points follow.":
 "Two planning points follow. First, the highest-consequence telemetry - the process/vital-state layer for EI and Kavach - is the hardest to obtain, because vital devices emit few logs and cannot be actively interrogated; it must be derived from passive observation and from data loggers and onboard recorders. The balise air-gap concerns the Kavach trackside RFID tag and onboard RFID reader; these are Kavach RFID balises, distinct from the ETCS Eurobalise and not to be conflated with it. Second, the identity layer (engineering and vendor/VPN access) is most often neglected yet most directly tied to the highest-value attack path, the EWS.",

"These constraints are grounded in the standards, not in preference.":
 "These constraints are grounded in the standards, not preference. IEC 62443 FR7 (Resource Availability) and the OT priority order (Safety, then Availability, then Integrity, then Confidentiality) forbid security mechanisms that degrade the process. EN 50159's treatment of delay and corruption as first-class threats means an in-line device on a SIL-4 bus would be a potential source of the very hazards the safety layer is designed to exclude. The Monitoring Paradox is thus the railway expression of the principle that in OT the cure must never be more dangerous than the disease.",

"Railway OT incident response inverts the IT assumption":
 "Railway OT incident response inverts the IT assumption that speed is paramount: it assesses safety impact before every action, contains only when containment does not create a more dangerous state, and recovers with Operations and Engineering. A four-gate model (Figure 4, Table 10) governs every action, and no gate permits an action that interrupts a vital path or induces a wrong-side state. The repertoire is graduated - observe and enrich, then alert the controller, then recommend a manual protective action - the last always executed by signalling operations under safe-working rules, never automatically by the SOC tool.",

"Reactive detection cannot find an adversary":
 "Reactive detection cannot find an adversary who has pre-positioned without triggering a rule; proactive threat hunting can. In the railway OT-SOC, hunting is hypothesis-led and anchored to MITRE ATT&CK for ICS [19] and the validated architecture: hunters look for the Table 8 techniques that would precede a wrong-side failure - rogue masters or hardware additions on a vital segment (UC-08, UC-13), program-download activity outside a change window (UC-05), and adversary-in-the-middle indicators on the Kavach bearer (UC-03). Hunts run entirely on passively-collected E2/E3 data and on data-logger and onboard-recorder archives, never by active interrogation of vital devices. Each hunt yielding a repeatable indicator is promoted to a new detection use case, feeding the new-use-cases KPI (Section 17). A minimum of two hunt campaigns per quarter is recommended, rising for high-risk corridors.",

"The immediate next step is field validation":
 "The immediate next step is field validation of the model-derived findings on a representative Kavach corridor and electronic interlocking under the TS 50701 lifecycle. Further work includes resolving the EN 50159 / IEC 62443 authentication tension in the cyber security case; curating railway-specific detection-content libraries (Kavach interfaces, MVB and balise telegrams) under RDSO; machine-assisted architecture-conformance monitoring that diffs the validated model against observed topology (extending UC-08); integration of the safety and security cases under TS 50701 and EN 50129 cross-acceptance; resilience engineering of the OT-SOC itself; a Railway-sector ISAC aligned to NCIIPC and CERT-In; and standardisation of OT-SOC interfaces to the Kavach Cyber Security SOP.",

"Severity in OT is defined by potential physical consequence":
 "Severity in OT is defined by potential physical consequence, not data value. Table 11 anchors severity to the wrong-side failure and the SIL of the affected asset. The escalation timelines are the authors' recommended targets and must be reconciled against the Kavach Cyber Security SOP [10] reporting timelines and the statutory CERT-In window [13] during Phase 0 of the programme.",

"OT-SOC performance is measured by safety-weighted":
 "OT-SOC performance is measured by safety-weighted, standards-anchored indicators rather than volume metrics, spanning leading, lagging and operational-health classes across three reporting audiences with concrete, time-bound targets (Table 12). The mean-time-to-detect target is anchored to IEC 62443-3-3 SR 6.2 (continuous monitoring) and the operational hazard-exposure budget, not to EN 50129. Quantitative staffing and uptime figures are adopted from cross-sector practice [17] for validation against an Indian Railways workload study during the pilot.",

"The cross-sector 90-day launch programme [17] is adapted":
 "The cross-sector 90-day launch programme [17] is adapted to a pilot on a representative Kavach corridor, assuming Phase 0 (validated architecture, conduit inventory and TS 50701 SL-T allocation) is complete - no sensor is deployed before the model that configures it. The guiding principle is that breadth never precedes vital depth, and deployment never precedes the validated model. An indicative budget envelope (passive TAP/NDR, OT-contextualised SIEM, RF monitors, staffing and RDSO certification) should accompany the programme for Railway Board approval; a corridor-scale pilot is modest relative to the Kavach programme and is the prerequisite for credible national costing.",
}


def body_prose_wc(doc):
    """Word count of narrative prose (Normal, non-caption, before References,
    excluding front matter labels and appendices)."""
    total = 0
    in_refs = False
    appendix = False
    for p in doc.paragraphs:
        st = p.style.name
        t = p.text.strip()
        if st == "Heading 1":
            if t.lower() == "references":
                in_refs = True
            if t.startswith("Appendix"):
                appendix = True
            continue
        if st == "Heading 2":
            continue
        if in_refs or appendix:
            continue
        # skip captions and bold/italic front-matter labels
        isfmt = any((r.bold or r.italic) and r.text.strip() for r in p.runs)
        if t.startswith("Table ") and isfmt:
            continue
        if isfmt:  # title, subtitle, author note, keywords, nomenclature label, note
            continue
        total += wc(t)
    return total


def apply(doc):
    applied = 0
    saved = 0
    for p in doc.paragraphs:
        t = p.text
        for lead, new in REPLACEMENTS.items():
            if t.strip().startswith(lead):
                old_w = wc(t)
                # replace text in first run, clear the rest (paragraphs are single-run)
                if not p.runs:
                    p.add_run(new)
                else:
                    p.runs[0].text = new
                    for r in p.runs[1:]:
                        r.text = ""
                applied += 1
                saved += old_w - wc(new)
                break
    return applied, saved


def main():
    doc = docx.Document(str(SRC))
    before = body_prose_wc(doc)
    applied, saved = apply(doc)
    after = body_prose_wc(doc)
    doc.save(str(DST))
    print(f"replacements applied : {applied}/{len(REPLACEMENTS)}")
    print(f"body prose before    : {before}")
    print(f"body prose after     : {after}")
    print(f"words removed        : {before - after}")
    print(f"in 4000-5000 band    : {4000 <= after <= 5000}")
    # preservation check
    chk = docx.Document(str(DST))
    print(f"tables  : {len(chk.tables)}  images: {len(chk.inline_shapes)}  paras: {len(chk.paragraphs)}")


if __name__ == "__main__":
    main()
