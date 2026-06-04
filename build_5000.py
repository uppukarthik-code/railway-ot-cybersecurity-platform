# -*- coding: utf-8 -*-
"""Build IRSE_OT_SOC_Conference_Paper_5000W.docx (conference-length, <=5000 words)."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

d = Document()
n = d.styles["Normal"]; n.font.name = "Times New Roman"; n.font.size = Pt(10.5)
for hs, sz in (("Heading 1", 13.5), ("Heading 2", 11.5)):
    st = d.styles[hs]; st.font.name = "Times New Roman"; st.font.size = Pt(sz)
    st.font.color.rgb = RGBColor(0x10, 0x2A, 0x43)

def title(t, sz=16):
    p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(t); r.bold = True; r.font.size = Pt(sz); r.font.color.rgb = RGBColor(0x0B,0x1F,0x33)
def sub(t):
    p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(t); r.italic = True; r.font.size = Pt(10.5)
def h1(t): d.add_heading(t, level=1)
def h2(t): d.add_heading(t, level=2)
def P(t, bold=False, italic=False):
    p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(t); r.bold = bold; r.italic = italic
def KW(label, t):
    p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(label + " "); r.bold = True; p.add_run(t)
def bullet(t):
    d.add_paragraph(t, style="List Bullet")
def tbl(num, caption, headers, rows):
    cap = d.add_paragraph(); rc = cap.add_run("Table %d. %s" % (num, caption)); rc.bold = True; rc.font.size = Pt(9.5)
    t = d.add_table(rows=1, cols=len(headers)); t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, hh in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = ""; r = c.paragraphs[0].add_run(hh); r.bold = True; r.font.size = Pt(8.5)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ""; r = cells[i].paragraphs[0].add_run(str(v)); r.font.size = Pt(8.5)
    d.add_paragraph()
def figure(path):
    d.add_picture(path, width=Inches(6.2)); d.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    d.add_paragraph()

# ---------------- front matter ----------------
title("Cyber Security Framework Planning and Establishment of Operational Technology Security "
      "Operations Centres (OT-SOCs) for Indian Railways")
sub("[Author(s) and affiliation(s) per IRSE India template]")
sub("Keywords audience: IRSE India, RDSO, IRSSE officers, railway signalling engineers, OT cyber security practitioners")
d.add_paragraph()

h1("Abstract")
P("Indian Railways' deployment of Kavach, electronic interlocking (EI), centralised Traffic Management "
  "Systems (TMS) and converged Multi-Protocol Label Switching (MPLS) backbones has turned signalling "
  "from electrically isolated vital circuits into a networked, safety-critical cyber-physical system, "
  "exposing Operational Technology (OT) to threats that enterprise Information Technology (IT) Security "
  "Operations Centres (SOCs) cannot competently monitor. This paper sets out an architecture-driven "
  "methodology for establishing dedicated OT-SOCs for Indian Railways. Its thesis is that a "
  "standards-validated model of the signalling architecture - its IEC 62443 zones and conduits, EN 50159 "
  "transmission categories and Safety Integrity Level (SIL) allocations - must precede and configure SOC "
  "design, sensor placement, detection engineering and incident response. Evidence is drawn from an "
  "independent assessment of a reference architecture model (not a fielded installation), which surfaced "
  "durable monitoring requirements: vital buses without cryptographic authentication, privileged "
  "engineering access without multi-factor authentication, open radio bearers without timeliness "
  "monitoring, and incompletely instrumented trust boundaries. The paper contributes a railway-specific "
  "three-echelon OT-SOC architecture, an explicit IEC 62443 / EN 50159 / TS 50701 integration, and a "
  "federated operating model mapped to the Railway Board, RDSO, NCIIPC, CERT-In, Zonal Railways and "
  "Divisions, with railway detection use cases, a four-gate incident-response doctrine and a SIL-4 "
  "residual-risk method.")
KW("Keywords:", "OT-SOC; Indian Railways; Kavach; IEC 62443; EN 50159; TS 50701; Electronic "
   "Interlocking; zones and conduits; detection engineering; SIL-4.")

# ---------------- 1 ----------------
h1("1. Introduction and Background")
P("For most of its history, railway signalling safety rested on physical isolation and fail-safe design; "
  "the cyber threat surface was the perimeter fence. That premise no longer holds. Electronic "
  "interlockings exchange vital state with object controllers over digital buses; Kavach, the indigenous "
  "Automatic Train Protection (ATP) system, carries vital movement authority across an open "
  "radio-frequency (RF) bearer between the Stationary Kavach (S-Kavach) at the wayside and the Loco "
  "Kavach (L-Kavach) onboard, and integrates with the EI at the station; TMS and Network Management "
  "Systems (NMS) supervise wide geographies over MPLS; and engineering workstations (EWS) able to "
  "download vital logic reach the safety domain through VPN access, jump hosts and an industrial "
  "de-militarised zone (IDMZ). Four phases mark this evolution - electromechanical isolation, "
  "first-generation electronics (with data loggers as the first incidental forensic record), networked OT "
  "over the backbone, and the present phase of safety carried over open bearers, defined by the "
  "nationwide rollout of Kavach.")
P("Extending an enterprise IT SOC to cover signalling is inadequate and, in places, dangerous. An IT "
  "SOC optimises for confidentiality and may quarantine an endpoint on detection; an OT-SOC must "
  "optimise for safety and availability, and an automated action that interrupts a vital path may itself "
  "cause the hazard it was meant to prevent. Cross-sector OT-SOC practice and the wider OT-security "
  "literature concur that the tools, telemetry, threat models, playbooks and expertise differ "
  "fundamentally, so that an extended IT SOC yields coverage that looks credible on paper but is blind "
  "to the threats that matter [13,16]. This paper argues that OT-SOC capability for Indian Railways must "
  "be a discipline in its own right, whose decisive precondition is an architecture validated against "
  "standards before it is monitored: one cannot monitor what one has not modelled, nor prioritise "
  "detection on a flow whose criticality, EN 50159 category and trust boundary are unknown. Prior work "
  "provides the components - CENELEC EN 50126/50129 and CLC/TS 50701 for safety-security lifecycle "
  "integration, IEC 62443 for zones and conduits, EN 50159 for the communication threat taxonomy, NIST "
  "SP 800-82 and MITRE ATT&CK for ICS for OT monitoring and adversary modelling, and cross-sector "
  "OT-SOC foundational practice for operating-model and incident-response frameworks. This paper's "
  "contribution is to bind these into a railway-specific, architecture-driven OT-SOC for Indian Railways.")

# ---------------- 2 ----------------
h1("2. Railway OT Threat Landscape")
P("The railway OT landscape differs from the enterprise landscape in its worst-case consequence: not "
  "data loss but a wrong-side failure - a signal cleared against a conflicting movement, a point moved "
  "under a train, a movement authority issued where none is safe, or an emergency brake command "
  "suppressed. The SIL of the affected asset is therefore the primary axis along which monitoring is "
  "prioritised. The adversary spectrum runs from commodity malware on an EWS or USB device, through "
  "the capable insider with privileged access and protocol knowledge, to the resourced nation-state "
  "actor; sectoral incident data indicate that most OT incidents begin with human error or "
  "supply-chain compromise rather than sophisticated attack [13,16]. The capable insider is acute for "
  "railways because vital buses rely, by standard, on closed-transmission assumptions and safety "
  "coding rather than cryptographic source authentication. Table 1 maps the five recurring OT threat "
  "categories to their railway manifestations.")
tbl(1, "OT threat categories mapped to Indian Railways",
    ["Category", "Railway manifestation", "Consequence"],
    [["IT-sourced lateral movement", "Enterprise-IDMZ-operations-interlocking traversal through weak boundary controls", "Disruption of TMS/NMS; staging toward vital zones"],
     ["EWS compromise (TRITON class)", "EWS with logic-download access to EI / S-Kavach", "Alteration of vital logic - wrong-side failure"],
     ["Supply-chain / vendor access", "Vendor laptop or update channel for Kavach/EI/TMS", "Implant in the vital supply chain"],
     ["Insider / sabotage", "Maintainer with privileged engineering access", "Logic/setpoint change; alarm suppression"],
     ["Nation-state pre-positioning", "Persistent access awaiting operational timing", "Coordinated disruption of national rail safety"]])

# ---------------- 3 ----------------
h1("3. Standards Integration: IEC 62443, EN 50159, TS 50701 and Kavach")
P("Three standards families and the Kavach regime form the structural grammar of railway OT security, "
  "and the OT-SOC must rest on all of them. IEC 62443 decomposes the system into zones connected by "
  "conduits and defines seven Foundational Requirements (FR1 Identification and Authentication, FR2 Use "
  "Control, FR3 System Integrity, FR4 Data Confidentiality, FR5 Restricted Data Flow, FR6 Timely "
  "Response to Events, FR7 Resource Availability); FR6 and SR 6.2 (continuous monitoring) are "
  "themselves a monitoring mandate. The reference architecture realises this as a zone hierarchy mapped "
  "to the Purdue model - Enterprise (L5), IDMZ and Security-Management (L3.5), Operations (L3), Telecom "
  "Core and Radio Access (telecom L2/L1), Interlocking (signalling L2/L1), Field (L0) and Onboard - the "
  "dual L1/L2 use for the telecom and signalling sub-domains being a deliberate railway extension.")
P("EN 50159 classifies safety-related transmission into three categories that fix the threat model and "
  "hence the detector set. In Category 1 (closed) the participants and characteristics are fixed and "
  "unauthorised access is excluded by design, so the defences are non-cryptographic (safety code, "
  "sequence number, time-out). In Category 2 (controlled) the medium is not fully controlled but the "
  "residual risk of unauthorised access is negligible, so cryptography is not mandated. In Category 3 "
  "(open) unauthorised access must be assumed, so cryptographic procedures are required in addition. "
  "The Kavach radio interface and the balise air-gap are Category 3; the managed MPLS/IP backhaul is "
  "Category 2; the vital station buses are Category 1. Two consequences follow. First, for Category-1 "
  "vital buses EN 50159 does not mandate cryptographic source authentication - the closed-medium "
  "assumption and the safety code carry the assurance - so the absence of authentication on those buses "
  "is standards-consistent provided the assumption holds, and the OT-SOC's role is to police it. Second, "
  "for Category-3 bearers cryptography is necessary but not sufficient: timeliness, replay and "
  "RF-anomaly monitoring remain obligatory because delay and masquerade are first-class EN 50159 "
  "threats that cryptography does not detect in operation.")
P("CLC/TS 50701 binds IEC 62443 to the EN 5012x safety lifecycle and requires a documented risk "
  "assessment, the allocation of target Security Levels (SL-T) to zones and conduits, and a cyber "
  "security case analogous to the safety case; the SL-T calibrates monitoring intensity and response "
  "targets. The RDSO Kavach specification and the Kavach Cyber Security SOP add railway obligations - "
  "cryptographic protection of the radio interface, integrity and replay protection of vital telegrams, "
  "key management, and continuous-monitoring and incident-reporting duties the OT-SOC discharges. The "
  "interlock is the conceptual core of this paper: IEC 62443 defines where the boundaries are; EN 50159 "
  "defines what to look for on each conduit; TS 50701 defines why and to what assurance target; and the "
  "Kavach SOP defines the obligations to report against. Because each conduit's required controls are "
  "simultaneously its monitoring specification, the trust-boundary inventory is the OT-SOC sensor "
  "placement plan (Table 2).")
tbl(2, "Trust-boundary monitoring matrix (EN 50159 category drives the detector set)",
    ["Trust boundary", "EN 50159", "Required controls", "OT-SOC monitoring obligation"],
    [["Enterprise (L5) - IDMZ (L3.5)", "IT/OT divide", "Firewall, inspection", "Verify controls; detect unauthorised ingress; diode log export"],
     ["IDMZ - Operations (L3)", "-", "Firewall, inspection", "Brokered-access audit; jump-host/VPN session + MFA-event verification"],
     ["Operations - Interlocking (L2)", "Cat 1", "Firewall, inspection, integrity, authentication", "Highest scrutiny; engineering-access and logic-download detection"],
     ["Telecom Core - Radio Access", "Cat 2", "Firewall, inspection, authentication", "Transport-security (IPsec/MACsec) health; participant verification"],
     ["Radio Access - Onboard", "Cat 3 (open RF)", "Integrity, replay, authentication, radio monitoring", "Timeliness, jamming, RF anomaly, replay"],
     ["Balise air-gap (Field - Onboard)", "Cat 3 passive", "Integrity (passive telegram); crypto exempt", "Telegram integrity and positional plausibility; no crypto false positives"]])

# ---------------- 4 ----------------
h1("4. The Architecture-Driven Imperative")
P("The methodological core of this paper is that architecture validation must precede OT-SOC "
  "deployment. A SOC detects deviation from expected behaviour, and the definition of expected is the "
  "architecture; if it is unmodelled, the baseline is whatever traffic existed when monitoring began, "
  "including any pre-existing compromise. In the reference assessment the architecture was expressed as "
  "three separated authorities: a semantic authority (the ontology of asset types, protocols, conduit "
  "classes, Purdue levels and protocol capabilities), a governance authority (zoning, trust boundaries, "
  "flow rules, conduit security profiles and EN 50159 categories) and an enforcement authority (the "
  "validators that consume the first two and emit findings). This separation lets a deviation be "
  "classified unambiguously as a real security gap, a standards-compliance gap, a modelling defect or an "
  "enforcement defect, rather than collapsing them into undifferentiated alerts.")
P("The assessment demonstrated the value concretely. Before validation, legitimate flows - operator "
  "workstation to TMS, the network-management plane, the telecom backhaul and the Loco-to-Stationary "
  "Kavach safety association - were unmodelled, while a flow that should never exist (S-Kavach directly "
  "cabled to the radio base station, bypassing the radio gateway and the rf-transition boundary) was "
  "present yet indistinguishable from the legitimate ones. Only after the governance model was "
  "completed could the forbidden flow be classified and prevented and the legitimate flows authorised "
  "and monitored. The validated architecture is thus not documentation; it is the configuration source "
  "for every sensor, detector and boundary control. The residual findings of the assessment - notably "
  "the absence of cryptographic authentication on vital Category-1 buses, the absence of multi-factor "
  "authentication (MFA) on privileged engineering access, the absence of timeliness monitoring on the "
  "open RF bearer, and incompletely instrumented trust boundaries - become the OT-SOC's standing "
  "monitoring requirements (Sections 8 and 11).")
P("Limitations. These findings were obtained by assessing a reference architecture model, not by "
  "instrumenting deployed Kavach or interlocking equipment; they are properties of the model and its "
  "governance logic, illustrating the classes of monitoring requirement an OT-SOC must satisfy. Whether "
  "any specific finding holds for a given fielded installation is a question of fact to be established by a "
  "field assessment under the TS 50701 lifecycle, which is the immediate next step.")

# ---------------- 5 ----------------
h1("5. Risk Assessment and SL-T Derivation")
P("TS 50701 and IEC 62443-3-2 require SL-T to be derived for each zone and conduit from a documented "
  "risk assessment. This paper recommends a consequence-by-likelihood method. Consequence is scored "
  "primarily by the safety significance of the affected function - its SIL and its potential to cause a "
  "wrong-side failure - and secondarily by operational impact, so that a SIL-4 vital function takes the "
  "highest band by default. Likelihood is scored from exposure (enterprise reachability, open RF, "
  "vendor/engineering access), the EN 50159 category of the conduit and the strength of existing "
  "controls. The zone SL-T is the level needed to reduce the residual risk of its highest-consequence "
  "credible scenario to acceptable, and conduits inherit the higher SL-T of the zones they join. For a "
  "Kavach corridor this yields SL-T 3 for the Interlocking and Onboard zones (SIL-4, engineering and "
  "open-RF exposure), SL-T 2-3 for Operations and the IDMZ, and SL-T 2 for the Telecom Core. The full "
  "zone/conduit register is an artefact of the cyber security case and is the prerequisite for any sensor "
  "deployment.")

# ---------------- 6 ----------------
h1("6. OT-SOC Operating Model")
P("Five operating-model archetypes cover the realistic spectrum; the right choice depends on scale, "
  "geography, regulatory obligation and OT maturity. Table 3 evaluates them for a national railway with "
  "thousands of dispersed wayside and Kavach sites, a zonal/divisional structure and SIL-4 assets.")
tbl(3, "OT-SOC operating-model evaluation and recommendation",
    ["Model", "Character", "Fit for Indian Railways"],
    [["A Dedicated", "Separate stack, OT-native staff", "National level under RDSO; justified by SIL-4 criticality"],
     ["B Fused IT/OT", "Shared SIEM, ring-fenced OT tier", "Zonal level; needs OT veto over IT actions affecting OT"],
     ["C MSSP-led", "Monitoring outsourced", "Transitional only; mandatory knowledge transfer"],
     ["D Collaborative (ISAC)", "Pooled sector telemetry", "Railway-sector ISAC, RDSO-anchored, NCIIPC/CERT-In aligned"],
     ["E Virtual/distributed", "Lightweight sensors per site", "Essential for the dispersed wayside/Kavach estate"]])
P("Recommendation - a federated hybrid. A dedicated national OT-SOC (Model A) holds the validated "
  "architecture, SL-T allocation, cyber security case, national threat intelligence and a Railway-sector "
  "ISAC (Model D); fused zonal tiers (Model B) provide regional correlation with explicit authority to "
  "override IT actions affecting signalling; and a distributed sensing layer (Model E) instruments the "
  "wayside, interlocking and Kavach sites. Model C is admissible only as a time-boxed bridge. Across all "
  "tiers the single non-negotiable control is a charter defining decision rights: the OT tier must be able "
  "to block any action, including from IT-SOC staff, that would affect a vital system.")

# ---------------- 7 ----------------
h1("7. Railway OT-SOC Architecture: Visibility Layers and Three Echelons")
P("Comprehensive visibility requires telemetry from five layers - Network, Asset/Endpoint, "
  "Application/Supervisory, Process/Vital-state, and Identity and Access. Programmes that instrument only "
  "the network layer see the backbone but miss the process and identity layers, which carry the "
  "highest-consequence and highest-value signals. Table 4 applies the layers to the railway assets; "
  "Figure 1 places them in the zone/conduit reference model. Two points are decisive. The "
  "process/vital-state telemetry for EI and Kavach is the hardest to obtain, because vital devices emit "
  "few logs and cannot be actively interrogated, so it must be derived from passive observation and "
  "from data loggers and onboard recorders. The identity layer - engineering and vendor/VPN access - is "
  "most often neglected yet most directly tied to the EWS, the highest-value attack path. The balise "
  "air-gap concerns the Kavach trackside RFID tag and the onboard RFID reader; these are Kavach RFID "
  "balises and are distinct from the ETCS Eurobalise, with which they should not be conflated.")
figure("figs/fig1_zones.png")
tbl(4, "Five visibility layers applied to Indian Railways OT assets",
    ["Asset", "Dominant layer(s)", "Collection / key telemetry"],
    [["Electronic Interlocking", "Process/Vital-state; Network", "Passive TAP; route/occupancy state, interface reads/writes, safety-code failures, state-machine violations"],
     ["Kavach (S-/L-Kavach)", "Network; Process/Vital-state", "TAP + RF monitor; onboard recorder; movement-authority frames, sequence/freshness, RF anomaly"],
     ["MPLS / PE routers", "Network", "Passive flow; IPsec/MACsec tunnel health, routing anomaly, unexpected participants"],
     ["Radio networks (BTS)", "Network; Process/Vital-state", "RF spectrum + conduit TAP; signal-strength/error-rate anomaly, latency deviation, replay"],
     ["TMS", "Application/Supervisory", "Syslog/TLS, HTTPS; operator route/setpoint commands, command integrity, failed logins"],
     ["NMS", "Application; Identity", "SNMPv3; unexpected set operations, config changes, off-hours access"],
     ["Data loggers", "Process/Vital-state", "Read-only ingestion; forensic stream, tamper and time-sync drift, gap/reorder"],
     ["Engineering workstations", "Asset/Endpoint; Identity", "Agentless read-only, PAM/VPN recording; logic-download and config-change actions, USB, privileged logins"]])
P("The architecture is organised in three echelons, mirroring the railway's distribution and the IEC "
  "62443 trust hierarchy (Figure 2). To avoid overloading the word tier, the geographic structure uses "
  "echelons E1-E3, the analyst workforce uses grades G1-G4, and the KPI framework uses reporting "
  "audiences. Echelon E1 (Station/Wayside), co-located with the relay room, carries passive TAP sensors "
  "on field and interlocking conduits, data-logger ingestion and a data-diode log export; it detects "
  "locally, buffers evidence, and never hosts an in-line control on a vital path. Echelon E2 "
  "(Zonal/Divisional) hosts the SIEM, correlation, IT/OT-boundary intrusion sensors in detection mode, "
  "vulnerability management for non-vital assets and case management, and correlates across stations - "
  "for example a corridor-wide pattern of balise-telegram anomalies no single station would see. Echelon "
  "E3 (National), under RDSO, is custodian of the validated architecture and detection content, the SL-T "
  "and cyber security case, cross-zone correlation, threat intelligence and the ISAC. The Monitoring "
  "Paradox governs the design: every monitoring device inside the OT network is itself an attack surface, "
  "so on SIL-4 and onboard segments the collector must be a passive optical or copper TAP that cannot "
  "transmit onto the link (SPAN mirroring, which can drop frames under load, is reserved for non-vital "
  "segments); telemetry leaves E1 strictly upward through a data diode; and sensor configuration is "
  "applied through a separate, scheduled, out-of-band management channel, so there is no online "
  "downward path into the vital tier. This is grounded in IEC 62443 FR7 and in EN 50159, under which an "
  "in-line device could introduce the very delay and corruption the safety layer excludes.")
figure("figs/fig2_echelons.png")

# ---------------- 8 ----------------
h1("8. Detection Engineering and Railway Use Cases")
P("Detection engineering converts the validated architecture and its findings into deployable content: "
  "every conduit requirement and every assessment finding becomes a use case. The vital Category-1 "
  "buses (the EI to S-Kavach interface, EI to object controller and speed-sensor to L-Kavach, and "
  "L-Kavach to driver-machine interface and brake-interface unit) provide integrity and replay through "
  "safety coding but not cryptographic authentication; the SOC cannot add authentication but must "
  "detect the symptoms its absence would allow - unexpected source identifiers, protocol state-machine "
  "violations, and rates or sequences inconsistent with the vital application - which is equivalent to "
  "continuously policing the closed-medium assumption. Privileged engineering access is not MFA-gated "
  "in the modelled flows, so every privileged session is a high-value target; the open RF conduits "
  "require timeliness, jamming and RF-anomaly monitoring; and the telecom backhaul carries traffic "
  "without an evidenced transport-security overlay. Table 5 is the use-case library; each use case maps "
  "to MITRE ATT&CK for ICS and, in the deployed repository, records its data source, detection logic, "
  "validated playbook and false-positive baseline. A standards-justified exemption shows why the "
  "validated architecture matters: the passive balise telegram is protected by the safety-coded telegram "
  "with freshness from location coding and odometry, so the OT-SOC must not raise a missing-encryption "
  "or missing-replay alarm on it (a false positive) while still monitoring telegram integrity and positional "
  "plausibility (Figure 3 distinguishes the legitimate end-to-end Kavach association from the forbidden "
  "direct edge).")
figure("figs/fig3_kavach.png")
tbl(5, "Railway OT-SOC detection use cases (with ATT&CK for ICS mapping)",
    ["ID / use case", "Trigger", "Sev", "ATT&CK for ICS"],
    [["UC-01 Spoofed vital field message", "Source-ID not in EI whitelist / state-machine violation", "SEV1", "Spoof Reporting Message"],
     ["UC-02 Kavach EI-interface integrity loss", "Rising CRC / safety-code failure", "SEV1", "Modify Parameter"],
     ["UC-03 Kavach RF replay / delay", "Sequence regression / latency beyond SIL threshold", "SEV1", "Adversary-in-the-Middle"],
     ["UC-04 RF jamming / loss of comms", "Signal-strength collapse / error-rate spike", "SEV2", "Denial of Service"],
     ["UC-05 Unauthorised logic download", "Logic/config change from EWS toward EI/S-Kavach", "SEV1", "Program Download"],
     ["UC-06 Privileged access without MFA", "Jump-host/VPN session, no MFA event", "SEV2", "Exploitation of Remote Services"],
     ["UC-07 Trust-boundary bypass", "Flow crossing enterprise-operations without required controls", "SEV2", "Lateral Movement"],
     ["UC-08 Forbidden topology edge", "S-Kavach to BTS (or any forbidden flow) detected", "SEV1", "Rogue Master / Hardware Additions"],
     ["UC-09 Backhaul transport-security failure", "Loss/absence of IPsec/MACsec on TMS-MPLS", "SEV3", "Manipulation of Control"],
     ["UC-10 Balise integrity / position anomaly", "Telegram safety-code failure / position implausibility", "SEV2", "Spoof Reporting Message"],
     ["UC-11 New device on vital segment", "MAC/IP not in inventory on interlocking zone", "SEV2", "Rogue Master"],
     ["UC-12 Removable media on EWS", "USB device-ID not in approved list", "SEV3", "Replication Through Removable Media"]])

# ---------------- 9 ----------------
h1("9. Incident Response and Severity")
P("Railway OT incident response inverts the IT assumption that speed is paramount; it assesses safety "
  "before every action, contains only when containment does not create a more dangerous state, and "
  "recovers with Operations and Engineering. A four-gate doctrine governs every action, and no gate "
  "permits an action that interrupts a vital path or induces a wrong-side state. G1 (safety assessment, "
  "owned by a functional-safety engineer) must be cleared before any containment; G2 (operational "
  "impact) is decided by Divisional Control with the SOC; G3 (containment) by the SOC Lead and "
  "Operations Lead, recording the minimum reversible action; and G4 (recovery authorisation) by the "
  "RDSO or divisional safety authority after the affected segment is re-validated against the model. The "
  "response is graduated - observe and enrich, alert the controller, then recommend a manual protective "
  "action under safe-working rules (caution orders, degraded-mode working, temporary speed "
  "restrictions) - never automated by the SOC tool. Decision rights are explicit: the OT-SOC owns "
  "detection, enrichment and recommendation; signalling operations owns any action affecting train "
  "movement; and governance owns escalation and external reporting. Forensic readiness is a "
  "precondition - data loggers and onboard recorders must have assured integrity and "
  "time-synchronisation before an incident. Severity is defined by physical consequence, not data value "
  "(Table 6); escalation timelines are recommended targets to be reconciled with the Kavach SOP and "
  "the statutory CERT-In window in Phase 0.")
tbl(6, "Railway OT incident severity model",
    ["Level", "Definition (consequence-based)", "Example", "Escalation"],
    [["SEV1 Safety-critical", "Vital/Kavach/EI compromise confirmed or probable; potential wrong-side failure", "Manipulated EI logic; spoofed Kavach movement authority", "All-hands; G1 invoked; functional-safety + executive notice <=15 min; CERT-In within statutory window; NCIIPC engaged"],
     ["SEV2 Operation-critical", "Confirmed OT compromise able to disrupt operations; active lateral movement", "Ransomware on TMS; compromised EWS with live EI/Kavach link", "Lead escalation <=15 min; Operations <=30 min; CISO/IRSSE <=1 hr"],
     ["SEV3 OT-zone compromise", "Malicious presence, not yet affecting the vital process", "Malware on isolated EWS; rogue device", "1-hr investigation; Operations informed; CISO same day"],
     ["SEV4 Anomaly / policy", "Policy violation or anomaly, no confirmed compromise", "Unauthorised media; out-of-window vendor access", "Next-business-day triage; trend-tracked"]])

# ---------------- 10 ----------------
h1("10. Governance, Staffing and KPIs")
P("Governance must span the railway hierarchy and the national cyber-security apparatus. Signalling and "
  "train-protection systems are Critical Information Infrastructure, so their protection engages NCIIPC "
  "for designation and oversight under section 70 of the Information Technology Act 2000 and CERT-In for "
  "statutory reporting under its 2022 Directions (reporting within six hours of a reportable incident), "
  "alongside the internal Railway Board, RDSO, Zonal and Divisional chain. The staffing model builds "
  "cyber capability atop IRSSE signalling expertise - the scarcer and harder-to-acquire half - over a "
  "12-24 month pipeline, and embeds three railway roles with no analogue in a generic OT-SOC: the "
  "IRSSE signalling subject-matter expert, the Kavach subject-matter expert and the functional-safety "
  "specialist who owns the G1 gate. Analyst grades G1-G4 (operator, investigator, senior/hunter, lead) "
  "are distinct from the geographic echelons E1-E3. A pilot corridor needs of the order of 14-18 "
  "full-time equivalents for round-the-clock cover, a figure to be validated against an Indian Railways "
  "workload study; RDSO is the custodian of the competency curriculum and certification. Table 7 maps "
  "responsibilities and Table 8 gives concrete, time-bound KPIs measured by leading, lagging and "
  "operational-health classes; the mean-time-to-detect target is anchored to IEC 62443-3-3 SR 6.2, not "
  "to EN 50129.")
tbl(7, "Governance, staffing and responsibility model",
    ["Entity / role", "Echelon / grade", "Principal OT-SOC responsibility"],
    [["Railway Board", "National policy", "Mandate, funding, national risk-acceptance authority for SIL-4 residual risk"],
     ["RDSO", "National (E3)", "Custody of validated architecture and SL-T; detection-content repository; competency certification; national OT-SOC"],
     ["NCIIPC / CERT-In", "National CII / CERT", "CII designation and oversight (IT Act s.70); statutory 6-hour incident reporting; national threat intelligence"],
     ["Zonal Railways", "Zonal (E2)", "Fused zonal OT-SOC; cross-zone correlation; OT veto over IT actions"],
     ["Divisions", "Local (E1)", "Station/wayside collection; local detection; first-line response under safe-working rules"],
     ["G1-G4 analysts + IRSSE/Kavach/functional-safety SMEs", "Grades G1-G4", "Triage to hunting; SME architecture interpretation; G1 safety-gate ownership"]])
tbl(8, "Three-audience KPI framework (concrete targets)",
    ["Audience: KPI", "Target", "Class"],
    [["Board: vital (SIL-3/4) conduit monitoring coverage", "100% within 18 months; sustained", "Leading"],
     ["Board: MTTD (vital-affecting known technique)", "< 1 hr", "Lagging"],
     ["Board: incidents contained without operational impact", "> 90%", "Lagging"],
     ["SOC-Lead: trust-boundary instrumentation", "100% within 12 months", "Leading"],
     ["SOC-Lead: EN 50159 detector completeness", "100% (match each conduit's category)", "Leading"],
     ["SOC-Lead: privileged-session audit coverage", "100% within 24 hr", "Op-health"],
     ["Analyst: TAP coverage of L1/L2 vital segments", "100%", "Op-health"],
     ["Analyst: architecture-conformance drift; exempt-flow false positives", "0; 0", "Leading / Op-health"]])

# ---------------- 11 ----------------
h1("11. Residual Risk and SIL-4 Acceptance")
P("Residual risk must be quantified, documented and owned, not buried in assumption logs. The open "
  "findings of the assessment are entered into a register with named ownership and a treatment decision "
  "and remain visible as standing monitoring requirements: A-01 no cryptographic authentication on "
  "vital Category-1 buses (mitigate by closure monitoring, UC-01/02/11; resolve in the cyber security "
  "case); A-02 missing integrity on field-vital train detection; A-03 no MFA on privileged engineering "
  "access (mitigate by privileged-session monitoring, UC-05/06; enforce MFA in the architecture); A-04 no "
  "timeliness or RF-anomaly monitoring on the open RF bearer (UC-03/04); A-05 forbidden vital-to-open-RF "
  "edge (avoid and detect, UC-08); and A-06 incompletely instrumented trust boundaries. Three "
  "constraints apply to SIL-4 systems. Residual risk affecting a SIL-4 vital function cannot be accepted "
  "at SOC or divisional level; it is escalated to the Railway Board with a safety and cyber-security "
  "argument. Where a control mandated by IEC 62443 cannot be implemented for sound engineering "
  "reasons - the EN 50159 / IEC 62443 authentication tension being the archetype - the residual risk is "
  "reduced as low as reasonably practicable through compensating controls, of which the OT-SOC's "
  "continuous monitoring is one, and is argued in the cyber security case. And no treatment may weaken "
  "a SIL-4 boundary, an EN 50159 protection, an authentication, MFA, integrity or replay-protection "
  "requirement, or a trust-boundary control.")

# ---------------- 12 ----------------
h1("12. Implementation and Maturity Roadmap")
P("Deployment follows a phased path in which architecture validation and the TS 50701 evidence base "
  "precede sensor deployment, and vital-asset depth precedes breadth. Phase 0 produces the validated "
  "architecture, conduit inventory and SL-T allocation. A 90-day pilot on a representative Kavach "
  "corridor then follows: Days 1-30 confirm the architecture and SL-T, ratify the SOC charter and "
  "decision rights, identify TAP points and assign signalling-first staff; Days 31-60 deploy OT-native "
  "network detection with railway-protocol parsing and an OT-contextualised SIEM, implement use cases "
  "UC-01 to UC-08 and table-top the four-gate playbooks with Operations and the functional-safety "
  "specialist; Days 61-90 go live, tune false positives (notably the passive-telegram exemption), run "
  "one severity-graded incident exercise under safe-working rules, publish the first KPI dashboards and "
  "route SIL-4 residual-risk items to the Board. Maturity advances from Level 1 (ad-hoc, no validated "
  "architecture) through Level 2 (validated architecture and core use cases, SL-T documented), Level 3 "
  "(three-echelon SOC operational, four-gate doctrine tested, category-aware detectors complete), Level 4 "
  "(national E3 SOC and Railway-sector ISAC, proactive threat hunting anchored to ATT&CK for ICS, "
  "automated conformance monitoring) to Level 5 (continuous architecture-to-detection synchronisation, "
  "leading-indicator driven). The realistic near-term target is Level 3 on pilot corridors, progressing to "
  "Level 4 nationally.")

# ---------------- 13 ----------------
h1("13. Conclusion")
P("The convergence that has modernised Indian Railways signalling has dissolved the isolation on which "
  "its security historically depended. Protecting this estate requires dedicated OT-SOCs that are "
  "passive-first, safety-aware, SIL-prioritised, protocol-literate and bound to a validated architecture. "
  "The contribution of this paper is the argument, evidenced by an independent standards-traceable "
  "model assessment, that architecture validation must precede OT-SOC deployment, together with a "
  "railway-specific three-echelon architecture, an explicit IEC 62443 / EN 50159 / TS 50701 integration in "
  "which each conduit's EN 50159 category fixes its detector set, and a federated operating model mapped "
  "to the Railway Board, RDSO, NCIIPC, CERT-In, Zonal Railways and Divisions. The recommendation is "
  "unambiguous: invest first in the validated architecture and the TS 50701 evidence base; adopt a "
  "federated hybrid model with explicit OT decision rights; instrument vital conduits and trust "
  "boundaries before broadening coverage; build the dual-competency workforce under RDSO "
  "certification; and run a 90-day Kavach-corridor pilot, validating the model-derived findings in the "
  "field. Isolation can no longer be assumed; it must now be observed.")

# ---------------- references ----------------
h1("References")
refs = [
 "IEC 62443-3-2, Security risk assessment for system design, IEC.",
 "IEC 62443-3-3, System security requirements and security levels, IEC.",
 "IEC 62443-2-1, Establishing an IACS security programme, IEC.",
 "CLC/TS 50701, Railway applications - Cybersecurity, CENELEC.",
 "EN 50159:2010, Safety-related communication in transmission systems, CENELEC.",
 "EN 50126, Railway applications - RAMS, CENELEC.",
 "EN 50129, Railway applications - Safety-related electronic systems for signalling, CENELEC.",
 "RDSO, Specification for Train Collision Avoidance System (Kavach), Ministry of Railways, Government of India.",
 "RDSO / Indian Railways, Kavach Cyber Security Standard Operating Procedure (SOP).",
 "Government of India, Information Technology Act 2000, s.70; NCIIPC, Guidelines for Protection of Critical Information Infrastructure.",
 "CERT-In, Directions under s.70B(6) of the IT Act, 28 April 2022 (six-hour incident reporting).",
 "NIST SP 800-82 Rev. 3, Guide to Operational Technology (OT) Security, NIST.",
 "MITRE ATT&CK for ICS, The MITRE Corporation.",
 "ENISA, Railway Cybersecurity / Transport Threat Landscape reports.",
 "IRSE, International Technical Committee guidance on cyber security for signalling systems.",
 "ShieldworkZ, OT SOC Foundational Guide - cross-sector OT-SOC operating-model, visibility, incident-response, KPI, maturity and residual-risk frameworks (adapted and railway-contextualised herein).",
 "Reference Railway OT Cyber Security Architecture Model and Independent Assessment - separated semantic, governance and enforcement authorities (primary evidence base).",
]
for i, r in enumerate(refs, 1):
    pp = d.add_paragraph(); pp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pp.paragraph_format.left_indent = Inches(0.28); pp.paragraph_format.first_line_indent = Inches(-0.28)
    rr = pp.add_run("[%d]  %s" % (i, r)); rr.font.size = Pt(9)

# ---------------- figure & table lists ----------------
h1("Figure List")
for t in ["Figure 1. IEC 62443 zone-and-conduit reference model for railway signalling (EN 50159 category annotated).",
          "Figure 2. Three-echelon OT-SOC topology with passive collection and unidirectional log export.",
          "Figure 3. Kavach communication architecture: legitimate end-to-end association vs forbidden direct edge."]:
    bullet(t)
h1("Table List")
for t in ["Table 1. OT threat categories mapped to Indian Railways.",
          "Table 2. Trust-boundary monitoring matrix.",
          "Table 3. OT-SOC operating-model evaluation and recommendation.",
          "Table 4. Five visibility layers applied to railway OT assets.",
          "Table 5. Railway OT-SOC detection use cases (ATT&CK for ICS mapping).",
          "Table 6. Railway OT incident severity model.",
          "Table 7. Governance, staffing and responsibility model.",
          "Table 8. Three-audience KPI framework."]:
    bullet(t)

d.save("IRSE_OT_SOC_Conference_Paper_5000W.docx")
print("SAVED IRSE_OT_SOC_Conference_Paper_5000W.docx")
