# -*- coding: utf-8 -*-
"""Generate IRSE_OT_SOC_Conference_Paper_Final.docx (python-docx)."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# ----- base styles -----
normal = doc.styles["Normal"]
normal.font.name = "Times New Roman"
normal.font.size = Pt(11)
for hs, sz in (("Heading 1", 15), ("Heading 2", 12.5), ("Heading 3", 11.5)):
    st = doc.styles[hs]
    st.font.name = "Times New Roman"
    st.font.size = Pt(sz)
    st.font.color.rgb = RGBColor(0x10, 0x2A, 0x43)

_tblno = [0]

def title(t):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(t)
    r.bold = True
    r.font.size = Pt(18)
    r.font.name = "Times New Roman"
    r.font.color.rgb = RGBColor(0x0B, 0x1F, 0x33)

def subtitle(t):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(t)
    r.italic = True
    r.font.size = Pt(11.5)

def byline(t):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(t)
    r.font.size = Pt(10.5)

def h1(t):
    doc.add_heading(t, level=1)

def h2(t):
    doc.add_heading(t, level=2)

def p(text, italic=False, bold=False):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = para.add_run(text)
    r.italic = italic
    r.bold = bold
    return para

def kw(label, text):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = para.add_run(label + " ")
    r.bold = True
    para.add_run(text)

def bullet(text):
    para = doc.add_paragraph(style="List Bullet")
    para.add_run(text)

def table(caption, headers, rows):
    _tblno[0] += 1
    cap = doc.add_paragraph()
    rc = cap.add_run("Table %d. %s" % (_tblno[0], caption))
    rc.bold = True
    rc.font.size = Pt(10)
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, htext in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(htext)
        run.bold = True
        run.font.size = Pt(9.5)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(val))
            run.font.size = Pt(9.5)
    doc.add_paragraph()

# =====================================================================
# FRONT MATTER
# =====================================================================
title("Cyber Security Framework Planning and Establishment of "
      "Operational Technology Security Operations Centres (OT-SOCs) "
      "for Indian Railways")
subtitle("An architecture-driven methodology integrating IEC 62443, EN 50159, "
         "TS 50701 and Kavach cyber security requirements with field-proven "
         "OT-SOC foundational practice")
byline("Prepared for the IRSE India Technical Convention  |  Audience: IRSSE Officers, RDSO, "
       "Railway Board, Signalling Engineers and OT Cyber Security Practitioners")
doc.add_paragraph()

h1("Abstract")
p("The nationwide deployment of the Train Collision Avoidance System (Kavach), electronic "
  "interlocking (EI), centralised Traffic Management Systems (TMS) and converged Multi-Protocol "
  "Label Switching (MPLS) telecom backbones has transformed Indian Railways signalling from a "
  "collection of electrically isolated vital circuits into a networked, safety-critical cyber-physical "
  "system. This convergence exposes Operational Technology (OT) to a threat surface that "
  "enterprise Information Technology (IT) Security Operations Centres (SOCs) are neither designed "
  "nor competent to monitor. This paper presents a structured, architecture-driven methodology for "
  "planning and establishing dedicated Operational Technology Security Operations Centres "
  "(OT-SOCs) for Indian Railways. Its central thesis is that a machine-validated, standards-traceable "
  "model of the signalling architecture - its IEC 62443 zones and conduits, its EN 50159 transmission "
  "categories and its Safety Integrity Level (SIL) allocations - must precede and continuously inform "
  "SOC operating-model selection, sensor placement, detection engineering, incident response and "
  "governance. The evidence base is an independent cyber security assessment of a reference railway "
  "OT architecture model in which the semantic, governance and enforcement authorities were "
  "formally separated and validated; the assessment surfaced durable engineering findings (absence "
  "of cryptographic authentication on vital Category-1 buses, absence of multi-factor authentication "
  "on privileged engineering access, absence of timeliness monitoring on open radio bearers, "
  "incompletely instrumented trust boundaries, and an unprotected telecom backhaul) together with "
  "standards-compliance findings (the EN 50159 / IEC 62443 authentication-crediting tension and the "
  "need for TS 50701 lifecycle evidence). The paper integrates this railway-specific evidence with "
  "field-proven OT-SOC foundational practice: an evaluation of five OT-SOC operating models with a "
  "recommendation for Indian Railways; a five-layer visibility architecture applied to Electronic "
  "Interlocking, Kavach, MPLS, radio networks, TMS, NMS, data loggers and engineering workstations; "
  "the Monitoring Paradox and its grounding in SIL-4 signalling, IEC 62443 and EN 50159; a four-gate, "
  "safety-first incident-response doctrine; a railway SEV1-SEV4 severity model; a three-tier "
  "(Board / SOC-Lead / Analyst) KPI framework; a railway OT-SOC staffing and competency model; a "
  "national governance model spanning Railway Board, RDSO, NCIIPC, CERT-In, Zonal Railways and "
  "Divisions; a residual-risk methodology with explicit risk-acceptance principles for SIL-4 systems; a "
  "five-level OT-SOC maturity roadmap; and a 90-day pilot-deployment programme. The paper is "
  "offered as publishable planning guidance for the establishment of OT-SOC capability across "
  "Indian Railways.")

kw("Keywords:", "OT-SOC; Indian Railways; Kavach; IEC 62443; EN 50159; TS 50701; Electronic "
   "Interlocking; Zones and Conduits; Detection Engineering; SIL-4; Incident Response; Monitoring "
   "Paradox; Residual Risk; SOC Maturity; NCIIPC; CERT-In; RDSO.")

# =====================================================================
h1("1. Introduction")
p("For most of its history, the safety of railway signalling rested upon two principles: physical "
  "isolation and fail-safe design. Vital functions were realised in relay logic, hard-wired track "
  "circuits and electrically isolated lineside cabling; the cyber threat surface was, in practice, the "
  "perimeter fence. The doctrine of security through isolation was tenable because the signalling "
  "system genuinely was isolated.")
p("That premise no longer holds. The contemporary Indian Railways signalling estate is a converged, "
  "IP-enabled, geographically distributed cyber-physical system. Electronic interlockings exchange "
  "vital state with object controllers over digital buses; the Kavach Automatic Train Protection (ATP) "
  "system extends vital communication across an open radio-frequency (RF) air interface between the "
  "Stationary Kavach (S-Kavach) at the wayside and the Loco Kavach (L-Kavach) onboard; TMS and "
  "Network Management Systems (NMS) supervise wide geographies over MPLS backbones; and "
  "engineering workstations capable of downloading vital logic connect, through jump hosts and "
  "remote-access infrastructure, into the heart of the safety domain. Each of these capabilities is a "
  "deliberate engineering advance; each also introduces a conduit through which a determined "
  "adversary may attempt to influence a safety function.")
p("The instinctive institutional response - to extend the existing enterprise IT SOC to also cover "
  "signalling - is inadequate and, in places, dangerous. Cross-sector OT-SOC experience is "
  "unambiguous: extending an IT SOC into OT consistently fails because the tools, telemetry sources, "
  "threat models, response playbooks and human expertise required are fundamentally different, "
  "producing a false sense of security - monitoring coverage that looks credible on paper but is blind "
  "to the threats that matter most. An IT SOC optimises for confidentiality and the protection of data; "
  "an OT-SOC must optimise for the preservation of safety and the availability of a physical process. "
  "An IT SOC may legitimately quarantine an endpoint or block a flow on detection; an OT-SOC that "
  "interrupts a vital communication path may itself precipitate the hazard it was deployed to prevent.")
p("This paper argues that the establishment of OT-SOC capability for Indian Railways must be a "
  "discipline in its own right, and that its single most important precondition is an architecture that "
  "has been validated against standards before it is monitored. One cannot monitor what one has not "
  "modelled; one cannot prioritise detection on a flow whose criticality, transmission category and "
  "trust boundary are unknown; and one cannot defend a zone whose conduits have not been "
  "enumerated. The argument is developed from concrete evidence - an independent, "
  "standards-traceable assessment of a reference railway OT architecture model - and fused with "
  "field-proven OT-SOC foundational practice, so that Indian Railways need not re-derive at cost what "
  "other critical-infrastructure sectors have already learned.")

# =====================================================================
h1("2. Evolution of Cyber Security in Indian Railways")
p("Four phases characterise the trajectory of railway signalling cyber security in India. Phase I - "
  "electromechanical isolation - comprised relay interlockings and token working; security was "
  "physical and procedural and no cyber threat surface of consequence existed. Phase II - "
  "first-generation electronics - introduced solid-state and electronic interlockings, data loggers and "
  "digital axle counters; digital logic entered the vital domain but communication remained local, "
  "proprietary and serial. Data loggers, installed for diagnostics, became almost incidentally the first "
  "forensic record of what would later be recognised as security events. Phase III - networked OT and "
  "the rise of the backbone - brought MPLS backbones, centralised TMS, NMS and remote "
  "diagnostics, the first genuine IT/OT convergence; engineering access, previously requiring physical "
  "presence in the relay room, became remotely brokered, and the threat surface expanded from the "
  "fence to the network.")
p("Phase IV - safety over open bearers - is the present, and its defining event is the nationwide "
  "rollout of Kavach. Kavach deliberately carries vital movement-authority information across an open "
  "RF bearer between wayside and locomotive and integrates with the electronic interlocking at the "
  "station. For the first time, a vital, SIL-4 function depends upon communication over a transmission "
  "system that, in EN 50159 terms, is open - one in which the population of participants and the "
  "medium itself cannot be fully controlled. The RDSO Kavach specification and the associated Kavach "
  "Cyber Security Standard Operating Procedure (SOP), together with Indian Railways' signalling cyber "
  "security policy direction, recognise that cryptographic protection of the Kavach interface is "
  "necessary but not sufficient: protection without monitoring provides no assurance that the "
  "protection is functioning, no detection of attempts to defeat it, and no evidence for incident "
  "response. The OT-SOC is the institutional mechanism through which protection is continuously "
  "verified, and its establishment is therefore no longer optional but a regulatory and safety "
  "imperative.")

# =====================================================================
h1("3. The Railway OT Threat Landscape")
p("The railway OT threat landscape differs from the enterprise landscape in its consequences, its "
  "adversaries and its constraints. Two decades of cross-sector incident response show that OT "
  "threats cluster into five operationally distinct categories; Table 1 maps these to their railway "
  "manifestations and to the consequence that matters in signalling - the wrong-side failure.")
table("Cross-sector OT threat categories mapped to Indian Railways",
      ["#", "Generic OT threat category", "Indian Railways manifestation", "Consequence"],
      [["1", "IT-sourced lateral movement (ransomware crossing a weak IT/OT boundary; Colonial Pipeline class)",
        "Enterprise to IDMZ to operations to interlocking traversal where boundary controls are incompletely instantiated",
        "Disruption of TMS/NMS; staging toward vital zones"],
       ["2", "Engineering-workstation compromise (TRITON/TRISIS modified safety logic, 2017)",
        "Compromise of an engineering workstation with logic-download access to EI / S-Kavach",
        "Alteration of vital interlocking logic - wrong-side failure"],
       ["3", "Supply-chain / vendor-access exploitation (SolarWinds scale)",
        "Vendor laptop or software-update channel for Kavach / EI / TMS subsystems",
        "Implant introduced into the vital supply chain"],
       ["4", "Insider threat and sabotage (Maroochy Water, 2000)",
        "Maintainer or contractor with privileged engineering access and protocol knowledge",
        "Setpoint/logic change; alarm suppression on a vital system"],
       ["5", "Nation-state pre-positioning and destructive attack (INDUSTROYER, TRITON, PIPEDREAM)",
        "Persistent access to Kavach/EI/telecom infrastructure awaiting operational timing",
        "Coordinated destructive disruption of national rail safety"]])
p("The worst-case outcome in railway OT is not data loss but a wrong-side failure: a signal cleared "
  "against a conflicting movement, a point moved under a train, a movement authority issued where "
  "none is safe, or the suppression of an emergency brake command. The safety-criticality of the "
  "affected asset - captured formally as its SIL - must therefore be the primary axis along which "
  "monitoring effort is prioritised. The relevant adversary spectrum runs from the opportunistic "
  "(commodity malware introduced via an engineering laptop or USB device) through the capable "
  "insider (a maintainer with privileged access and detailed protocol knowledge) to the resourced "
  "nation-state actor. The capable insider is of particular concern because the vital buses rely, by "
  "design and by standard, on closed-transmission assumptions and safety coding rather than "
  "cryptographic source authentication - a defence robust against the outsider but materially weaker "
  "against an actor who has obtained access to the closed segment. Critically, the majority of OT "
  "incidents that require SOC response begin with human error or supply-chain compromise, not "
  "sophisticated attack: a well-designed railway OT-SOC must be equally capable of detecting "
  "configuration drift, accidental connectivity changes and unauthorised removable media as of "
  "hunting nation-state implants.")

# =====================================================================
h1("4. Standards Framework: IEC 62443, EN 50159, TS 50701 and Kavach")
p("Four standards and policy families provide the structural grammar of railway OT security, and the "
  "OT-SOC must be built upon all of them simultaneously.")
h2("4.1 IEC 62443 - Zones, conduits and security levels")
p("IEC 62443-3-2 (zone and conduit risk assessment) and IEC 62443-3-3 (system security requirements "
  "and security levels) decompose a system into zones - groupings of assets sharing common security "
  "requirements - connected by conduits, the controlled communication paths between zones. Its "
  "seven Foundational Requirements - FR1 Identification and Authentication Control, FR2 Use Control, "
  "FR3 System Integrity, FR4 Data Confidentiality, FR5 Restricted Data Flow, FR6 Timely Response to "
  "Events, and FR7 Resource Availability - give a complete vocabulary for security controls. Crucially "
  "for the OT-SOC, FR6 is itself a monitoring mandate. The reference architecture realises the IEC "
  "62443 model as an explicit zone hierarchy - Enterprise (L5), IDMZ and Security-Management (L3.5), "
  "Operations (L3), Telecom Core (L2), Radio Access (L1), Interlocking (L2/L1), Field (L0) and Onboard - "
  "and as a set of typed conduits, each carrying a security profile.")
h2("4.2 EN 50159 - Safety-related communication and transmission categories")
p("EN 50159 governs safety-related communication over transmission systems and classifies them into "
  "three categories. Category 1 (closed) systems have a fixed number of participants, fixed "
  "transmission characteristics and no unauthorised access to the medium. Category 2 (controlled) "
  "systems are not under full control but the threats are partly mitigated by design. Category 3 (open) "
  "systems have an unknown number of participants and an unknown medium, and the possibility of "
  "unauthorised access must be assumed. The category dictates the required defences against the "
  "canonical EN 50159 threats - repetition, deletion, insertion, re-sequencing, corruption, delay and "
  "masquerade - and therefore directly determines what an OT-SOC must monitor: a Category-1 vital "
  "bus demands integrity and sequence monitoring; a Category-3 open RF bearer additionally demands "
  "timeliness, replay, masquerade and RF-anomaly monitoring.")
h2("4.3 TS 50701 - Railway cyber security lifecycle")
p("CLC/TS 50701 adapts the IEC 62443 approach to the railway domain and binds it to the EN 5012x "
  "safety lifecycle (EN 50126 RAMS, EN 50129 safety case). It requires a documented risk assessment, "
  "the allocation of target Security Levels (SL-T) to zones and conduits, and the maintenance of a "
  "cyber security case analogous to the safety case. For the OT-SOC, the SL-T per zone is the input "
  "that sets monitoring intensity, alert prioritisation and response-time targets; TS 50701 is therefore "
  "the standard that makes monitoring an auditable lifecycle obligation and supplies the calibration "
  "data for the SOC.")
h2("4.4 Kavach cyber security requirements")
p("The RDSO Kavach specification and the Kavach Cyber Security SOP impose railway-specific "
  "obligations: cryptographic protection of the Kavach radio interface, integrity and replay protection "
  "of vital telegrams, controlled key management, and - of direct relevance here - the continuous "
  "monitoring and incident-reporting obligations that an OT-SOC discharges. The interlock between "
  "the four families is the conceptual core of this paper: IEC 62443 tells the OT-SOC where the "
  "boundaries are; EN 50159 tells it what to look for on each conduit; TS 50701 tells it why, to what "
  "assurance target and with what evidence; and the Kavach SOP tells it the railway-specific "
  "obligations it must satisfy and report against.")

# =====================================================================
h1("5. The Architecture-Driven Imperative: Validate Before You Monitor")
p("The central methodological proposition of this paper is that architecture validation must precede "
  "OT-SOC deployment. A SOC detects deviation from expected behaviour, and the definition of "
  "expected is the architecture. If the architecture is not modelled, the SOC's baseline is whatever "
  "traffic happened to exist when monitoring began - including any pre-existing compromise or "
  "misconfiguration. In the reference assessment, the architecture was expressed as three separated "
  "authorities: a semantic authority (the ontology of asset types, protocols, conduit classes, Purdue "
  "levels and protocol capabilities), a governance authority (the zoning, trust boundaries, flow rules, "
  "conduit security profiles and EN 50159 transmission categories) and an enforcement authority (the "
  "validators that consume the first two and emit findings). This separation is precisely what allows a "
  "deviation to be classified unambiguously as a real security gap, a standards-compliance gap, a "
  "modelling defect or an enforcement defect.")
p("The assessment demonstrated this concretely. Before validation, several legitimate architectural "
  "flows - operator workstation to TMS, the network-management plane, the telecom backhaul, and the "
  "Loco-to-Stationary Kavach safety association - were unmodelled, and a vital-to-open-RF flow that "
  "should never exist (S-Kavach directly cabled to the radio base station) was present in the topology "
  "but indistinguishable from the legitimate ones. Only after the governance model was completed "
  "could the forbidden flow be classified as forbidden and prevented, and the legitimate flows be "
  "classified as authorised-but-monitored. An OT-SOC fed by the pre-validation model would have had "
  "no basis on which to alarm on the forbidden edge. The validated architecture is thus not "
  "documentation; it is the configuration source for every sensor, detector and boundary control the "
  "OT-SOC will deploy. The residual findings of that assessment - referenced throughout this paper as "
  "evidence and consolidated in the residual-risk register of Section 16 - become the OT-SOC's "
  "standing monitoring requirements rather than items to be quietly closed.")

# =====================================================================
h1("6. OT-SOC Operating Models: Evaluation and Recommendation for Indian Railways")
p("There is no single correct OT-SOC operating model; the right choice depends on scale, geographic "
  "footprint, regulatory obligation, existing IT-security capability and the operational maturity of the "
  "OT environment. Five archetypes cover the realistic spectrum. Table 2 evaluates each against the "
  "specific circumstances of Indian Railways - a national operator with thousands of geographically "
  "dispersed wayside and Kavach sites, a tiered zonal/divisional administrative structure, and "
  "safety-critical SIL-4 assets.")
table("Evaluation of the five OT-SOC operating models for Indian Railways",
      ["Model", "Description", "Strengths", "Suitability for Indian Railways"],
      [["A - Dedicated OT-SOC (greenfield)",
        "Physically separate SIEM, analyst workstations and network, OT-native or cross-trained staff, independent stack",
        "Highest detection fidelity; full OT focus",
        "Appropriate at NATIONAL (Tier-3) level under RDSO / Railway Board for the vital estate; high cost justified by SIL-4 criticality"],
       ["B - Fused IT/OT SOC with dedicated OT tier",
        "Shared SIEM/SOAR/ticketing but a ring-fenced OT tier: separate rules, queues, playbooks and at least one OT analyst per shift",
        "Strong capability at manageable cost; reuses existing IT-SOC investment",
        "Appropriate at ZONAL/DIVISIONAL (Tier-2) level; requires a charter giving the OT tier authority to veto IT-SOC actions affecting OT"],
       ["C - MSSP-led remote monitoring",
        "Lean internal team (2-4 FTE) for governance and escalation; day-to-day monitoring outsourced to a specialist OT MSSP",
        "Fast to stand up; compensates for talent scarcity",
        "Transitional only; institutional knowledge must be contractually transferred - unsuitable as the steady state for vital safety"],
       ["D - Asset-owner collaborative (sector-ISAC aligned)",
        "Multiple asset owners contribute telemetry to a sector-level capability; shared L2/L3 threat intelligence",
        "Excellent cross-network threat detection and intelligence",
        "Strongly relevant as a Railway-sector ISAC function, RDSO-anchored and CERT-In / NCIIPC aligned"],
       ["E - Virtual / distributed OT-SOC",
        "Lightweight sensors at each site aggregating to a central analysis platform; analysts across multiple locations",
        "Practical coverage of highly distributed estates",
        "Essential pattern for the thousands of distributed wayside / Kavach / interlocking sites"]])
p("Recommendation - a federated hybrid. No single archetype fits a national railway. The "
  "recommended model for Indian Railways is a federated hybrid that combines the strengths of four "
  "archetypes. A dedicated national OT-SOC (Model A) at Tier-3, operated under RDSO and Railway "
  "Board governance, holds the validated architecture, the SL-T allocation, the cyber security case, "
  "national threat intelligence and a Railway-sector ISAC function (Model D). Divisional and zonal OT "
  "tiers (Model B), fused with existing IT-SOC infrastructure but with explicit authority to override IT "
  "actions affecting signalling, provide regional correlation and case management. A distributed "
  "sensing layer (Model E) instruments the geographically dispersed wayside, interlocking and Kavach "
  "sites with passive collectors that aggregate upward. Model C (MSSP) is admissible only as a "
  "time-boxed transitional measure to bridge the talent pipeline, with mandatory knowledge-transfer "
  "obligations. Across all tiers, the single non-negotiable governance control is a charter defining "
  "decision rights: the OT tier must be able to block any proposed action - including from IT-SOC staff "
  "- that would affect a vital system.")

# =====================================================================
h1("7. Visibility Architecture: The Five Layers Applied to Railway Assets")
p("Comprehensive OT-SOC visibility requires collecting telemetry from five distinct layers. Many "
  "programmes focus exclusively on the network layer and miss the process and identity layers that "
  "are most relevant to OT-specific threats. The five layers are: the Network layer (L2/L3 switches, "
  "routers, firewalls and IDS sensors); the Asset/Endpoint layer (engineering and operator "
  "workstations, jump hosts); the Application/Supervisory layer (TMS, NMS, SOC servers); the "
  "Process/Vital-state layer (EI vital state, data loggers, onboard event recorders); and the "
  "Identity and Access layer (IDMZ directory services, remote-access gateways, vendor jump servers). "
  "Table 3 applies these layers concretely to the railway assets named in the assessment.")
table("Five visibility layers applied to Indian Railways OT assets",
      ["Railway asset", "Dominant visibility layer(s)", "Collection method", "Key railway telemetry"],
      [["Electronic Interlocking (EI)", "Process/Vital-state; Network",
        "Passive SPAN/TAP on EI vital buses; read-only vital-state observation",
        "Route/occupancy state, RaSTA/Kavach-interface reads-writes, safety-code failures, state-machine violations"],
       ["Kavach (S-Kavach / L-Kavach)", "Network; Process/Vital-state",
        "Passive TAP plus RF monitor (wayside); onboard event recorder (loco)",
        "Movement-authority frames, MCOMM integrity, sequence/freshness, brake-command integrity, RF anomaly"],
       ["MPLS / PE routers (backbone)", "Network",
        "Passive flow monitoring on the telecom core conduit",
        "Transport-security (IPsec/MACsec) tunnel health, routing anomaly, unexpected participants"],
       ["Radio networks (BTS / radio gateway)", "Network; Process/Vital-state",
        "RF spectrum monitor plus conduit TAP",
        "Signal-strength and error-rate anomaly (jamming), timeliness/latency deviation, replay indicators"],
       ["Traffic Management System (TMS)", "Application/Supervisory",
        "Syslog/TLS forwarding; HTTPS session monitoring",
        "Operator route/setpoint commands, supervisory-command integrity, failed logins, alarm-state changes"],
       ["Network Management System (NMS)", "Application/Supervisory; Identity",
        "SNMPv3 event capture; management-plane monitoring",
        "Unexpected SNMP set operations, configuration changes, off-hours management access"],
       ["Data loggers", "Process/Vital-state",
        "Read-only log ingestion (never inline)",
        "Forensic event stream; tamper indicators; time-synchronisation drift; gap/reorder detection"],
       ["Engineering workstations (EWS)", "Asset/Endpoint; Identity",
        "Endpoint agent where supported, else agentless read-only; PAM session recording",
        "Logic-download and configuration-change actions, USB insertion, anomalous/privileged logins"]])
p("The mapping makes two planning points concrete. First, the highest-consequence telemetry - the "
  "process/vital-state layer for EI and Kavach - is precisely the telemetry that is hardest to obtain, "
  "because vital devices emit few logs and cannot be actively interrogated; it must be derived from "
  "passive network observation and from the data loggers and onboard recorders. Second, the identity "
  "layer (engineering and vendor access) is the layer most often neglected yet most directly tied to the "
  "highest-value attack path, the engineering workstation. A railway OT-SOC that instruments only the "
  "network layer will see the backbone but be blind to a logic download in progress on the EI.")

# =====================================================================
h1("8. The Monitoring Paradox: Why Railway OT Monitoring Must Remain Passive")
p("Every piece of monitoring infrastructure deployed inside the OT network is itself an attack surface. "
  "A compromised sensor or monitoring agent sitting on a vital network is a threat actor's bridge into "
  "the most sensitive environment a railway operates. This is the Monitoring Paradox, and in the "
  "railway context it has a sharper edge than in any other sector, because the monitored assets are "
  "SIL-4 vital systems whose disturbance can directly cause a wrong-side failure.")
p("Three consequences follow. First, monitoring must be passive by default. Active scanning of a vital "
  "network - port scans, authenticated polling, vulnerability probes - can reset or destabilise "
  "controllers and disrupt time-critical communication; it must require explicit signalling-engineering "
  "sign-off for each target and must never be the default mode on the interlocking, field or onboard "
  "zones. Telemetry is therefore obtained predominantly through SPAN/TAP and data-diode export. "
  "Second, the SOC must remain outside the vital network: the footprint inside the interlocking and "
  "onboard zones is minimised, and aggregation and analysis are performed in higher, less-trusted "
  "tiers. Third, monitoring must not interfere with the EN 50159 safety layer or the EN 50129 safety "
  "case: any device inserted into a vital path could, in principle, introduce delay or corruption, which "
  "are themselves EN 50159 threats - so the monitoring apparatus must be demonstrably incapable of "
  "affecting the vital communication it observes.")
p("These constraints are not merely operational preferences; they are grounded in the standards. IEC "
  "62443 FR7 (Resource Availability) and the OT priority order (Safety, then Availability, then Integrity, "
  "then Confidentiality) forbid security mechanisms that degrade the process. EN 50159's treatment of "
  "delay and corruption as first-class threats means that an in-line monitoring device on a SIL-4 bus "
  "would be a potential source of the very hazards the safety layer is designed to exclude. The "
  "Monitoring Paradox, properly understood, is therefore the railway expression of the principle that in "
  "OT the cure must never be more dangerous than the disease.")

# =====================================================================
h1("9. Three-Tier OT-SOC Architecture")
p("The federated hybrid recommended in Section 6 is realised as a three-tier architecture organised "
  "by geography, latency and function, mirroring both the physical distribution of the railway and the "
  "IEC 62443 trust hierarchy. Tier 1 (Station / Wayside) is co-located with the relay or equipment "
  "room and comprises passive sensors on the field and interlocking conduits, SPAN/TAP feeds from "
  "the EI and object-controller buses, data-logger ingestion and a unidirectional (data-diode) log "
  "export toward the upper tiers; it performs local detection and buffers evidence so that connectivity "
  "loss does not lose forensic data, and it never hosts an in-line control on a vital path. Tier 2 "
  "(Divisional / Zonal OT-SOC) hosts the SIEM, correlation and analytics, intrusion sensors on the "
  "IT/OT boundary operating in detection mode toward OT, vulnerability management for non-vital "
  "assets and case management; it correlates events across stations in its geography - for example a "
  "corridor-wide pattern of balise-telegram anomalies that no single station would detect - and "
  "interfaces with divisional signalling control for response. Tier 3 (Central / National OT-SOC) is the "
  "custodian of the validated architecture and detection-content repository, the SL-T and cyber "
  "security case, cross-zone correlation, national threat intelligence, the Railway-sector ISAC, KPI "
  "consolidation and regulatory reporting; when the architecture changes - a new Kavach section, a "
  "re-zoned IDMZ - Tier 3 re-validates and pushes updated detection baselines downward. Log flow is "
  "strictly upward and, across the IT/OT boundary, diode-enforced; control flow (sensor configuration) "
  "is downward and authenticated, preserving the FR5 restricted-data-flow principle within the SOC's "
  "own architecture and respecting the Monitoring Paradox.")

# =====================================================================
h1("10. Detection Engineering and Railway-Specific Use Cases")
p("Detection engineering converts the validated architecture and its findings into deployable content. "
  "The governing principle is that every conduit security requirement and every assessment finding "
  "becomes a detection use case: a control that the architecture requires but the traffic may lack is "
  "precisely the condition the SOC must detect. The reference assessment surfaced findings that are "
  "restated here as durable monitoring requirements. The vital Category-1 buses - EI to S-Kavach over "
  "the Kavach EI interface, EI to object controller and speed-sensor to L-Kavach over CAN, and "
  "L-Kavach to driver-machine interface and brake-interface unit over MVB - provide integrity and "
  "replay protection through safety coding but not cryptographic source authentication. The SOC "
  "cannot add authentication, but it can and must detect the symptoms its absence would allow: "
  "unexpected source identifiers, protocol state-machine violations and message rates or sequences "
  "inconsistent with the vital application. Privileged engineering access is not, in the modelled flows, "
  "gated by multi-factor authentication, so every privileged session becomes a high-value detection "
  "target. The open RF safety conduits require timeliness, jamming and RF-anomaly monitoring; "
  "several trust-boundary crossings are not fully instrumented; and the telecom backhaul carries "
  "operational traffic without an evidenced transport-security overlay. Table 4 maps the EN 50159 and "
  "IEC 62443 threats to railway detections, and Table 5 expresses the deployable use-case library.")
table("Threat-to-detection mapping (EN 50159 / IEC 62443)",
      ["Threat (standard reference)", "Railway manifestation", "Detection technique", "Sensor / tier"],
      [["Masquerade / insertion (FR1)", "Spoofed axle-counter-evaluator or object-controller message",
        "Source-ID whitelist; state-machine validation; rate anomaly", "Interlocking TAP / T1"],
       ["Corruption (FR3; EN 50159)", "Manipulated EI-to-Kavach telegram",
        "Safety-code/CRC failure rate; integrity-flag monitoring", "EI/Kavach TAP / T1"],
       ["Repetition / replay (EN 50159)", "Replayed Kavach RF movement-authority frame",
        "Sequence-number regression; freshness deviation", "Radio conduit, recorder / T1-T2"],
       ["Delay / timeliness (FR6; EN 50159)", "Stale movement authority over RF",
        "End-to-end latency vs SIL-bounded threshold", "Radio conduit / T1-T2"],
       ["Jamming / availability (FR7)", "RF denial against L-Kavach to BTS",
        "RF signal-strength / error-rate anomaly; loss-of-comms", "RF monitor / T1"],
       ["Privileged misuse (FR2)", "Unauthorised logic download from EWS",
        "Engineering-session and configuration-change detection; MFA-absence alert", "Ops-to-interlocking / T2"],
       ["Lateral movement (FR5)", "Enterprise to operations traversal",
        "Cross-boundary flow correlation; IT-alert time correlation", "Trust-boundary sensors / T2"],
       ["Backhaul tampering (FR3/FR4)", "Injection on unprotected MPLS/IP core",
        "Transport-security health; unexpected-participant detection", "Telecom conduit / T2"],
       ["Forbidden topology (FR5)", "Direct vital-to-open-RF edge (S-Kavach to BTS)",
        "Topology-conformance vs validated model", "Conformance engine / T3"]])
table("Railway-specific OT-SOC detection use-case library",
      ["ID", "Use case", "Trigger", "Severity", "Safety-aware response"],
      [["UC-01", "Spoofed vital field message", "Source-ID not in EI whitelist / state-machine violation",
        "SEV1", "Alert control; do not block; manual verification under safe-working rules"],
       ["UC-02", "Kavach EI-interface integrity loss", "Rising CRC / safety-code failure", "SEV1",
        "Alert; correlate EI health; engineering investigation"],
       ["UC-03", "Kavach RF replay / delay", "Sequence regression / latency beyond SIL threshold", "SEV1",
        "Alert; verify RF environment; no automated RF action"],
       ["UC-04", "RF jamming / loss of comms", "Signal-strength collapse / error-rate spike", "SEV2",
        "Alert; invoke degraded-mode working procedures"],
       ["UC-05", "Unauthorised logic download", "Logic/config change from EWS toward EI/S-Kavach", "SEV1",
        "Alert; freeze change control; verify authorisation"],
       ["UC-06", "Privileged access without MFA", "Jump-host to workstation session, no MFA event", "SEV2",
        "Alert; challenge session; review access policy"],
       ["UC-07", "Trust-boundary bypass", "Flow crossing enterprise to operations without required controls", "SEV2",
        "Alert; correlate lateral movement"],
       ["UC-08", "Forbidden topology edge", "S-Kavach to BTS (or any forbidden flow) detected", "SEV1",
        "Alert; emergency architecture review"],
       ["UC-09", "Backhaul transport-security failure", "Loss/absence of IPsec/MACsec on TMS-MPLS path", "SEV3",
        "Alert; verify tunnel health"],
       ["UC-10", "Balise integrity / position anomaly", "Telegram safety-code failure / position implausibility", "SEV2",
        "Alert; NO crypto-encryption alarm (EN 50159 passive-telegram exemption)"],
       ["UC-11", "Data-logger tampering / time drift", "Gap, reorder or clock skew in logger stream", "SEV3",
        "Alert; preserve forensic chain"],
       ["UC-12", "Management-plane anomaly", "Unexpected SNMPv3 set on NMS to elements", "SEV3",
        "Alert; verify change authorisation"],
       ["UC-13", "New device on vital segment", "MAC/IP not in asset inventory on interlocking zone", "SEV2",
        "Alert; verify closed-medium assumption (links to UC-01)"],
       ["UC-14", "Removable media on EWS", "USB device-ID not in approved list", "SEV3",
        "Alert; isolate offline EWS per procedure"]])
p("A deliberately exempt condition illustrates why the validated architecture matters. The passive "
  "balise telegram - the trackside RFID/EUROBALISE air-gap and the reader-to-L-Kavach leg - is "
  "protected by the passive safety-coded telegram, with freshness derived from location coding and "
  "onboard odometry, not by cryptographic encryption or replay control (EN 50159). The OT-SOC must "
  "therefore not raise a missing-encryption or missing-replay alarm on the passive telegram, which "
  "would be a false positive (UC-10), while still monitoring telegram integrity and positional "
  "plausibility. Encoding standards-justified exemptions as carefully as detections is a discipline that "
  "only a validated, authority-separated architecture makes reliable.")
table("Trust-boundary monitoring matrix",
      ["Trust boundary", "EN 50159 relevance", "Required controls", "OT-SOC monitoring obligation"],
      [["Enterprise (L5) - IDMZ (L3.5)", "IT/OT divide", "Firewall, inspection",
        "Verify controls present; detect unauthorised ingress; diode log export"],
       ["IDMZ (L3.5) - Operations (L3)", "-", "Firewall, inspection",
        "Brokered-access audit; jump-host session and MFA-event verification"],
       ["Operations (L3) - Interlocking (L2)", "Cat 1 downstream", "Firewall, inspection, integrity, authentication",
        "Highest scrutiny; engineering-access and logic-download detection"],
       ["Telecom Core (L2) - Radio Access (L1)", "Cat 2", "Firewall, inspection, authentication",
        "Transport-security health; participant verification"],
       ["Radio Access (L1) - Onboard", "Cat 3 (open RF)", "Integrity, replay, authentication, radio monitoring",
        "Full Cat-3 set: timeliness, jamming, RF anomaly, replay"],
       ["Balise air-gap (Field - Onboard)", "Cat 3 passive", "Integrity (passive telegram); crypto exempt",
        "Telegram integrity and positional plausibility; no crypto false positives"]])

# =====================================================================
h1("11. Incident Response Doctrine: The Four-Gate, Safety-First Model")
p("Railway OT incident response inverts the core IT assumption that speed is paramount. IT doctrine is "
  "built around speed - contain fast, eradicate fast, recover fast. OT doctrine is built around safety: "
  "assess safety impact before every action, contain only when containment does not create a more "
  "dangerous state, and recover in coordination with Operations and Engineering. A four-gate model "
  "governs every railway OT incident-response action, and no gate permits an action that interrupts a "
  "vital path or induces a wrong-side state. The response repertoire is graduated - observe and "
  "enrich, then alert the controller, then recommend a manual, procedurally-governed protective "
  "action - the last always executed by signalling operations under existing safe-working rules "
  "(caution orders, degraded-mode working, temporary speed restrictions), never automatically by the "
  "SOC tool.")
table("Four-gate OT incident-response doctrine adapted for railway operations",
      ["Gate", "Question to answer", "Decision authority (railway)"],
      [["G1 - Safety assessment",
        "Does the incident, or any proposed response, create or worsen a wrong-side or physical-safety risk?",
        "Signalling / Functional-Safety engineer - mandatory; SOC cannot proceed to containment without a documented safety assessment"],
       ["G2 - Operational-impact assessment",
        "What is the consequence of the incident continuing versus of the proposed response on live train operations?",
        "Divisional Control / Station Master jointly with SOC Tier-3"],
       ["G3 - Containment decision",
        "What is the minimum viable action that halts threat progression without unacceptable operational impact, and is it reversible?",
        "OT-SOC Lead and Operations Lead - recorded in the incident log"],
       ["G4 - Recovery authorisation",
        "Is the threat remediated, the architecture segment re-validated against the model, and the system safe to return to service?",
        "RDSO / Divisional safety authority - formal return-to-service sign-off"]])
p("Decision rights are explicit: the OT-SOC owns detection, enrichment and recommendation; signalling "
  "operations owns any action affecting train movement; and RDSO / Railway Board governance owns "
  "escalation and external reporting, including the statutory incident notification to CERT-In and the "
  "advisories from NCIIPC. Forensic readiness is a precondition, not an afterthought: data loggers and "
  "onboard event recorders must have assured integrity and time-synchronisation before an incident, "
  "so that evidence is admissible afterward.")

# =====================================================================
h1("12. Railway Incident Severity Model (SEV1-SEV4)")
p("Severity in OT must be defined by potential physical consequence, not data value. The following "
  "railway-specific scheme anchors severity to the wrong-side failure and to the SIL of the affected "
  "asset, and it drives the response SLAs and the escalation chain of Section 14.")
table("Railway OT incident severity classification",
      ["Level / name", "Definition (consequence-based)", "Railway examples", "Response SLA and escalation"],
      [["SEV1 - SAFETY-CRITICAL",
        "Active threat to physical safety; vital / Kavach / EI compromise confirmed or highly probable; potential wrong-side failure",
        "Manipulation of EI vital logic; spoofed Kavach movement authority; attack on a SIL-4 vital controller",
        "Immediate all-hands; G1 safety assessment invoked; executive and functional-safety notification within 15 min; CERT-In / NCIIPC notification within statutory window"],
       ["SEV2 - OPERATION-CRITICAL",
        "Confirmed OT compromise with capability to disrupt train operations; active lateral movement in OT; containment will affect services",
        "Ransomware on TMS or historian; compromised EWS with live EI/Kavach connection; C2 from an operations server",
        "Tier-3 and SOC-Lead escalation within 15 min; Operations within 30 min; CISO/IRSSE within 1 hour"],
       ["SEV3 - OT-ZONE COMPROMISE",
        "Confirmed malicious presence in OT, not yet demonstrating intent or capability to affect the vital process",
        "Malware on an isolated EWS; phished signalling engineer; rogue device on the OT network",
        "One-hour investigation; Operations informed; CISO notified same business day"],
       ["SEV4 - ANOMALY / POLICY",
        "Policy violation or anomaly with OT scope and no confirmed compromise",
        "Unauthorised removable media; out-of-window vendor access; configuration drift",
        "Next-business-day triage; trend-tracked for tuning and hunting"]])

# =====================================================================
h1("13. KPI Framework: Board, SOC-Lead and Analyst Tiers")
p("OT-SOC performance must be measured by safety-weighted, standards-anchored indicators rather "
  "than enterprise-IT volume metrics; metrics are only valuable if they drive decisions. The framework "
  "distinguishes three classes of metric - leading indicators (predictive of future posture), lagging "
  "indicators (measuring past performance) and operational-health indicators (measuring SOC "
  "sustainability) - and organises them into three reporting tiers, each railway-anchored.")
table("Three-tier KPI framework adapted for Indian Railways",
      ["Tier / KPI", "Target", "Class", "Railway meaning and standards anchor"],
      [["Board: Vital-asset monitoring coverage", "approaching 100% of SIL-3/4 conduits", "Leading",
        "Cannot detect what is unseen; FR6, TS 50701"],
       ["Board: Detection coverage by use case", "> 95% active", "Leading",
        "Defined use cases with validated, periodically-tested rules"],
       ["Board: Mean Time to Detect (vital)", "< 1 hr known TTP; < 4 hr", "Lagging",
        "Adversary action to first alert; EN 50129 / FR6"],
       ["Board: Mean Time to Safe Advisory", "minimise (safety-first)", "Lagging",
        "Detection to actionable advisory to signalling control; Kavach SOP"],
       ["Board: Incidents contained without operational impact", "> 90%", "Lagging",
        "Contained without unplanned service stop or safety activation"],
       ["Board: Regulatory-evidence availability", "100% on demand", "Op-health",
        "Log retention, access records, incident reports; TS 50701, CERT-In"],
       ["SOC-Lead: Trust-boundary instrumentation", "approaching 100% of governance-defined boundaries", "Leading",
        "FR5; verify, do not assume"],
       ["SOC-Lead: EN 50159 detector completeness", "approaching 100%", "Leading",
        "Deployed detectors match each conduit's transmission category"],
       ["SOC-Lead: Sensor / collector uptime", "> 99.5%", "Op-health",
        "Each gap is a vital blind spot"],
       ["SOC-Lead: Privileged-session audit coverage", "100%", "Op-health",
        "Every engineering / vendor session to vital zones reviewed within 24 hr; FR1/FR2"],
       ["SOC-Lead: New use cases per quarter", "> 1", "Leading",
        "Landscape evolves; sourced from intel, post-mortems, red team"],
       ["SOC-Lead: Threat-hunt campaigns per quarter", "> 2", "Leading",
        "Only mechanism to find pre-positioned threats"],
       ["Analyst: SPAN/TAP coverage completeness", "100% of L1/L2 vital segments", "Op-health",
        "Controller-level visibility is the most important telemetry"],
       ["Analyst: Architecture-conformance drift", "approaching 0", "Leading",
        "Observed flows deviating from the validated model; forbidden-flow detections"],
       ["Analyst: False positives on exempt flows", "approaching 0", "Op-health",
        "No crypto alarms on the passive balise telegram (EN 50159)"],
       ["Analyst: Detection-content freshness", "minimise lag", "Op-health",
        "Time since baseline last synchronised to an architecture change"]])

# =====================================================================
h1("14. Staffing, Competency and Governance Model")
p("The scarcest resource in railway OT security is the individual who is simultaneously fluent in "
  "signalling safety and in cyber security. Cross-sector experience confirms a 12-24 month "
  "development pipeline and warns that neither shortcut - IT-only analysts, or signalling engineers "
  "with no security training - produces effective OT-SOC practitioners. The railway staffing model "
  "therefore builds cyber capability atop IRSSE signalling expertise, treating the latter as the "
  "harder-to-acquire half, and embeds dedicated railway subject-matter expertise that has no analogue "
  "in a generic OT-SOC: the IRSSE signalling SME, the Kavach SME and the Functional Safety "
  "specialist.")
table("Railway OT-SOC staffing and competency model",
      ["Role", "Tier", "Core function", "Required background"],
      [["Tier-1 Analyst (Monitoring Operator)", "T1",
        "Alert triage, initial classification, playbook execution for known scenarios, escalation",
        "Signalling-aware; EI, Kavach, track circuits, axle counters, data loggers; EN 50159 categories; fail-safe principles"],
       ["Tier-2 Analyst (Investigator / Detection Engineer)", "T2",
        "In-depth investigation, network forensics, correlation, detection-content authoring, containment recommendation",
        "Reads railway protocol captures (RaSTA, MCOMM, MVB, CAN); IEC 62443 FRs; liaises with divisional control"],
       ["Tier-3 Analyst (Senior / Threat Hunter)", "T3",
        "Threat hunting, campaign tracking, SIEM tuning, escalation authority for safety-impacting incidents",
        "Deep OT protocol expertise; ICS malware analysis; authority to invoke safety-response procedures"],
       ["IRSSE Signalling SME", "T2/T3",
        "Authoritative interpretation of signalling architecture; conduit and zone context; G2 operational-impact input",
        "IRSSE officer; interlocking and ATP design; SIL allocation; operational railway context"],
       ["Kavach SME", "T2/T3",
        "Kavach RF and EI-interface behaviour; MCOMM and balise telegram semantics; Kavach SOP compliance",
        "Kavach system engineering; RDSO Kavach specification; radio and onboard subsystems"],
       ["Functional Safety Specialist", "T3/T4",
        "Owns the G1 safety-assessment gate; bridges the cyber and safety cases; residual-risk safety judgement",
        "EN 50126 / EN 50129; SIL and safety-case authorship; hazard analysis"],
       ["OT-SOC Lead / IRSSE-CISO interface", "T4",
        "SOC governance, KPI oversight, regulatory reporting, risk-acceptance recommendation, vendor management",
        "Security management (e.g. CISM/CISSP) AND OT context (ISA/IEC 62443 Cybersecurity-Expert level)"]])
p("A representative pilot establishment for a single Kavach corridor requires of the order of 14-18 "
  "full-time equivalents for continuous 24x7 coverage (Tier-1 to Tier-3 plus a threat hunter and a "
  "leave/training reserve), supplemented by the IRSSE, Kavach and functional-safety SMEs who may "
  "be shared across the zone. At national scale the Tier-3 dedicated SOC adds threat-intelligence, "
  "ISAC and architecture-custody functions. RDSO is the natural custodian of the competency "
  "curriculum and the certification scheme that sustains this workforce.")
p("Governance must span the full institutional hierarchy of Indian Railways and the national "
  "cyber-security apparatus. Signalling and train-protection systems are Critical Information "
  "Infrastructure; their protection therefore engages NCIIPC for designation and national protection "
  "oversight and CERT-In for statutory incident reporting and national coordination, in addition to the "
  "internal railway hierarchy. Table 8 defines the governance model and the principal responsibility of "
  "each entity.")
table("National OT-SOC governance and responsibility model",
      ["Entity", "Tier / scope", "Principal OT-SOC responsibility"],
      [["Railway Board", "National policy",
        "Policy mandate, funding, national risk-acceptance authority for SIL-4 residual risk, accountability to government"],
       ["RDSO", "National technical authority (Tier-3)",
        "Custody of the validated architecture and SL-T; detection-content repository; competency certification; standards and Kavach-SOP interpretation; national OT-SOC operation"],
       ["NCIIPC", "National CII protection",
        "Designation of signalling as Critical Information Infrastructure; protection oversight; national-level advisories and audits"],
       ["CERT-In", "National CERT",
        "Statutory incident reporting and timelines; national threat intelligence; coordination during national-scale incidents"],
       ["Zonal Railways", "Zonal (Tier-2)",
        "Operation of the divisional/zonal fused OT-SOC tier; correlation across the zone; OT veto authority over IT-SOC actions"],
       ["Divisions", "Local (Tier-1)",
        "Station and wayside collection; local detection and evidence buffering; first-line response under safe-working rules"]])

# =====================================================================
h1("15. Residual-Risk Methodology and Risk Acceptance for SIL-4 Systems")
p("Residual risk must be quantified, documented and owned - not buried in assumption logs. The "
  "OT-SOC is an assurance instrument, and the open findings of the reference assessment are precisely "
  "such residual risk; they are entered into a register with named ownership and a treatment decision, "
  "and they remain visible as standing monitoring requirements rather than being silently closed. For "
  "SIL-4 systems the methodology imposes additional constraints derived from the EN 50129 safety "
  "case and the TS 50701 cyber security case. First, residual risk affecting a SIL-4 vital function "
  "cannot be accepted at the SOC or divisional level; it must be escalated to the Railway Board with a "
  "supporting safety and cyber-security argument. Second, where a cyber control mandated by IEC "
  "62443 cannot be implemented for sound engineering reasons - the EN 50159 / IEC 62443 "
  "authentication tension on vital Category-1 buses is the archetype - the residual risk must be "
  "reduced As Low As Reasonably Practicable (ALARP) through compensating controls, of which the "
  "OT-SOC's continuous monitoring is itself one, and the residual must be formally argued in the cyber "
  "security case. Third, no risk treatment may weaken a SIL-4 boundary, an EN 50159 protection, an "
  "authentication or MFA requirement, an integrity or replay-protection requirement, or a "
  "trust-boundary control. Table 9 is an extract of the residual-risk register populated from the "
  "assessment evidence.")
table("Residual-risk register (extract) with SIL-4 acceptance routing",
      ["Ref", "Residual risk", "Standing monitoring lesson", "Treatment / acceptance route"],
      [["A-01", "No cryptographic source authentication on vital Cat-1 buses",
        "Police the closed-medium assumption (UC-01/02/13)",
        "Mitigate by monitoring; ALARP; resolve in cyber security case; SIL-4 - Board-level acceptance"],
       ["A-02", "Missing integrity on field-vital train detection (EI to track circuit)",
        "Occupancy-integrity monitoring", "Mitigate; SIL-4 - Board-level acceptance"],
       ["A-03", "No MFA on privileged engineering access",
        "Privileged-session monitoring (UC-05/06)", "Mitigate now; enforce MFA in architecture; divisional risk owner"],
       ["A-04", "No timeliness / RF-anomaly monitoring on open RF",
        "Category-3 RF instrumentation (UC-03/04)", "Mitigate; SIL-4 - Board-level acceptance"],
       ["A-05", "Forbidden vital-to-open-RF edge present in topology",
        "Topology-conformance monitoring (UC-08)", "Avoid (remove edge) and detect; architecture authority"],
       ["A-06", "Incompletely instrumented trust boundaries",
        "Boundary verification (UC-07)", "Mitigate; OT-SOC and network owners"],
       ["B-01", "EN 50159 / IEC 62443 authentication crediting unresolved",
        "Compensating closure-monitoring", "Mitigate by documentation; cyber-safety assurance lead"],
       ["B-02", "TS 50701 SL-T / risk-assessment / cyber-security-case evidence absent",
        "Calibrates all monitoring intensity", "Mitigate by producing evidence; RDSO / Railway Board"]])

# =====================================================================
h1("16. OT-SOC Maturity Model for Indian Railways")
p("A maturity model gives the programme a measurable trajectory. The five-level roadmap below "
  "adapts cross-sector maturity practice to the railway context; Indian Railways' realistic near-term "
  "target is Level 3 across pilot corridors, progressing to Level 4 nationally as the Tier-3 SOC and the "
  "Railway-sector ISAC mature.")
table("Five-level OT-SOC maturity roadmap for Indian Railways",
      ["Level", "Name", "Defining characteristics"],
      [["Level 1", "Initial",
        "Ad-hoc passive monitoring at a few sites; no validated architecture; reactive only; reliance on extended IT SOC"],
       ["Level 2", "Defined",
        "Validated architecture and conduit inventory; Tier-1/2 sensors on vital conduits; core use cases (UC-01 to UC-08); SL-T documented"],
       ["Level 3", "Managed",
        "Three-tier OT-SOC operational; four-gate IR doctrine tested against the real process; KPI framework live; EN 50159 category-aware detectors complete"],
       ["Level 4", "Integrated",
        "National Tier-3 SOC and Railway-sector ISAC; proactive threat hunting; cyber security case maintained; architecture-conformance monitoring automated"],
       ["Level 5", "Optimised",
        "Continuous architecture-to-detection synchronisation; leading-indicator-driven; cross-sector intelligence exchange with NCIIPC and CERT-In"]])

# =====================================================================
h1("17. The 90-Day Pilot-Deployment Programme")
p("To convert strategy into a credible first step, the cross-sector 90-day launch programme is adapted "
  "to a realistic Indian Railways pilot on a representative Kavach corridor. The programme assumes "
  "that Phase 0 - the production of the validated architecture, the conduit inventory and the TS 50701 "
  "SL-T allocation - has been completed, because no sensor should be deployed before the model that "
  "configures it exists. The guiding sequencing principle is that breadth never precedes vital depth, "
  "and deployment never precedes the validated model.")
table("90-day OT-SOC pilot-deployment programme (Kavach corridor)",
      ["Window", "Focus", "Key activities"],
      [["Days 1-30", "Foundation and assessment",
        "Confirm the validated architecture and SL-T; passive asset discovery on the corridor; ratify the SOC charter and decision rights with the zone; identify SPAN/TAP points on EI, field and boundary conduits; assign initial Tier-1 to Tier-3 staff and SMEs (signalling-first hiring)"],
       ["Days 31-60", "Technology and process",
        "Deploy OT-native network detection with railway-protocol parsing; stand up the OT-contextualised SIEM with asset, zone and criticality context; implement UC-01 to UC-08; draft and table-top the four-gate IR playbooks with divisional Operations and the functional-safety specialist"],
       ["Days 61-90", "Operationalisation and governance",
        "Go live with monitoring; tune false positives (notably the passive-telegram exemption, UC-10); conduct one severity-graded IR exercise against the real process under safe-working rules; publish the first Board and SOC-Lead KPI dashboards; enter the assessment's A and B findings into the residual-risk register and route SIL-4 items to the Board"]])

# =====================================================================
h1("18. Lessons Learned")
p("The synthesis of railway-specific assessment evidence with field-proven OT-SOC practice yields "
  "lessons at two levels. At the technical level, the assessment's findings are durable monitoring "
  "requirements: the vital buses rely on a closed-medium assumption that the SOC must continuously "
  "police; field-vital integrity must be observed, not assumed; privileged engineering access is the "
  "principal cyber pathway to a wrong-side failure and must be monitored as such until MFA is "
  "universal; the open RF bearer demands timeliness and RF-anomaly monitoring that cryptography "
  "alone does not provide; forbidden topologies must be both prevented and detected; trust boundaries "
  "must be instrumented rather than merely declared; the telecom backhaul is in scope; and "
  "standards-justified exemptions, such as the passive balise telegram, must be encoded as carefully "
  "as detections. At the programme level, an OT-SOC is an operational capability, not a technology "
  "purchase - tooling without trained analysts, tested playbooks and decision-rights governance yields "
  "visibility without the ability to act, which in an OT environment can be more dangerous than no "
  "visibility at all. The mundane and the sophisticated both matter; talent is the binding constraint; the "
  "SOC must remain outside the vital network; and residual risk on SIL-4 systems must be owned at "
  "Board level, not buried in an assumption log.")

# =====================================================================
h1("19. Future Work")
p("Several avenues warrant development. The EN 50159 / IEC 62443 authentication tension on vital "
  "Category-1 buses should be resolved in the TS 50701 cyber security case, either through a "
  "documented closed-zone argument that credits the safety coding and formally accepts the residual, "
  "or through the introduction of cryptographically authenticated vital bearers, with the OT-SOC's "
  "compensating monitoring formally credited. Railway-specific detection-content libraries - "
  "peer-reviewed signatures and anomaly models for RaSTA, the Kavach interfaces, MVB and balise "
  "telegrams - should be curated and shared by RDSO. Machine-assisted architecture-conformance "
  "monitoring should continuously diff the validated model against the observed topology to detect drift "
  "and forbidden flows automatically, extending UC-08. The safety and security cases should be "
  "integrated under TS 50701 and EN 50129 cross-acceptance; the OT-SOC itself should be subjected to "
  "resilience engineering so that loss of monitoring does not silently erode assurance; a Railway-sector "
  "ISAC should be established and aligned to NCIIPC and CERT-In; and OT-SOC interfaces to the Kavach "
  "Cyber Security SOP should be standardised so that detection, escalation and reporting are uniform "
  "nationwide.")

# =====================================================================
h1("20. Conclusion")
p("The convergence that has modernised Indian Railways signalling - Kavach over open radio, "
  "networked electronic interlockings, centralised TMS and NMS over MPLS, and remotely brokered "
  "engineering access - has dissolved the isolation on which signalling security historically depended. "
  "Protecting this estate requires dedicated Operational Technology Security Operations Centres "
  "whose design is fundamentally different from enterprise IT SOCs: passive-first, safety-aware, "
  "SIL-prioritised, protocol-literate and bound to a validated architecture. The central contribution of "
  "this paper is the argument, evidenced by an independent standards-traceable assessment and "
  "reinforced by field-proven OT-SOC practice, that architecture validation must precede OT-SOC "
  "deployment, together with the demonstration of how a validated architecture mechanically yields the "
  "SOC's operating-model choice, its five-layer visibility plan, its sensor placement, its detection "
  "specification, its incident-response and severity models, its KPIs, its staffing and competency "
  "needs, its governance responsibilities, its residual-risk register and its maturity roadmap. The "
  "recommendation to Indian Railways, RDSO and the Railway Board is unambiguous: invest first in the "
  "validated architecture and the TS 50701 evidence base; adopt a federated hybrid SOC model with "
  "explicit OT decision rights; instrument the vital conduits and trust boundaries before broadening "
  "coverage; build the rare dual-competency workforce on a multi-year horizon under RDSO "
  "certification; govern across the Railway Board, RDSO, NCIIPC, CERT-In, the Zonal Railways and the "
  "Divisions; and constitute the OT-SOC as a permanent, standards-anchored assurance organ that "
  "continuously verifies - through monitoring - the protection on which the safety of the travelling "
  "public increasingly depends. Isolation can no longer be assumed; it must now be observed.")

# =====================================================================
h1("References")
refs = [
 "IEC 62443-3-2, Security for industrial automation and control systems - Part 3-2: Security risk assessment for system design, International Electrotechnical Commission.",
 "IEC 62443-3-3, Security for industrial automation and control systems - Part 3-3: System security requirements and security levels, IEC.",
 "IEC 62443-2-1, Establishing an industrial automation and control system security programme, IEC.",
 "IEC 62443-4-2, Technical security requirements for IACS components, IEC.",
 "CLC/TS 50701, Railway applications - Cybersecurity, CENELEC.",
 "EN 50159:2010, Railway applications - Communication, signalling and processing systems - Safety-related communication in transmission systems, CENELEC.",
 "EN 50126, Railway applications - The specification and demonstration of Reliability, Availability, Maintainability and Safety (RAMS), CENELEC.",
 "EN 50129, Railway applications - Safety-related electronic systems for signalling, CENELEC.",
 "RDSO, Specification for Train Collision Avoidance System (Kavach), Research Designs and Standards Organisation, Ministry of Railways, Government of India.",
 "RDSO / Indian Railways, Kavach Cyber Security Standard Operating Procedure (SOP).",
 "Indian Railways / Railway Board, Signalling Cyber Security Policy and Guidelines.",
 "NCIIPC, Guidelines for the Protection of Critical Information Infrastructure, National Critical Information Infrastructure Protection Centre, Government of India.",
 "CERT-In, Directions on information security practices and incident reporting, Indian Computer Emergency Response Team.",
 "NIST SP 800-82, Guide to Operational Technology (OT) Security, National Institute of Standards and Technology.",
 "ISA/IEC 62443 Cybersecurity Expert competency programme; ISO/IEC 27001 (enterprise-tier governance reference).",
 "IRSE, International Technical Committee guidance on cyber security for signalling systems, Institution of Railway Signal Engineers.",
 "ShieldworkZ, OT SOC Foundational Guide (shieldworkz.com) - cross-sector OT-SOC foundations, operating-model archetypes, five-layer visibility model, the Monitoring Paradox, four-gate incident-response doctrine, severity, KPI, maturity and residual-risk frameworks (primary source adapted and railway-contextualised herein).",
 "Reference Railway OT Cyber Security Architecture Model and Independent Assessment - separated semantic, governance and enforcement authorities; zone/conduit, EN 50159 transmission-category and SIL traceability (primary evidence base for this paper).",
 "OT incident references cited for threat context: CRASHOVERRIDE/INDUSTROYER (2016); TRITON/TRISIS (2017); PIPEDREAM/INCONTROLLER (2022); Maroochy Water (2000).",
]
for i, r in enumerate(refs, 1):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.paragraph_format.left_indent = Inches(0.3)
    para.paragraph_format.first_line_indent = Inches(-0.3)
    rr = para.add_run("[%d]  %s" % (i, r))
    rr.font.size = Pt(9.5)

doc.save("IRSE_OT_SOC_Conference_Paper_Final.docx")
print("SAVED IRSE_OT_SOC_Conference_Paper_Final.docx")
