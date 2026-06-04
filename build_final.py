# -*- coding: utf-8 -*-
"""Build IRSE_OT_SOC_Conference_Paper_Reviewed.docx (editor's revised manuscript)
and Review_Resolution_Matrix.docx."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ============================================================ helpers
def new_doc():
    d = Document()
    n = d.styles["Normal"]; n.font.name = "Times New Roman"; n.font.size = Pt(11)
    for hs, sz in (("Heading 1", 15), ("Heading 2", 12.5), ("Heading 3", 11.5)):
        st = d.styles[hs]; st.font.name = "Times New Roman"; st.font.size = Pt(sz)
        st.font.color.rgb = RGBColor(0x10, 0x2A, 0x43)
    return d

def title(d, t, sz=18):
    p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(t); r.bold = True; r.font.size = Pt(sz); r.font.color.rgb = RGBColor(0x0B,0x1F,0x33)

def sub(d, t):
    p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(t); r.italic = True; r.font.size = Pt(11)

def byl(d, t):
    p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(t); r.font.size = Pt(10.5)

def h1(d, t): d.add_heading(t, level=1)
def h2(d, t): d.add_heading(t, level=2)

def P(d, text, bold=False, italic=False):
    para = d.add_paragraph(); para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = para.add_run(text); r.bold = bold; r.italic = italic
    return para

def KW(d, label, text):
    para = d.add_paragraph(); para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = para.add_run(label + " "); r.bold = True; para.add_run(text)

def bullet(d, text, lead=None):
    para = d.add_paragraph(style="List Bullet")
    if lead:
        r = para.add_run(lead); r.bold = True
    para.add_run(text)

def tbl(d, n, caption, headers, rows):
    cap = d.add_paragraph()
    label = ("%s" % caption) if n is None else ("Table %d. %s" % (n, caption))
    rc = cap.add_run(label); rc.bold = True; rc.font.size = Pt(10)
    t = d.add_table(rows=1, cols=len(headers)); t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, hh in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = ""; r = c.paragraphs[0].add_run(hh); r.bold = True; r.font.size = Pt(9)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ""; r = cells[i].paragraphs[0].add_run(str(v)); r.font.size = Pt(9)
    d.add_paragraph()

def figure(d, path, ref):
    d.add_picture(path, width=Inches(6.3))
    d.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    d.add_paragraph()

def add_toc(d):
    para = d.add_paragraph(); run = para.add_run()
    f1 = OxmlElement('w:fldChar'); f1.set(qn('w:fldCharType'), 'begin')
    it = OxmlElement('w:instrText'); it.set(qn('xml:space'), 'preserve'); it.text = 'TOC \\o "1-2" \\h \\z \\u'
    f2 = OxmlElement('w:fldChar'); f2.set(qn('w:fldCharType'), 'separate')
    tx = OxmlElement('w:t'); tx.text = "Table of Contents - right-click and Update Field in Word to populate."
    f3 = OxmlElement('w:fldChar'); f3.set(qn('w:fldCharType'), 'end')
    for e in (f1, it, f2, tx, f3): run._r.append(e)

# ============================================================ MANUSCRIPT
d = new_doc()
title(d, "Cyber Security Framework Planning and Establishment of Operational Technology Security "
         "Operations Centres (OT-SOCs) for Indian Railways")
sub(d, "An architecture-driven methodology integrating IEC 62443, EN 50159, TS 50701 and Kavach "
       "cyber security requirements with field-proven OT-SOC foundational practice")
byl(d, "[Author(s) and affiliation(s) to be completed per IRSE template]   |   Reviewed manuscript "
       "(incorporates peer-review revisions)")
byl(d, "Audience: IRSSE Officers, RDSO, Railway Board, Signalling Engineers and OT Cyber Security Practitioners")
d.add_paragraph()

h1(d, "Abstract")
P(d, "The nationwide deployment of the Train Collision Avoidance System (Kavach), electronic "
     "interlocking (EI), centralised Traffic Management Systems (TMS) and converged Multi-Protocol "
     "Label Switching (MPLS) telecom backbones has transformed Indian Railways signalling from a "
     "collection of electrically isolated vital circuits into a networked, safety-critical cyber-physical "
     "system. This convergence exposes Operational Technology (OT) to a threat surface that enterprise "
     "Information Technology (IT) Security Operations Centres (SOCs) are neither designed nor "
     "competent to monitor. This paper presents a structured, architecture-driven methodology for "
     "planning and establishing dedicated Operational Technology Security Operations Centres "
     "(OT-SOCs) for Indian Railways. Its central thesis is that a machine-validated, standards-traceable "
     "model of the signalling architecture - its IEC 62443 zones and conduits, its EN 50159 transmission "
     "categories and its Safety Integrity Level (SIL) allocations - must precede and continuously inform "
     "SOC operating-model selection, sensor placement, detection engineering, incident response and "
     "governance. The evidence base is an independent cyber security assessment of a reference railway "
     "OT architecture MODEL - a standards-traceable representation in which the semantic, governance "
     "and enforcement authorities were formally separated and validated; the findings reported here are "
     "therefore properties of that model and of the governance logic it encodes, and are presented as "
     "illustrative monitoring requirements rather than as measured deficiencies of any deployed Indian "
     "Railways installation. Field validation on representative Kavach and electronic-interlocking "
     "installations is identified as essential future work. With that scope understood, the assessment "
     "surfaced durable engineering findings (absence of cryptographic authentication on vital Category-1 "
     "buses, absence of multi-factor authentication on privileged engineering access, absence of "
     "timeliness monitoring on open radio bearers, incompletely instrumented trust boundaries and an "
     "unprotected telecom backhaul) and standards-compliance findings (the EN 50159 / IEC 62443 "
     "authentication-crediting tension and the need for TS 50701 lifecycle evidence). The paper "
     "integrates this railway-specific evidence with field-proven OT-SOC foundational practice: an "
     "evaluation of five OT-SOC operating models with a recommendation for Indian Railways; a "
     "five-layer visibility architecture applied to EI, Kavach, MPLS, radio networks, TMS, NMS, data "
     "loggers and engineering workstations; a risk-assessment and SL-T derivation method; the "
     "Monitoring Paradox grounded in SIL-4 signalling, IEC 62443 and EN 50159; a four-gate, safety-first "
     "incident-response doctrine; a railway SEV1-SEV4 severity model; threat-hunting practice; a "
     "three-audience (Board / SOC-Lead / Analyst) KPI framework; a railway OT-SOC staffing and "
     "competency model; a national governance model spanning Railway Board, RDSO, NCIIPC, CERT-In, "
     "Zonal Railways and Divisions; a residual-risk methodology with explicit risk-acceptance principles "
     "for SIL-4 systems; a five-level maturity roadmap; and a 90-day pilot-deployment programme.")
KW(d, "Keywords:", "OT-SOC; Indian Railways; Kavach; IEC 62443; EN 50159; TS 50701; Electronic "
   "Interlocking; Zones and Conduits; Detection Engineering; MITRE ATT&CK for ICS; SIL-4; Incident "
   "Response; Monitoring Paradox; Residual Risk; SOC Maturity; NCIIPC; CERT-In; RDSO.")

# ---- Nomenclature ----
h1(d, "Nomenclature")
tbl(d, None, "Acronyms and key terms",
    ["Term", "Meaning"],
    [["ATP", "Automatic Train Protection (Kavach)"],
     ["BTS", "Base Transceiver Station (radio)"],
     ["EI", "Electronic Interlocking (SIL-4 vital)"],
     ["EWS", "Engineering Workstation"],
     ["IDMZ", "Industrial De-Militarised Zone (IEC 62443 boundary zone, L3.5)"],
     ["L-Kavach / S-Kavach", "Loco Kavach (onboard) / Stationary Kavach (wayside)"],
     ["MVB / CAN", "Multifunction Vehicle Bus / Controller Area Network (onboard/field buses)"],
     ["NMS / TMS", "Network / Traffic Management System"],
     ["SL-T", "Target Security Level (IEC 62443 / TS 50701)"],
     ["SPAN / TAP", "Switched Port Analyser (mirror) / passive Test Access Point"],
     ["E1/E2/E3", "OT-SOC echelons: Station, Zonal, National (this paper)"],
     ["G1-G4", "Analyst grades: Operator, Investigator, Senior/Hunter, Lead (this paper)"]])
P(d, "Note on protocol nomenclature. Protocol identifiers used for the vital interfaces (the "
     "Kavach-EI interface, and the rolling-stock buses denoted CAN and MVB) and the Kavach radio "
     "protocol are those of the reference architecture model and of common rolling-stock practice; the "
     "authoritative definitions are given in the RDSO Kavach specification [9]. Claims about these "
     "interfaces should be read against that specification rather than as independent protocol "
     "assertions.", italic=True)

# ---- TOC ----
h1(d, "Table of Contents")
add_toc(d)
d.add_page_break()

# 1 Introduction
h1(d, "1. Introduction")
P(d, "For most of its history, the safety of railway signalling rested upon two principles: physical "
     "isolation and fail-safe design. Vital functions were realised in relay logic, hard-wired track "
     "circuits and electrically isolated lineside cabling; the cyber threat surface was, in practice, the "
     "perimeter fence. The doctrine of security through isolation was tenable because the signalling "
     "system genuinely was isolated.")
P(d, "That premise no longer holds. The contemporary Indian Railways signalling estate is a converged, "
     "IP-enabled, geographically distributed cyber-physical system. Electronic interlockings exchange "
     "vital state with object controllers over digital buses; the Kavach Automatic Train Protection (ATP) "
     "system extends vital communication across an open radio-frequency (RF) air interface between the "
     "Stationary Kavach (S-Kavach) at the wayside and the Loco Kavach (L-Kavach) onboard; TMS and "
     "Network Management Systems (NMS) supervise wide geographies over MPLS backbones; and "
     "engineering workstations (EWS) capable of downloading vital logic connect, through VPN access, "
     "jump hosts and the industrial de-militarised zone (IDMZ), into the heart of the safety domain. Each "
     "of these capabilities is a deliberate engineering advance; each also introduces a conduit through "
     "which a determined adversary may attempt to influence a safety function.")
P(d, "The instinctive institutional response - to extend the existing enterprise IT SOC to also cover "
     "signalling - is inadequate and, in places, dangerous. Cross-sector OT-SOC experience [17], and the "
     "wider OT-security literature [14], are unambiguous: extending an IT SOC into OT consistently fails "
     "because the tools, telemetry sources, threat models, response playbooks and human expertise "
     "required are fundamentally different, producing a false sense of security. An IT SOC optimises for "
     "confidentiality; an OT-SOC must optimise for the preservation of safety and the availability of a "
     "physical process. An IT SOC may legitimately quarantine an endpoint on detection; an OT-SOC that "
     "interrupts a vital communication path may itself precipitate the hazard it was deployed to prevent.")
P(d, "This paper argues that the establishment of OT-SOC capability for Indian Railways must be a "
     "discipline in its own right, and that its single most important precondition is an architecture that "
     "has been validated against standards before it is monitored. One cannot monitor what one has not "
     "modelled; one cannot prioritise detection on a flow whose criticality, transmission category and "
     "trust boundary are unknown; and one cannot defend a zone whose conduits have not been "
     "enumerated. The argument is developed from concrete evidence and fused with field-proven "
     "OT-SOC practice, so that Indian Railways need not re-derive at cost what other critical-infrastructure "
     "sectors have already learned.")

# 2 Evolution
h1(d, "2. Evolution of Cyber Security in Indian Railways")
P(d, "Four phases characterise the trajectory of railway signalling cyber security in India. Phase I - "
     "electromechanical isolation - comprised relay interlockings and token working; security was "
     "physical and procedural. Phase II - first-generation electronics - introduced solid-state and "
     "electronic interlockings, data loggers and digital axle counters; digital logic entered the vital "
     "domain but communication remained local, proprietary and serial, and data loggers became, almost "
     "incidentally, the first forensic record of events later recognised as security-relevant. Phase III - "
     "networked OT - brought MPLS backbones, centralised TMS, NMS and remote diagnostics, the first "
     "genuine IT/OT convergence; engineering access became remotely brokered and the threat surface "
     "expanded from the fence to the network. Phase IV - safety over open bearers - is the present, and "
     "its defining event is the nationwide rollout of Kavach, which carries vital movement-authority "
     "information across an open RF bearer and integrates with the EI at the station. For the first time a "
     "vital, SIL-4 function depends upon a transmission system that, in EN 50159 terms, is open. The "
     "RDSO Kavach specification [9] and the Kavach Cyber Security SOP [10] recognise that "
     "cryptographic protection of the Kavach interface is necessary but not sufficient: protection without "
     "monitoring provides no assurance that it is functioning and no evidence for incident response. The "
     "OT-SOC is the institutional mechanism through which protection is continuously verified.")

# 3 Related Work (NEW)
h1(d, "3. Related Work")
P(d, "Railway cyber security has an established literature on which this paper builds. The CENELEC "
     "lifecycle standards EN 50126 [7] and EN 50129 [8] and the dedicated cyber standard CLC/TS 50701 "
     "[5] provide the safety-security integration framework; IEC 62443 [1-4] supplies the zone/conduit "
     "and security-level model adopted here; and EN 50159 [6] supplies the communication-threat "
     "taxonomy. EU research programmes (CYRAIL, X2Rail, SAFETY4RAILS) [21] have addressed "
     "railway-specific threat assessment and resilience, and ENISA has published transport and railway "
     "threat-landscape analyses [20]. In the operational OT-SOC domain, NIST SP 800-82r3 [14] and the "
     "ISA/IEC 62443 series define monitoring and management-system practice, MITRE ATT&CK for ICS "
     "[19] provides the adversary-technique reference now standard in detection engineering, and "
     "cross-sector OT-SOC foundational practice [17] supplies the operating-model, visibility, "
     "incident-response, KPI and maturity frameworks adapted in Sections 8-21. The distinctive "
     "contribution of the present paper relative to this body of work is twofold: the argument that a "
     "machine-validated, authority-separated architecture model must precede and configure the OT-SOC, "
     "and the end-to-end railway contextualisation - from operating-model selection to a 90-day "
     "programme - anchored to Kavach and to Indian Railways' institutional structure. The structural "
     "OT-SOC frameworks adapted here are drawn from cross-sector practice and are cited as such; they "
     "are not claimed as original.")

# 4 Threat landscape
h1(d, "4. The Railway OT Threat Landscape")
P(d, "The railway OT threat landscape differs from the enterprise landscape in its consequences, "
     "adversaries and constraints. Cross-sector incident response [14,17,19,20] shows that OT threats "
     "cluster into five operationally distinct categories; Table 1 maps these to their railway "
     "manifestations and to the consequence that matters in signalling - the wrong-side failure.")
tbl(d, 1, "Cross-sector OT threat categories mapped to Indian Railways",
    ["#", "Generic OT threat category", "Indian Railways manifestation", "Consequence"],
    [["1", "IT-sourced lateral movement (ransomware crossing a weak IT/OT boundary)",
      "Enterprise to IDMZ to operations to interlocking traversal where boundary controls are incompletely instantiated",
      "Disruption of TMS/NMS; staging toward vital zones"],
     ["2", "Engineering-workstation compromise (TRITON/TRISIS, 2017)",
      "Compromise of an EWS with logic-download access to EI / S-Kavach",
      "Alteration of vital interlocking logic - wrong-side failure"],
     ["3", "Supply-chain / vendor-access exploitation",
      "Vendor laptop or update channel for Kavach / EI / TMS subsystems",
      "Implant introduced into the vital supply chain"],
     ["4", "Insider threat and sabotage (Maroochy Water, 2000)",
      "Maintainer/contractor with privileged engineering access and protocol knowledge",
      "Setpoint/logic change; alarm suppression on a vital system"],
     ["5", "Nation-state pre-positioning (INDUSTROYER, TRITON, PIPEDREAM)",
      "Persistent access to Kavach/EI/telecom infrastructure awaiting timing",
      "Coordinated destructive disruption of national rail safety"]])
P(d, "The worst-case outcome is not data loss but a wrong-side failure: a signal cleared against a "
     "conflicting movement, a point moved under a train, a movement authority issued where none is "
     "safe, or the suppression of an emergency brake command. The SIL of the affected asset is therefore "
     "the primary axis along which monitoring effort is prioritised. The adversary spectrum runs from the "
     "opportunistic (commodity malware via an EWS or USB device) through the capable insider (a "
     "maintainer with privileged access and protocol knowledge) to the resourced nation-state actor. The "
     "capable insider is of particular concern because the vital buses rely, by standard, on "
     "closed-transmission assumptions and safety coding rather than cryptographic source "
     "authentication - a defence robust against the outsider but weaker against an actor inside the closed "
     "segment. Sectoral incident data indicate that the majority of OT incidents requiring SOC response "
     "begin with human error or supply-chain compromise rather than sophisticated attack [14,17,20]; a "
     "railway OT-SOC must therefore detect configuration drift, accidental connectivity changes and "
     "unauthorised removable media as reliably as it hunts nation-state implants.")

# 5 Standards
h1(d, "5. Standards Framework: IEC 62443, EN 50159, TS 50701 and Kavach")
h2(d, "5.1 IEC 62443 - Zones, conduits and security levels")
P(d, "IEC 62443-3-2 (risk assessment for system design) [1] and IEC 62443-3-3 (system security "
     "requirements and security levels) [2] decompose a system into zones connected by conduits, and "
     "define seven Foundational Requirements: FR1 Identification and Authentication, FR2 Use Control, "
     "FR3 System Integrity, FR4 Data Confidentiality, FR5 Restricted Data Flow, FR6 Timely Response to "
     "Events, and FR7 Resource Availability. FR6 (and SR 6.2 continuous monitoring) is itself a "
     "monitoring mandate. The reference architecture realises the model as a zone hierarchy mapped to "
     "the Purdue model - Enterprise (Level 5), IDMZ and Security-Management (Level 3.5), Operations "
     "(Level 3), Telecom Core and Radio Access (Levels 2 and 1 of the telecom sub-domain), Interlocking "
     "(Levels 2/1 of the signalling sub-domain), Field (Level 0) and Onboard. The dual use of Levels 1-2 "
     "for the telecom and signalling sub-domains is a railway-specific extension of the Purdue model and "
     "is annotated as such in Figure 1.")
h2(d, "5.2 EN 50159 - Safety-related communication and transmission categories")
P(d, "EN 50159 [6] classifies safety-related transmission systems into three categories defined by the "
     "controllability of the medium and the consequent need for cryptographic protection. In a Category 1 "
     "(closed) system the participants and transmission characteristics are fixed and unauthorised access "
     "to the medium is excluded by design; the defences are non-cryptographic (safety code, sequence "
     "number, time-out). In a Category 2 (controlled) system the transmission system is not under the full "
     "control of the designer, but the residual risk of unauthorised access is assessed as negligible and "
     "can be excluded by design or procedure, so cryptographic protection is not mandated. In a "
     "Category 3 (open) system the possibility of unauthorised access must be assumed; cryptographic "
     "procedures are therefore required in addition to the non-cryptographic defences. The Kavach radio "
     "interface and the balise air-gap are Category 3; the managed MPLS/IP backhaul is Category 2; the "
     "vital station buses are Category 1. Because the category fixes the threat model, it fixes the detector "
     "set the OT-SOC must field; Table 2 makes the threat-defence-detection relationship explicit.")
tbl(d, 2, "EN 50159 threats mapped to safety defences and to OT-SOC detections",
    ["EN 50159 threat", "EN 50159 defence(s)", "OT-SOC detection"],
    [["Repetition", "Sequence number; time stamp; time-out", "Duplicate-frame and freshness analysis (UC-03)"],
     ["Deletion", "Sequence number", "Sequence-gap and loss-of-message detection (UC-04)"],
     ["Insertion", "Sequence number; source/destination identifier; identification procedure", "Source-ID whitelist; insertion/rate anomaly (UC-01, UC-13)"],
     ["Re-sequencing", "Sequence number; time stamp", "Out-of-order detection (UC-03)"],
     ["Corruption", "Safety code; cryptographic techniques", "Safety-code/CRC failure-rate and integrity-flag monitoring (UC-02, UC-10)"],
     ["Delay", "Time stamp; time-out", "End-to-end latency vs SIL-bounded threshold (UC-03)"],
     ["Masquerade", "Source/destination identifier; cryptographic identification", "Closure/identity and unexpected-participant detection (UC-01, UC-13)"]])
P(d, "Two consequences follow and resolve the central evidence finding. First, for Category-1 vital buses "
     "EN 50159 does not mandate cryptographic source authentication - the closed-medium assumption "
     "and the safety code carry the assurance - so the reference model's vital protocols provide integrity "
     "and replay but not authentication; this is standards-consistent PROVIDED the closed-medium "
     "assumption holds, and the OT-SOC's role is to police that assumption continuously. Second, for "
     "Category-3 open bearers cryptographic protection is mandatory but not sufficient: timeliness, replay "
     "and RF-anomaly monitoring remain obligatory because delay and masquerade are first-class "
     "EN 50159 threats that cryptography alone does not detect in operation.")
h2(d, "5.3 TS 50701 - Railway cyber security lifecycle")
P(d, "CLC/TS 50701 [5] adapts IEC 62443 to railways and binds it to the EN 5012x safety lifecycle. It "
     "requires a documented risk assessment, the allocation of target Security Levels (SL-T) to zones and "
     "conduits, and a cyber security case analogous to the safety case. The SL-T per zone is the input "
     "that calibrates monitoring intensity, alert prioritisation and response targets (Section 7).")
h2(d, "5.4 Kavach cyber security requirements")
P(d, "The RDSO Kavach specification [9] and the Kavach Cyber Security SOP [10] impose railway-specific "
     "obligations: cryptographic protection of the Kavach radio interface, integrity and replay protection "
     "of vital telegrams, controlled key management, and the continuous-monitoring and "
     "incident-reporting obligations that an OT-SOC discharges. The interlock is the conceptual core of "
     "this paper: IEC 62443 tells the OT-SOC where the boundaries are; EN 50159 tells it what to look for "
     "on each conduit; TS 50701 tells it why and to what assurance target; and the Kavach SOP defines "
     "the railway-specific obligations it must satisfy and report against.")

# 6 Architecture-driven + limitations
h1(d, "6. The Architecture-Driven Imperative: Validate Before You Monitor")
P(d, "The central methodological proposition is that architecture validation must precede OT-SOC "
     "deployment. A SOC detects deviation from expected behaviour, and the definition of expected is the "
     "architecture. If the architecture is not modelled, the baseline is whatever traffic existed when "
     "monitoring began - including any pre-existing compromise. In the reference assessment [18] the "
     "architecture was expressed as three separated authorities: a semantic authority (the ontology of "
     "asset types, protocols, conduit classes, Purdue levels and protocol capabilities), a governance "
     "authority (zoning, trust boundaries, flow rules, conduit security profiles and EN 50159 transmission "
     "categories) and an enforcement authority (the validators that consume the first two and emit "
     "findings). This separation is what allows a deviation to be classified unambiguously as a real "
     "security gap, a standards-compliance gap, a modelling defect or an enforcement defect.")
P(d, "The assessment demonstrated this concretely. Before validation, several legitimate flows - operator "
     "workstation to TMS, the network-management plane, the telecom backhaul and the Loco-to-Stationary "
     "Kavach safety association - were unmodelled, and a vital-to-open-RF flow that should never exist "
     "(S-Kavach directly cabled to the radio base station, Figure 3) was present but indistinguishable "
     "from the legitimate ones. Only after the governance model was completed could the forbidden flow "
     "be classified as forbidden and prevented, and the legitimate flows be classified as "
     "authorised-but-monitored. The validated architecture is thus not documentation; it is the "
     "configuration source for every sensor, detector and boundary control the OT-SOC deploys, and the "
     "residual findings become the OT-SOC's standing monitoring requirements (consolidated in the "
     "residual-risk register of Section 19).")
h2(d, "6.1 Limitations of the evidence base")
P(d, "The findings cited here were obtained by assessing a reference architecture MODEL, not by "
     "instrumenting deployed Kavach or electronic-interlocking equipment. They are properties of the "
     "model and of the governance rules it encodes; their value is to illustrate, traceably, the classes of "
     "monitoring requirement an OT-SOC must satisfy. Whether any specific finding holds for a given "
     "fielded installation is a question of fact to be established by a field assessment under the TS 50701 "
     "lifecycle. Such field validation - on a representative Kavach corridor and a representative electronic "
     "interlocking - is the immediate next step and is recorded as future work (Section 23).")

# 7 Risk assessment & SL-T (NEW)
h1(d, "7. Risk Assessment and SL-T Derivation")
P(d, "TS 50701 [5] and IEC 62443-3-2 [1] require that target Security Levels be derived for each zone and "
     "conduit from a documented risk assessment, and the OT-SOC uses SL-T to calibrate monitoring "
     "intensity and response targets. This paper recommends a consequence-by-likelihood method. "
     "Consequence is scored primarily by the safety significance of the affected function - its SIL and its "
     "potential to cause a wrong-side failure - and secondarily by operational and reputational impact; a "
     "SIL-4 vital function takes the highest consequence band by default. Likelihood is scored from "
     "exposure (enterprise/Internet reachability, open RF, vendor/engineering access), the EN 50159 "
     "transmission category of the conduit and the strength of existing controls. The SL-T for a zone is "
     "the level required to reduce the residual risk of its highest-consequence credible scenario to "
     "acceptable; conduits inherit the higher SL-T of the zones they join. Table 3 is a worked extract for "
     "a Kavach corridor; the full register is an artefact of the cyber security case and is the Phase-0 "
     "prerequisite for any sensor deployment.")
tbl(d, 3, "Worked SL-T derivation (Kavach-corridor extract)",
    ["Zone", "Consequence (SIL / wrong-side)", "Likelihood drivers", "SL-T", "Monitoring implication"],
    [["Interlocking (EI / S-Kavach)", "Very high (SIL-4)", "Engineering access; Cat-3 RF exposure via Kavach", "SL-T 3",
      "Highest intensity; privileged-access and closure monitoring mandatory"],
     ["Onboard (L-Kavach)", "Very high (SIL-4)", "Open RF; mobile", "SL-T 3", "Cat-3 RF + timeliness monitoring; onboard recorder"],
     ["Operations (TMS/NMS/EWS)", "High (safety-adjacent)", "EWS privilege; IDMZ-brokered access", "SL-T 2-3", "Privileged-session and config-change monitoring"],
     ["Telecom Core (MPLS)", "Medium (no direct vital)", "Managed Cat-2 core", "SL-T 2", "Transport-security health monitoring"],
     ["IDMZ / Security-Mgmt", "Medium-high", "Internet-adjacent; brokers access", "SL-T 2-3", "Boundary-control and access-broker auditing"],
     ["Enterprise", "Low-medium", "Internet-facing", "SL-T 1-2", "Lateral-movement detection at the boundary"]])

# 8 Operating models
h1(d, "8. OT-SOC Operating Models: Evaluation and Recommendation")
P(d, "There is no single correct OT-SOC operating model; the right choice depends on scale, geographic "
     "footprint, regulatory obligation, existing IT-security capability and OT maturity. Five archetypes "
     "[17] cover the realistic spectrum. Table 4 evaluates each against the circumstances of Indian "
     "Railways - a national operator with thousands of dispersed wayside and Kavach sites, a "
     "zonal/divisional structure and SIL-4 assets.")
tbl(d, 4, "Evaluation of the five OT-SOC operating models for Indian Railways",
    ["Model", "Description", "Suitability for Indian Railways"],
    [["A - Dedicated OT-SOC", "Physically separate stack and OT-native staff", "National (E3) level under RDSO/Railway Board; cost justified by SIL-4 criticality"],
     ["B - Fused IT/OT SOC with dedicated OT tier", "Shared SIEM/SOAR with ring-fenced OT tier and OT analyst per shift", "Zonal/Divisional (E2); requires charter giving OT veto over IT actions affecting OT"],
     ["C - MSSP-led", "Lean internal team; monitoring outsourced", "Transitional only; mandatory knowledge transfer; unsuitable as steady state for vital safety"],
     ["D - Asset-owner collaborative (ISAC)", "Pooled sector telemetry; shared intelligence", "Railway-sector ISAC, RDSO-anchored, NCIIPC/CERT-In aligned"],
     ["E - Virtual/distributed", "Lightweight sensors per site, central analysis", "Essential for the thousands of distributed wayside/Kavach sites"]])
P(d, "Recommendation - a federated hybrid. A dedicated national OT-SOC (Model A) at echelon E3, under "
     "RDSO and Railway Board governance, holds the validated architecture, the SL-T allocation, the "
     "cyber security case, national threat intelligence and a Railway-sector ISAC function (Model D). "
     "Zonal OT tiers (Model B), fused with existing IT-SOC infrastructure but with explicit authority to "
     "override IT actions affecting signalling, provide regional correlation. A distributed sensing layer "
     "(Model E) instruments the dispersed wayside, interlocking and Kavach sites. Model C is admissible "
     "only as a time-boxed transitional measure. Across all echelons the single non-negotiable control is "
     "a charter defining decision rights: the OT tier must be able to block any action - including from "
     "IT-SOC staff - that would affect a vital system.")

# 9 Visibility
h1(d, "9. Visibility Architecture: The Five Layers Applied to Railway Assets")
P(d, "Comprehensive visibility requires telemetry from five layers [17]: Network (L2/L3 switches, routers, "
     "firewalls, IDS); Asset/Endpoint (engineering and operator workstations, jump hosts); "
     "Application/Supervisory (TMS, NMS, SIEM/SOC servers); Process/Vital-state (EI vital state, data "
     "loggers, onboard recorders); and Identity and Access (IDMZ directory, VPN/remote-access "
     "gateways, vendor jump servers). Table 5 applies these to the railway assets; Figure 1 places them in "
     "the zone/conduit reference model.")
figure(d, "figs/fig1_zones.png", "Figure 1")
tbl(d, 5, "Five visibility layers applied to Indian Railways OT assets",
    ["Railway asset", "Dominant layer(s)", "Collection method", "Key telemetry"],
    [["Electronic Interlocking (EI)", "Process/Vital-state; Network", "Passive TAP on EI vital buses", "Route/occupancy state, Kavach-EI-interface reads/writes, safety-code failures, state-machine violations"],
     ["Kavach (S-/L-Kavach)", "Network; Process/Vital-state", "Passive TAP + RF monitor; onboard recorder", "Movement-authority frames, integrity, sequence/freshness, brake-command integrity, RF anomaly"],
     ["MPLS / PE routers", "Network", "Passive flow monitoring on telecom-core conduit", "IPsec/MACsec tunnel health, routing anomaly, unexpected participants"],
     ["Radio networks (BTS/gateway)", "Network; Process/Vital-state", "RF spectrum monitor + conduit TAP", "Signal-strength/error-rate anomaly (jamming), latency deviation, replay indicators"],
     ["TMS", "Application/Supervisory", "Syslog/TLS; HTTPS session monitoring", "Operator route/setpoint commands, supervisory-command integrity, failed logins"],
     ["NMS", "Application/Supervisory; Identity", "SNMPv3 event capture", "Unexpected SNMP set operations, config changes, off-hours management access"],
     ["Data loggers", "Process/Vital-state", "Read-only ingestion (never inline)", "Forensic stream; tamper indicators; time-sync drift; gap/reorder detection"],
     ["Engineering workstations (EWS)", "Asset/Endpoint; Identity", "Agent or agentless read-only; PAM/VPN session recording", "Logic-download and config-change actions, USB insertion, anomalous/privileged logins"]])
P(d, "Two planning points follow. First, the highest-consequence telemetry - the process/vital-state "
     "layer for EI and Kavach - is the hardest to obtain, because vital devices emit few logs and cannot "
     "be actively interrogated; it must be derived from passive observation and from the data loggers and "
     "onboard recorders. The balise air-gap concerns the Kavach trackside RFID tag and the onboard "
     "RFID reader; these are Kavach RFID balises and are distinct from the ETCS Eurobalise, with which "
     "they should not be conflated. Second, the identity layer (engineering and vendor/VPN access) is "
     "most often neglected yet most directly tied to the highest-value attack path, the EWS.")

# 10 Monitoring paradox
h1(d, "10. The Monitoring Paradox: Why Railway OT Monitoring Must Remain Passive")
P(d, "Every piece of monitoring infrastructure deployed inside the OT network is itself an attack surface "
     "[17]; a compromised sensor on a vital network is a bridge into the most sensitive environment a "
     "railway operates. In the railway context this Monitoring Paradox has a sharper edge than in any "
     "other sector, because the monitored assets are SIL-4 vital systems whose disturbance can directly "
     "cause a wrong-side failure. The passivity requirement has a precise engineering meaning: on vital "
     "(Category-1) and onboard segments the collection device must be a passive optical or copper TAP "
     "that is electrically and logically incapable of transmitting onto the monitored link; SPAN mirroring, "
     "which is a function of the production switch and can silently drop frames under load, is excluded "
     "from vital segments and reserved for non-vital ones. The Monitoring Paradox is then enforced "
     "architecturally (Figure 2): telemetry leaves the Station echelon strictly upward through a data "
     "diode, there is no online downward control path into the vital tier, and sensor configuration is "
     "applied through a separate, scheduled, out-of-band management channel under change control.")
P(d, "These constraints are grounded in the standards, not in preference. IEC 62443 FR7 (Resource "
     "Availability) and the OT priority order (Safety, then Availability, then Integrity, then "
     "Confidentiality) forbid security mechanisms that degrade the process. EN 50159's treatment of delay "
     "and corruption as first-class threats means that an in-line monitoring device on a SIL-4 bus would "
     "be a potential source of the very hazards the safety layer is designed to exclude. The Monitoring "
     "Paradox is therefore the railway expression of the principle that in OT the cure must never be more "
     "dangerous than the disease.")

# 11 Three-echelon architecture
h1(d, "11. Three-Echelon OT-SOC Architecture")
P(d, "Terminology. To avoid ambiguity this paper distinguishes three uses of the word formerly rendered "
     "as 'tier'. The geographic SOC structure is described in ECHELONS - E1 Station/Wayside, E2 "
     "Zonal/Divisional, E3 National; the analyst workforce in GRADES - G1 Operator, G2 Investigator, G3 "
     "Senior/Hunter, G4 Lead; and the KPI framework by reporting AUDIENCE - Board, SOC-Lead, "
     "Analyst. The federated hybrid of Section 8 is realised as a three-echelon architecture (Figure 2). "
     "E1 (Station/Wayside) is co-located with the relay room and comprises passive TAP sensors on field "
     "and interlocking conduits, data-logger ingestion and a data-diode log export; it performs local "
     "detection and buffers evidence, and never hosts an in-line control on a vital path. E2 (Zonal) hosts "
     "the SIEM, correlation, IT/OT-boundary intrusion sensors in detection mode, vulnerability "
     "management for non-vital assets and case management, and correlates events across stations - for "
     "example a corridor-wide pattern of balise-telegram anomalies. E3 (National) is custodian of the "
     "validated architecture and detection-content repository, the SL-T and cyber security case, cross-zone "
     "correlation, national threat intelligence, the Railway-sector ISAC and KPI consolidation; when the "
     "architecture changes it re-validates and pushes updated baselines downward. Log flow is strictly "
     "upward and diode-enforced across the IT/OT boundary, preserving FR5 within the SOC's own "
     "architecture.")
figure(d, "figs/fig2_echelons.png", "Figure 2")

# 12 Detection engineering
h1(d, "12. Detection Engineering and Railway-Specific Use Cases")
P(d, "Detection engineering converts the validated architecture and its findings into deployable content: "
     "every conduit security requirement and every assessment finding becomes a detection use case. The "
     "vital Category-1 buses - the EI to S-Kavach interface, EI to object controller and speed-sensor to "
     "L-Kavach over CAN, and L-Kavach to driver-machine interface and brake-interface unit over MVB "
     "(identifiers per the reference model and [9]) - provide integrity and replay through safety coding "
     "but not cryptographic source authentication; the SOC cannot add authentication but must detect the "
     "symptoms its absence would allow. Privileged engineering access is not, in the modelled flows, "
     "MFA-gated, so every privileged session is a high-value target. The open RF conduits require "
     "timeliness, jamming and RF-anomaly monitoring; several trust boundaries are not fully "
     "instrumented; and the telecom backhaul carries traffic without an evidenced transport-security "
     "overlay. Table 6 maps the threats to detections; Table 7 is the use-case library; and Table 8 maps "
     "the use cases to MITRE ATT&CK for ICS [19] so that coverage can be assessed against a recognised "
     "adversary model. Each use case additionally records, in the deployed content repository, its data "
     "source, detection logic, validated playbook reference and documented false-positive baseline.")
tbl(d, 6, "Threat-to-detection mapping (EN 50159 / IEC 62443)",
    ["Threat (reference)", "Railway manifestation", "Detection technique", "Sensor / echelon"],
    [["Masquerade/insertion (FR1)", "Spoofed axle-counter or object-controller message", "Source-ID whitelist; state-machine validation; rate anomaly", "Interlocking TAP / E1"],
     ["Corruption (FR3; EN 50159)", "Manipulated EI-to-Kavach telegram", "Safety-code/CRC failure rate; integrity-flag", "EI/Kavach TAP / E1"],
     ["Repetition/replay (EN 50159)", "Replayed Kavach RF movement-authority frame", "Sequence regression; freshness deviation", "Radio conduit, recorder / E1-E2"],
     ["Delay/timeliness (FR6)", "Stale movement authority over RF", "Latency vs SIL-bounded threshold", "Radio conduit / E1-E2"],
     ["Jamming/availability (FR7)", "RF denial vs L-Kavach to BTS", "RF signal-strength/error-rate anomaly", "RF monitor / E1"],
     ["Privileged misuse (FR2)", "Unauthorised logic download from EWS", "Engineering-session + config-change detection", "Ops-to-interlocking / E2"],
     ["Lateral movement (FR5)", "Enterprise to operations traversal", "Cross-boundary flow correlation", "Boundary sensors / E2"],
     ["Backhaul tampering (FR3/FR4)", "Injection on unprotected MPLS/IP core", "Transport-security health; unexpected participant", "Telecom conduit / E2"],
     ["Forbidden topology (FR5)", "Direct vital-to-RF edge (S-Kavach to BTS)", "Topology-conformance vs validated model", "Conformance engine / E3"]])
tbl(d, 7, "Railway-specific OT-SOC detection use-case library",
    ["ID", "Use case", "Trigger", "Sev", "Safety-aware response"],
    [["UC-01", "Spoofed vital field message", "Source-ID not in EI whitelist / state-machine violation", "SEV1", "Alert control; do not block; manual verification under safe-working rules"],
     ["UC-02", "Kavach EI-interface integrity loss", "Rising CRC/safety-code failure", "SEV1", "Alert; correlate EI health; engineering investigation"],
     ["UC-03", "Kavach RF replay / delay", "Sequence regression / latency beyond SIL threshold", "SEV1", "Alert; verify RF; no automated RF action"],
     ["UC-04", "RF jamming / loss of comms", "Signal-strength collapse / error-rate spike", "SEV2", "Alert; invoke degraded-mode working"],
     ["UC-05", "Unauthorised logic download", "Logic/config change from EWS toward EI/S-Kavach", "SEV1", "Alert; freeze change control; verify authorisation"],
     ["UC-06", "Privileged access without MFA", "Jump-host/VPN session, no MFA event", "SEV2", "Alert; challenge session; review policy"],
     ["UC-07", "Trust-boundary bypass", "Flow crossing enterprise to operations without required controls", "SEV2", "Alert; correlate lateral movement"],
     ["UC-08", "Forbidden topology edge", "S-Kavach to BTS (or any forbidden flow) detected", "SEV1", "Alert; emergency architecture review"],
     ["UC-09", "Backhaul transport-security failure", "Loss/absence of IPsec/MACsec on TMS-MPLS path", "SEV3", "Alert; verify tunnel health"],
     ["UC-10", "Balise integrity / position anomaly", "Telegram safety-code failure / position implausibility", "SEV2", "Alert; NO crypto-encryption alarm (EN 50159 passive-telegram exemption)"],
     ["UC-11", "Data-logger tampering / time drift", "Gap, reorder or clock skew", "SEV3", "Alert; preserve forensic chain"],
     ["UC-12", "Management-plane anomaly", "Unexpected SNMPv3 set on NMS", "SEV3", "Alert; verify change authorisation"],
     ["UC-13", "New device on vital segment", "MAC/IP not in inventory on interlocking zone", "SEV2", "Alert; verify closed-medium assumption"],
     ["UC-14", "Removable media on EWS", "USB device-ID not in approved list", "SEV3", "Alert; isolate offline EWS per procedure"]])
tbl(d, 8, "Use cases mapped to MITRE ATT&CK for ICS",
    ["Use case", "ATT&CK for ICS tactic", "Representative technique"],
    [["UC-01 Spoofed vital message", "Impair Process Control", "Spoof Reporting Message; Modify Parameter"],
     ["UC-03 RF replay/delay", "Inhibit Response Function", "Adversary-in-the-Middle; Manipulation of Control"],
     ["UC-04 Jamming", "Inhibit Response Function", "Denial of Service; Block Reporting Message"],
     ["UC-05 Logic download", "Lateral Movement / Execution", "Program Download; Modify Controller Tasking"],
     ["UC-07 Boundary bypass", "Lateral Movement", "Exploitation of Remote Services"],
     ["UC-08 Forbidden edge", "Initial Access / Persistence", "Hardware Additions; Rogue Master"],
     ["UC-13 New device", "Initial Access", "Rogue Master; Hardware Additions"],
     ["UC-14 Removable media", "Initial Access", "Replication Through Removable Media"]])
P(d, "A deliberately exempt condition illustrates why the validated architecture matters. The passive "
     "balise telegram is protected by the passive safety-coded telegram, with freshness from location "
     "coding and odometry, not by cryptographic encryption or replay (EN 50159). The OT-SOC must "
     "therefore not raise a missing-encryption or missing-replay alarm on the passive telegram (a false "
     "positive, UC-10) while still monitoring telegram integrity and positional plausibility. Encoding "
     "standards-justified exemptions as carefully as detections is a discipline only a validated, "
     "authority-separated architecture makes reliable.")

# 13 Trust boundaries
h1(d, "13. Trust Boundaries and Monitoring Matrix")
P(d, "The trust boundary is the unit of OT-SOC instrumentation; each boundary's required controls are "
     "simultaneously its monitoring specification. The assessment found several boundary controls "
     "declared in governance but not yet instantiated, so the OT-SOC must verify control presence rather "
     "than presume it. Table 9 gives the matrix.")
tbl(d, 9, "Trust-boundary monitoring matrix",
    ["Trust boundary", "EN 50159 relevance", "Required controls", "OT-SOC monitoring obligation"],
    [["Enterprise (L5) - IDMZ (L3.5)", "IT/OT divide", "Firewall, inspection", "Verify controls; detect unauthorised ingress; diode log export"],
     ["IDMZ (L3.5) - Operations (L3)", "-", "Firewall, inspection", "Brokered-access audit; jump-host/VPN session + MFA-event verification"],
     ["Operations (L3) - Interlocking (L2)", "Cat 1 downstream", "Firewall, inspection, integrity, authentication", "Highest scrutiny; engineering-access and logic-download detection"],
     ["Telecom Core (L2) - Radio Access (L1)", "Cat 2", "Firewall, inspection, authentication", "Transport-security health; participant verification"],
     ["Radio Access (L1) - Onboard", "Cat 3 (open RF)", "Integrity, replay, authentication, radio monitoring", "Full Cat-3 set: timeliness, jamming, RF anomaly, replay"],
     ["Balise air-gap (Field - Onboard)", "Cat 3 passive", "Integrity (passive telegram); crypto exempt", "Telegram integrity and positional plausibility; no crypto false positives"]])

# 14 IR doctrine
h1(d, "14. Incident Response Doctrine: The Four-Gate, Safety-First Model")
P(d, "Railway OT incident response inverts the IT assumption that speed is paramount: it assesses safety "
     "impact before every action, contains only when containment does not create a more dangerous "
     "state, and recovers in coordination with Operations and Engineering. A four-gate model (Figure 4, "
     "Table 10) governs every action, and no gate permits an action that interrupts a vital path or induces "
     "a wrong-side state. The repertoire is graduated - observe and enrich, then alert the controller, then "
     "recommend a manual protective action - the last always executed by signalling operations under "
     "safe-working rules, never automatically by the SOC tool.")
figure(d, "figs/fig4_fourgate.png", "Figure 4")
tbl(d, 10, "Four-gate OT incident-response doctrine adapted for railway operations",
    ["Gate", "Question", "Decision authority"],
    [["G1 Safety assessment", "Does the incident or any response create/worsen a wrong-side or safety risk?", "Functional-Safety engineer - mandatory before any containment"],
     ["G2 Operational-impact", "Consequence of the incident continuing vs of the response on live operations?", "Divisional Control / Station Master with SOC"],
     ["G3 Containment", "Minimum viable action that halts progression without unacceptable impact; reversible?", "OT-SOC Lead + Operations Lead - logged"],
     ["G4 Recovery authorisation", "Threat remediated, segment re-validated, safe to return to service?", "RDSO / Divisional safety authority - formal sign-off"]])
P(d, "Decision rights are explicit: the OT-SOC owns detection, enrichment and recommendation; signalling "
     "operations owns any action affecting train movement; and RDSO / Railway Board governance owns "
     "escalation and external reporting. Forensic readiness is a precondition: data loggers and onboard "
     "recorders must have assured integrity and time-synchronisation before an incident.")
figure(d, "figs/fig3_kavach.png", "Figure 3")

# 15 Severity
h1(d, "15. Railway Incident Severity Model (SEV1-SEV4)")
P(d, "Severity in OT is defined by potential physical consequence, not data value. Table 11 anchors "
     "severity to the wrong-side failure and the SIL of the affected asset. The escalation timelines below "
     "are the authors' recommended targets; they must be reconciled against the specific reporting "
     "timelines of the Kavach Cyber Security SOP [10] and the statutory CERT-In window [13] during "
     "Phase 0 of the programme.")
tbl(d, 11, "Railway OT incident severity classification",
    ["Level", "Definition (consequence-based)", "Examples", "Response / escalation"],
    [["SEV1 Safety-critical", "Active threat to safety; vital/Kavach/EI compromise confirmed or highly probable; potential wrong-side failure", "Manipulation of EI vital logic; spoofed Kavach movement authority", "All-hands; G1 invoked; functional-safety + executive notice within 15 min; CERT-In within statutory window; NCIIPC engaged"],
     ["SEV2 Operation-critical", "Confirmed OT compromise able to disrupt operations; active lateral movement; containment will affect services", "Ransomware on TMS/historian; compromised EWS with live EI/Kavach link", "G3/Lead escalation within 15 min; Operations within 30 min; CISO/IRSSE within 1 hr"],
     ["SEV3 OT-zone compromise", "Confirmed malicious presence, not yet able/intending to affect the vital process", "Malware on isolated EWS; phished engineer; rogue device", "1-hr investigation; Operations informed; CISO same business day"],
     ["SEV4 Anomaly/policy", "Policy violation or anomaly, no confirmed compromise", "Unauthorised media; out-of-window vendor access; config drift", "Next-business-day triage; trend-tracked"]])

# 16 Threat hunting (NEW)
h1(d, "16. Threat Hunting")
P(d, "Reactive detection cannot find an adversary who has pre-positioned without triggering a rule; "
     "proactive threat hunting is the only mechanism that can. In the railway OT-SOC, hunting is "
     "hypothesis-led and anchored to MITRE ATT&CK for ICS [19] and to the validated architecture: "
     "hunters look for evidence of the techniques in Table 8 that would precede a wrong-side failure - "
     "rogue masters or hardware additions on a vital segment (UC-08, UC-13), program-download activity "
     "outside a change window (UC-05), and adversary-in-the-middle indicators on the Kavach bearer "
     "(UC-03). Hunts are conducted entirely on the passively-collected E2/E3 data and on data-logger and "
     "onboard-recorder archives, never by active interrogation of vital devices. Each hunt that yields a "
     "repeatable indicator is promoted to a new detection use case, feeding the KPI on new use cases per "
     "quarter (Section 17). A minimum of two hunt campaigns per quarter is recommended, rising for "
     "high-risk corridors.")

# 17 KPI
h1(d, "17. KPI Framework: Board, SOC-Lead and Analyst Audiences")
P(d, "OT-SOC performance is measured by safety-weighted, standards-anchored indicators rather than "
     "volume metrics, distinguishing leading, lagging and operational-health classes across three "
     "reporting audiences. Targets are concrete and time-bound (Table 12). The mean-time-to-detect "
     "target is anchored to IEC 62443-3-3 SR 6.2 (continuous monitoring) and the operational "
     "hazard-exposure budget, not to EN 50129. Quantitative staffing and uptime figures are adopted from "
     "cross-sector practice [17] and are to be validated against an Indian Railways workload study during "
     "the pilot.")
tbl(d, 12, "Three-audience KPI framework adapted for Indian Railways",
    ["Audience: KPI", "Concrete target", "Class"],
    [["Board: Vital-asset (SIL-3/4 conduit) monitoring coverage", "100% within 18 months of go-live; sustained", "Leading"],
     ["Board: MTTD (vital-affecting known technique)", "< 1 hr; < 4 hr unknown", "Lagging"],
     ["Board: Mean time to safe advisory (SEV1)", "< 1 hr", "Lagging"],
     ["Board: Incidents contained without operational impact", "> 90%", "Lagging"],
     ["Board: Regulatory-evidence availability", "100% on demand", "Op-health"],
     ["SOC-Lead: Trust-boundary instrumentation", "100% within 12 months", "Leading"],
     ["SOC-Lead: EN 50159 detector completeness", "100% (detectors match each conduit's category)", "Leading"],
     ["SOC-Lead: Sensor/TAP uptime", ">= 99.5% monthly", "Op-health"],
     ["SOC-Lead: Privileged-session audit coverage", "100% within 24 hr of session", "Op-health"],
     ["SOC-Lead: New use cases / threat hunts per quarter", ">= 1 / >= 2", "Leading"],
     ["Analyst: SPAN/TAP coverage of L1/L2 vital segments", "100%", "Op-health"],
     ["Analyst: Architecture-conformance drift; exempt-flow false positives", "0; 0", "Leading / Op-health"]])

# 18 Staffing, competency, governance
h1(d, "18. Staffing, Competency and Governance Model")
P(d, "The scarcest resource is the individual fluent in both signalling safety and cyber security; "
     "cross-sector experience confirms a 12-24-month development pipeline and warns that neither "
     "IT-only analysts nor signalling engineers without security training produce effective practitioners "
     "[17]. The model builds cyber capability atop IRSSE signalling expertise and embeds dedicated "
     "railway expertise with no analogue in a generic OT-SOC: the IRSSE signalling SME, the Kavach SME "
     "and the Functional Safety specialist (Table 13). Analyst grades G1-G4 are distinct from the "
     "geographic echelons E1-E3; the SMEs are grade-G2/G3 roles serving across echelons. A pilot "
     "corridor requires of the order of 14-18 full-time equivalents for 24x7 cover (a figure adopted from "
     "[17] for validation during the pilot); RDSO is the custodian of the competency curriculum and "
     "certification.")
tbl(d, 13, "Railway OT-SOC staffing and competency model",
    ["Role (grade)", "Core function", "Required background"],
    [["G1 Monitoring Operator", "Triage, classification, playbook execution, escalation", "Signalling-aware; EI, Kavach, track circuits, data loggers; EN 50159 categories"],
     ["G2 Investigator / Detection Engineer", "Investigation, forensics, correlation, content authoring", "Reads railway protocol captures (Kavach-EI iface, MVB, CAN); IEC 62443 FRs"],
     ["G3 Senior Analyst / Threat Hunter", "Hunting, campaign tracking, SIEM tuning, safety-response authority", "Deep OT protocol expertise; ICS malware analysis"],
     ["IRSSE Signalling SME (G2/G3)", "Authoritative architecture interpretation; G2-gate operational input", "IRSSE officer; interlocking/ATP design; SIL allocation"],
     ["Kavach SME (G2/G3)", "Kavach RF and EI-interface behaviour; balise semantics; SOP compliance", "Kavach system engineering; RDSO Kavach specification [9]"],
     ["Functional Safety Specialist (G3/G4)", "Owns the G1 safety gate; bridges cyber and safety cases", "EN 50126/50129; SIL and safety-case authorship"],
     ["G4 OT-SOC Lead / IRSSE-CISO interface", "Governance, KPI, reporting, risk-acceptance recommendation", "Security management (CISM/CISSP) and ISA/IEC 62443 expert level"]])
P(d, "Governance spans the full institutional hierarchy and the national cyber-security apparatus. "
     "Signalling and train-protection systems are Critical Information Infrastructure; their protection "
     "therefore engages NCIIPC for designation and protection oversight under section 70 of the "
     "Information Technology Act 2000 [12], and CERT-In for statutory incident reporting under the "
     "CERT-In Directions of 28 April 2022, which require reporting within six hours of noticing a "
     "reportable incident [13], in addition to the internal railway hierarchy. The SOC charter must define "
     "the interplay: the OT-SOC notifies the internal RDSO/Railway-Board chain and, for incidents "
     "meeting statutory thresholds, CERT-In within the mandated window, with NCIIPC engaged for "
     "CII-level events. Table 14 defines the governance model.")
tbl(d, 14, "National OT-SOC governance and responsibility model",
    ["Entity", "Scope (echelon)", "Principal OT-SOC responsibility"],
    [["Railway Board", "National policy", "Mandate, funding, national risk-acceptance authority for SIL-4 residual risk"],
     ["RDSO", "National technical (E3)", "Custody of validated architecture and SL-T; detection-content repository; competency certification; national OT-SOC operation"],
     ["NCIIPC", "National CII protection", "CII designation (IT Act s.70); protection oversight; national advisories and audits"],
     ["CERT-In", "National CERT", "Statutory 6-hour incident reporting; national threat intelligence; coordination"],
     ["Zonal Railways", "Zonal (E2)", "Operation of the zonal fused OT-SOC; cross-zone correlation; OT veto over IT actions"],
     ["Divisions", "Local (E1)", "Station/wayside collection; local detection and evidence buffering; first-line response under safe-working rules"]])

# 19 Residual risk
h1(d, "19. Residual-Risk Methodology and Risk Acceptance for SIL-4 Systems")
P(d, "Residual risk must be quantified, documented and owned. The open findings of the reference "
     "assessment are entered into a register with named ownership and a treatment decision (Table 15) "
     "and remain visible as standing monitoring requirements. For SIL-4 systems three constraints apply. "
     "First, residual risk affecting a SIL-4 vital function cannot be accepted at SOC or divisional level; it "
     "is escalated to the Railway Board with a supporting safety and cyber-security argument. Second, "
     "where a control mandated by IEC 62443 cannot be implemented for sound engineering reasons - the "
     "EN 50159 / IEC 62443 authentication tension on vital Category-1 buses is the archetype - the "
     "residual risk is reduced As Low As Reasonably Practicable through compensating controls, of which "
     "the OT-SOC's continuous monitoring is one, and is argued in the cyber security case. Third, no risk "
     "treatment may weaken a SIL-4 boundary, an EN 50159 protection, an authentication or MFA "
     "requirement, an integrity or replay-protection requirement, or a trust-boundary control.")
tbl(d, 15, "Residual-risk register (extract) with SIL-4 acceptance routing",
    ["Ref", "Residual risk", "Standing monitoring lesson", "Treatment / acceptance route"],
    [["A-01", "No cryptographic authentication on vital Cat-1 buses", "Police closed-medium assumption (UC-01/02/13)", "Mitigate; ALARP; cyber security case; SIL-4 - Board acceptance"],
     ["A-02", "Missing integrity on field-vital train detection", "Occupancy-integrity monitoring", "Mitigate; SIL-4 - Board acceptance"],
     ["A-03", "No MFA on privileged engineering access", "Privileged-session monitoring (UC-05/06)", "Mitigate now; enforce MFA in architecture; divisional owner"],
     ["A-04", "No timeliness/RF-anomaly monitoring on open RF", "Cat-3 RF instrumentation (UC-03/04)", "Mitigate; SIL-4 - Board acceptance"],
     ["A-05", "Forbidden vital-to-open-RF edge present", "Topology-conformance monitoring (UC-08)", "Avoid (remove) + detect; architecture authority"],
     ["A-06", "Incompletely instrumented trust boundaries", "Boundary verification (UC-07)", "Mitigate; OT-SOC + network owners"],
     ["B-01", "EN 50159 / IEC 62443 authentication crediting unresolved", "Compensating closure-monitoring", "Mitigate by documentation; assurance lead"],
     ["B-02", "TS 50701 SL-T / risk-assessment / cyber-security-case evidence absent", "Calibrates monitoring intensity", "Mitigate by producing evidence; RDSO / Railway Board"]])

# 20 Maturity
h1(d, "20. OT-SOC Maturity Model for Indian Railways")
P(d, "The five-level roadmap below adapts cross-sector maturity practice [17] and IEC 62443-2-1 "
     "management-system concepts [3] to the railway context; the realistic near-term target is Level 3 "
     "across pilot corridors, progressing to Level 4 nationally.")
tbl(d, 16, "Five-level OT-SOC maturity roadmap for Indian Railways",
    ["Level", "Name", "Defining characteristics"],
    [["1", "Initial", "Ad-hoc passive monitoring at a few sites; no validated architecture; reactive only"],
     ["2", "Defined", "Validated architecture and conduit inventory; E1/E2 sensors on vital conduits; core use cases; SL-T documented"],
     ["3", "Managed", "Three-echelon OT-SOC operational; four-gate IR tested against the real process; KPI framework live; category-aware detectors complete"],
     ["4", "Integrated", "National E3 SOC and Railway-sector ISAC; proactive threat hunting; cyber security case maintained; conformance monitoring automated"],
     ["5", "Optimised", "Continuous architecture-to-detection synchronisation; leading-indicator-driven; intelligence exchange with NCIIPC and CERT-In"]])

# 21 90-day
h1(d, "21. The 90-Day Pilot-Deployment Programme")
P(d, "The cross-sector 90-day launch programme [17] is adapted to a pilot on a representative Kavach "
     "corridor, assuming Phase 0 (validated architecture, conduit inventory and TS 50701 SL-T allocation) "
     "is complete - no sensor is deployed before the model that configures it. The guiding principle is "
     "that breadth never precedes vital depth, and deployment never precedes the validated model. An "
     "indicative budget envelope (passive TAP/NDR, OT-contextualised SIEM, RF monitors, staffing and "
     "RDSO certification) should accompany the programme for Railway Board approval; a corridor-scale "
     "pilot is modest relative to the Kavach programme itself and is the prerequisite for credible national "
     "costing.")
tbl(d, 17, "90-day OT-SOC pilot-deployment programme (Kavach corridor)",
    ["Window", "Focus", "Key activities"],
    [["Days 1-30", "Foundation and assessment", "Confirm validated architecture and SL-T; passive asset discovery; ratify SOC charter and decision rights; identify TAP points on EI/field/boundary conduits; assign G1-G3 staff and SMEs (signalling-first hiring)"],
     ["Days 31-60", "Technology and process", "Deploy OT-native NDR with railway-protocol parsing; stand up OT-contextualised SIEM with asset/zone/criticality context; implement UC-01 to UC-08; draft and table-top the four-gate IR playbooks with Operations and the functional-safety specialist"],
     ["Days 61-90", "Operationalisation and governance", "Go live; tune false positives (passive-telegram exemption, UC-10); run one severity-graded IR exercise under safe-working rules; publish first Board and SOC-Lead KPI dashboards; enter A/B findings into the residual-risk register and route SIL-4 items to the Board"]])

# 22 Lessons learned
h1(d, "22. Lessons Learned")
P(d, "At the technical level the assessment's findings are durable monitoring requirements: vital buses "
     "rely on a closed-medium assumption the SOC must continuously police; field-vital integrity must be "
     "observed, not assumed; privileged engineering access is the principal cyber pathway to a "
     "wrong-side failure and must be monitored until MFA is universal; the open RF bearer demands "
     "timeliness and RF-anomaly monitoring that cryptography alone does not provide; forbidden "
     "topologies must be both prevented and detected; trust boundaries must be instrumented, not merely "
     "declared; the telecom backhaul is in scope; and standards-justified exemptions must be encoded as "
     "carefully as detections. At the programme level, an OT-SOC is an operational capability, not a "
     "technology purchase; the mundane and the sophisticated both matter; talent is the binding "
     "constraint; the SOC must remain outside the vital network; and residual risk on SIL-4 systems must "
     "be owned at Board level.")

# 23 Future work
h1(d, "23. Future Work")
P(d, "The immediate next step is field validation of the model-derived findings on a representative "
     "Kavach corridor and electronic interlocking under the TS 50701 lifecycle. Further work includes: "
     "resolving the EN 50159 / IEC 62443 authentication tension in the cyber security case; curating "
     "railway-specific detection-content libraries (signatures and anomaly models for the Kavach "
     "interfaces, MVB and balise telegrams) under RDSO; machine-assisted architecture-conformance "
     "monitoring that diffs the validated model against observed topology (extending UC-08); integration "
     "of the safety and security cases under TS 50701 and EN 50129 cross-acceptance; resilience "
     "engineering of the OT-SOC itself; establishment of a Railway-sector ISAC aligned to NCIIPC and "
     "CERT-In; and standardisation of OT-SOC interfaces to the Kavach Cyber Security SOP.")

# 24 Conclusion
h1(d, "24. Conclusion")
P(d, "The convergence that has modernised Indian Railways signalling has dissolved the isolation on "
     "which signalling security historically depended. Protecting this estate requires dedicated OT-SOCs "
     "that are passive-first, safety-aware, SIL-prioritised, protocol-literate and bound to a validated "
     "architecture. The central contribution of this paper is the argument, evidenced by an independent "
     "standards-traceable model assessment and reinforced by field-proven practice, that architecture "
     "validation must precede OT-SOC deployment, together with the demonstration of how a validated "
     "architecture mechanically yields the SOC's operating-model choice, its five-layer visibility plan, its "
     "SL-T calibration, its detection specification, its incident-response and severity models, its KPIs, its "
     "staffing and governance, its residual-risk register and its maturity roadmap. The recommendation "
     "to Indian Railways, RDSO and the Railway Board is unambiguous: invest first in the validated "
     "architecture and the TS 50701 evidence base; adopt a federated hybrid SOC model with explicit OT "
     "decision rights; instrument the vital conduits and trust boundaries before broadening coverage; build "
     "the dual-competency workforce under RDSO certification; govern across the Railway Board, RDSO, "
     "NCIIPC, CERT-In, the Zonal Railways and the Divisions; and constitute the OT-SOC as a permanent, "
     "standards-anchored assurance organ. Isolation can no longer be assumed; it must now be observed.")

# References
h1(d, "References")
refs = [
 "IEC 62443-3-2, Security risk assessment for system design, IEC.",
 "IEC 62443-3-3, System security requirements and security levels, IEC.",
 "IEC 62443-2-1, Establishing an industrial automation and control system security programme, IEC.",
 "IEC 62443-4-2, Technical security requirements for IACS components, IEC.",
 "CLC/TS 50701, Railway applications - Cybersecurity, CENELEC.",
 "EN 50159:2010, Railway applications - Safety-related communication in transmission systems, CENELEC.",
 "EN 50126, Railway applications - RAMS, CENELEC.",
 "EN 50129, Railway applications - Safety-related electronic systems for signalling, CENELEC.",
 "RDSO, Specification for Train Collision Avoidance System (Kavach), Ministry of Railways, Government of India.",
 "RDSO / Indian Railways, Kavach Cyber Security Standard Operating Procedure (SOP) [document number and issue to be confirmed].",
 "Indian Railways / Railway Board, Signalling Cyber Security Policy and Guidelines [document number and issue to be confirmed].",
 "Government of India, Information Technology Act 2000, section 70; NCIIPC, Guidelines for the Protection of Critical Information Infrastructure.",
 "CERT-In, Directions under sub-section (6) of section 70B of the IT Act, 28 April 2022 (six-hour incident-reporting obligation).",
 "NIST SP 800-82 Rev. 3, Guide to Operational Technology (OT) Security, NIST.",
 "ISA/IEC 62443 Cybersecurity Expert competency programme; ISO/IEC 27001 (enterprise-tier governance reference).",
 "IRSE, International Technical Committee guidance on cyber security for signalling systems, Institution of Railway Signal Engineers.",
 "ShieldworkZ, OT SOC Foundational Guide (shieldworkz.com) - cross-sector OT-SOC foundations: operating-model archetypes, five-layer visibility model, Monitoring Paradox, four-gate IR doctrine, severity, KPI, maturity and residual-risk frameworks (adapted and railway-contextualised herein).",
 "Reference Railway OT Cyber Security Architecture Model and Independent Assessment - separated semantic, governance and enforcement authorities; zone/conduit, EN 50159 transmission-category and SIL traceability (primary evidence base).",
 "MITRE ATT&CK for ICS, The MITRE Corporation - adversary tactics and techniques for industrial control systems.",
 "ENISA, Railway Cybersecurity and Transport Threat Landscape reports, European Union Agency for Cybersecurity.",
 "EU railway-cyber research outputs: CYRAIL, X2Rail and SAFETY4RAILS project deliverables.",
]
for i, r in enumerate(refs, 1):
    para = d.add_paragraph(); para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.paragraph_format.left_indent = Inches(0.3); para.paragraph_format.first_line_indent = Inches(-0.3)
    rr = para.add_run("[%d]  %s" % (i, r)); rr.font.size = Pt(9.5)

d.save("IRSE_OT_SOC_Conference_Paper_Reviewed.docx")
print("SAVED IRSE_OT_SOC_Conference_Paper_Reviewed.docx")

# ============================================================ RESOLUTION MATRIX
m = new_doc()
title(m, "Review Comment Resolution Matrix", 17)
sub(m, "Paper: Cyber Security Framework Planning and Establishment of OT-SOCs for Indian Railways")
sub(m, "Editor disposition of the independent peer review (Reviewer_Assessment.docx)")
m.add_paragraph()
h1(m, "1. Major Comment Resolution")
tbl(m, 1, "Major comment resolution (all major comments are mandatory unless technically incorrect)",
    ["ID", "Reviewer comment (quoted/condensed)", "Disposition", "Technical justification", "Section(s) modified"],
    [["M1", "Evidence base is a model, not deployed systems - risk of over-claim.", "Accepted",
      "Abstract and new Section 6.1 now state the evidence is model-derived; field validation added to Future Work. No claim of measured deficiency on fielded assets remains.", "Abstract; 6.1; 23"],
     ["M2", "No figures/diagrams.", "Accepted",
      "Four figures added and referenced: zone/conduit model, three-echelon topology, Kavach legitimate-vs-forbidden, four-gate IR flow.", "Figs 1-4; 9, 10, 11, 14"],
     ["M3", "Over-reliance on a single commercial source; weak novelty/independence.", "Accepted",
      "New Related Work section positions the paper against primary literature (NIST 800-82r3, IEC 62443-2-1, ENISA, MITRE ATT&CK for ICS, EU projects); the novel validated-architecture thesis is foregrounded; adapted frameworks are explicitly cited as such.", "3; References [19-21]"],
     ["M4", "Technical inaccuracies: Eurobalise/RFID; MCOMM/RaSTA naming; Purdue labelling.", "Accepted",
      "Kavach RFID balise distinguished from ETCS Eurobalise; protocol identifiers qualified as model/[9] terms in a new Nomenclature note; dual L1/L2 Purdue use justified as a railway extension and annotated in Fig 1.", "Nomenclature; 5.1; 9; 12"],
     ["M5", "Detection-engineering methodology incomplete (no ATT&CK, no EN 50159 threat-defence table).", "Accepted",
      "EN 50159 threat-defence-detection table (Table 2) and MITRE ATT&CK-for-ICS mapping (Table 8) added; per-use-case data-source/playbook/false-positive fields mandated in text.", "5.2; 12"],
     ["M6", "No risk-assessment / SL-T derivation method.", "Accepted",
      "New Section 7 gives a consequence-by-likelihood method and a worked Kavach SL-T table (Table 3).", "7"],
     ["M7", "Cross-reference errors; 'Tier' overloaded.", "Accepted",
      "Cross-references corrected; terminology disambiguated into Echelons (E1-E3), Grades (G1-G4) and KPI Audiences; applied throughout.", "11; 17; 18; 19; Nomenclature"],
     ["M8", "Regulatory obligations (CERT-In, NCIIPC) asserted, not substantiated; SLAs vs Kavach SOP.", "Partially Accepted",
      "CERT-In 6-hour rule (28 Apr 2022 Directions) and NCIIPC CII basis (IT Act s.70) now cited [12,13]. Severity SLAs are presented as recommended targets to be reconciled with the Kavach SOP in Phase 0, because the SOP timelines are controlled and cannot be quoted verbatim here - full acceptance (quoting SOP figures) is therefore not technically possible.", "15; 18; References [12,13]"]])
P(m, "Summary: of eight Major Comments, seven are Accepted and one (M8) is Partially Accepted; none "
     "was Rejected. The single partial acceptance reflects the unavailability of the controlled Kavach "
     "SOP timelines for verbatim citation, not a disagreement with the reviewer.")
h1(m, "2. Minor Comment Resolution")
tbl(m, 2, "Minor comment resolution (recommended improvements)",
    ["ID", "Comment", "Disposition", "Action"],
    [["m1", "SPAN vs TAP conflated", "Accepted", "Section 9/10 now mandate passive TAP on SIL-4 segments; SPAN reserved for non-vital"],
     ["m2", "Diode upward vs control downward contradiction", "Accepted", "Section 10/11: out-of-band, scheduled management path; no online downward path to vital tier (Fig 2)"],
     ["m3", "Unsupported quantitative figures", "Accepted", "FTE/pipeline/uptime attributed to [17] and flagged for pilot validation (Sections 17, 18)"],
     ["m4", "Vague KPI targets", "Accepted", "Table 12 gives concrete, time-bound targets"],
     ["m5", "IDMZ undefined", "Accepted", "Defined in Nomenclature and on first use (Section 1)"],
     ["m6", "MTTD anchored to EN 50129", "Accepted", "Re-anchored to IEC 62443-3-3 SR 6.2 (Section 17)"],
     ["m7", "EN 50159 Cat 2 wording loose", "Accepted", "Normative wording adopted (Section 5.2)"],
     ["m8", "No cost model", "Partially Accepted", "Indicative budget envelope noted (Section 21); full costing deferred to post-pilot"],
     ["m9", "Byline placeholder", "Accepted", "Author/affiliation placeholder per IRSE template retained for completion"],
     ["m10", "Reference doc numbers missing", "Partially Accepted", "[10],[11] marked for document-number confirmation; not publicly available to the authors"],
     ["m11", "Human-error claim uncited", "Accepted", "Cited to [14,17,20] (Section 4)"],
     ["m12", "Acronyms inconsistent / undefined", "Accepted", "Nomenclature table added; acronyms standardised throughout"]])
P(m, "Summary: of twelve Minor Comments, ten are fully Accepted and two (m8, m10) Partially Accepted "
     "owing to information not available to the authors at this stage; none was Rejected.")
m.save("Review_Resolution_Matrix.docx")
print("SAVED Review_Resolution_Matrix.docx")
