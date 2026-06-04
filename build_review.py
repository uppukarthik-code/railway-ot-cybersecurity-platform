# -*- coding: utf-8 -*-
"""Generate Reviewer_Assessment.docx (independent review + rewritten sections)."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()
normal = doc.styles["Normal"]
normal.font.name = "Times New Roman"
normal.font.size = Pt(11)
for hs, sz in (("Heading 1", 15), ("Heading 2", 12.5), ("Heading 3", 11.5)):
    st = doc.styles[hs]; st.font.name = "Times New Roman"; st.font.size = Pt(sz)
    st.font.color.rgb = RGBColor(0x10, 0x2A, 0x43)
_t = [0]

def title(t):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(t); r.bold = True; r.font.size = Pt(17); r.font.color.rgb = RGBColor(0x0B,0x1F,0x33)

def sub(t):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(t); r.italic = True; r.font.size = Pt(11)

def h1(t): doc.add_heading(t, level=1)
def h2(t): doc.add_heading(t, level=2)

def p(text, bold=False, italic=False):
    para = doc.add_paragraph(); para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = para.add_run(text); r.bold = bold; r.italic = italic
    return para

def bullet(text, bold_lead=None):
    para = doc.add_paragraph(style="List Bullet")
    if bold_lead:
        r = para.add_run(bold_lead); r.bold = True
    para.add_run(text)

def num(text, bold_lead=None):
    para = doc.add_paragraph(style="List Number")
    if bold_lead:
        r = para.add_run(bold_lead); r.bold = True
    para.add_run(text)

def table(caption, headers, rows):
    _t[0] += 1
    cap = doc.add_paragraph(); rc = cap.add_run("Table R-%d. %s" % (_t[0], caption))
    rc.bold = True; rc.font.size = Pt(10)
    tb = doc.add_table(rows=1, cols=len(headers)); tb.style = "Table Grid"
    tb.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = tb.rows[0].cells[i]; c.text = ""
        run = c.paragraphs[0].add_run(h); run.bold = True; run.font.size = Pt(9.5)
    for row in rows:
        cells = tb.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ""; run = cells[i].paragraphs[0].add_run(str(v)); run.font.size = Pt(9.5)
    doc.add_paragraph()

# ============================ HEADER ============================
title("Independent Peer Review and Required Revisions")
sub("Paper under review: \"Cyber Security Framework Planning and Establishment of Operational "
    "Technology Security Operations Centres (OT-SOCs) for Indian Railways\"")
sub("Review conducted jointly in the roles of IRSE Technical Committee Reviewer, RDSO Reviewer, "
    "IEC 62443 Assessor, EN 50159 Assessor, TS 50701 Reviewer and Railway OT-SOC Architect")
doc.add_paragraph()

# ============================ A. EXECUTIVE REVIEW ============================
h1("A. Executive Review")
p("The manuscript is ambitious, well-organised and addresses a genuinely important and timely "
  "problem: the establishment of dedicated Operational Technology Security Operations Centres "
  "(OT-SOCs) for Indian Railways at the moment the Kavach programme makes vital train protection "
  "dependent on an open radio bearer. Its central thesis - that a standards-validated architecture "
  "model must precede and configure OT-SOC deployment - is sound, defensible and represents the "
  "manuscript's principal original contribution. The standards grammar (IEC 62443 zones and "
  "conduits, EN 50159 transmission categories, TS 50701 lifecycle) is correctly identified and used "
  "as the organising spine, and the railway contextualisation of generic OT-SOC practice "
  "(operating-model selection, five-layer visibility, the Monitoring Paradox, four-gate incident "
  "response, severity, KPIs, staffing, governance, residual risk, maturity and a 90-day programme) is "
  "comprehensive and internally coherent.")
p("However, the manuscript is not yet publishable in its current form. Three classes of weakness must "
  "be remedied. First, a validity concern: the entire evidence base is an independent assessment of a "
  "reference architecture MODEL, not a measurement campaign on deployed Indian Railways "
  "installations; the manuscript repeatedly presents model-derived findings (absent authentication "
  "on vital buses, absent MFA, a forbidden topology edge) in language that a reader could mistake for "
  "empirical facts about fielded Kavach/EI systems. This must be explicitly bounded and a Limitations "
  "section added. Second, technical-accuracy concerns: the conflation of Kavach RFID balises with the "
  "ETCS Eurobalise, the unqualified presentation of model-internal protocol names (MCOMM, RaSTA, "
  "Kavach EI interface) as established Kavach protocol identifiers, and a non-standard Purdue labelling "
  "scheme. Third, completeness and scholarship concerns: there are no figures (an IRSE paper "
  "requires architecture and flow diagrams), the detection engineering is not mapped to MITRE "
  "ATT&CK for ICS, the EN 50159 threat-to-defence mapping is asserted but not tabulated, no "
  "risk-assessment / SL-T derivation method is given, a large fraction of the structural frameworks is "
  "drawn from a single commercial source without corroboration, related work is absent, and several "
  "internal cross-references are wrong. None of these is fatal; collectively they require more than "
  "cosmetic change. The recommendation is therefore Major Revision.")
p("On the individual criteria: technical correctness - good but with specific errors (Section B/C); "
  "railway applicability - strong; IEC 62443 compliance - strong conceptually, weak on the "
  "risk-assessment method (62443-3-2); EN 50159 consistency - mostly correct, but the category "
  "definitions are loose and the threat-defence mapping is missing; TS 50701 alignment - the lifecycle "
  "intent is right but the SL-T derivation is unsupported; OT-SOC architecture quality - high; detection "
  "engineering - good structure, lacks ATT&CK-for-ICS mapping and per-use-case data-source / "
  "false-positive specification; incident response - excellent (the four-gate, safety-first model is the "
  "manuscript's strongest operational contribution); governance - strong, but regulatory obligations "
  "(CERT-In timelines, NCIIPC CII designation) are asserted rather than substantiated; competency "
  "framework - strong and appropriately railway-specific; novelty - moderate (concentrated in the "
  "validated-architecture thesis; the frameworks are adapted from prior art); publishability - achievable "
  "after major revision.")

# ============================ B. MAJOR COMMENTS ============================
h1("B. Major Comments (ranked by severity)")
mc = [
 ("M1 - Evidence base is a model, not deployed systems (validity / over-claim).",
  "The findings referenced throughout (no cryptographic authentication on vital Category-1 buses, "
  "no MFA on engineering access, the forbidden S-Kavach-to-base-station edge, the unprotected "
  "backhaul) derive from an independent assessment of a REFERENCE ARCHITECTURE MODEL "
  "(reference [18]), not from instrumentation of fielded Kavach or electronic-interlocking "
  "installations. As written, the abstract and Sections 3, 5 and 10 invite the reader to treat these as "
  "empirical properties of Indian Railways' deployed estate. This is the most serious issue. The "
  "authors must (a) state plainly, in the abstract and at first use, that the evidence is model-derived; "
  "(b) add a Limitations section; and (c) identify field validation on representative installations as "
  "essential future work. Without this the paper risks asserting unverified deficiencies of national "
  "safety systems."),
 ("M2 - No figures or diagrams (publishability for an IRSE venue).",
  "The manuscript contains fourteen tables but not a single figure. An IRSE technical paper on a "
  "system architecture is expected to include, at minimum: (i) an IEC 62443 zone-and-conduit "
  "reference diagram for railway signalling; (ii) the three-echelon OT-SOC topology showing TAP/SPAN "
  "collection, data-diode log flow and the IT/OT boundary; (iii) a Kavach communication-architecture "
  "diagram distinguishing the legitimate Loco-to-Stationary association from the forbidden direct "
  "vital-to-RF edge; and (iv) a four-gate incident-response flowchart. Section H of this assessment "
  "specifies these figures."),
 ("M3 - Over-reliance on a single commercial source for the structural frameworks (independence / "
  "scholarship).",
  "The operating-model archetypes, five-layer visibility model, Monitoring Paradox, four-gate doctrine, "
  "severity model, three-tier KPI philosophy, maturity model and 90-day programme are all adapted "
  "from one commercial guide (reference [17]). For a peer-reviewed venue this concentration weakens "
  "both the perceived independence and the novelty of the contribution. The frameworks must be "
  "corroborated against, and positioned within, the primary literature: NIST SP 800-82r3, IEC "
  "62443-2-1 (CSMS / maturity), ENISA railway-sector threat-landscape work, and MITRE ATT&CK for "
  "ICS. The genuinely novel element - the validated-architecture / authority-separation thesis - should "
  "be foregrounded and clearly distinguished from the adapted material."),
 ("M4 - Technical inaccuracies in railway protocol and topology nomenclature (technical correctness).",
  "(a) The manuscript repeatedly writes 'RFID/EUROBALISE' for the Kavach balise air-gap. The ETCS "
  "Eurobalise and the Kavach RFID trackside tag are different technologies; Kavach uses RFID tags, "
  "not Eurobalises. Conflating them is a factual error for an Indian Railways audience. (b) The "
  "protocol names MCOMM, RaSTA and 'Kavach EI interface' are presented as established identifiers, "
  "but they originate in the reference model; the actual Kavach radio and EI-interface protocols are "
  "defined in the RDSO Kavach specification and must be cited as such or explicitly flagged as "
  "model abstractions. (c) The Purdue labelling ('L1 Telecom', 'L2 Telecom', 'L2/L1 Interlocking') is "
  "non-standard; the Purdue Enterprise Reference Architecture uses Levels 0-5 without domain "
  "suffixes. Either justify the railway-specific mapping explicitly or normalise it."),
 ("M5 - Detection-engineering methodology incomplete (detection quality).",
  "The use-case library (Table 5) is a good start but is not engineering-complete: each use case "
  "should specify its data source, detection logic, validated playbook reference and a documented "
  "false-positive baseline, and the library should be mapped to MITRE ATT&CK for ICS techniques so "
  "that coverage can be assessed against a recognised adversary model. In addition, the EN 50159 "
  "threat-to-defence relationship is asserted in prose but never tabulated; a normative "
  "threat-defence-detection mapping table is required (provided in Section H, Table R-2)."),
 ("M6 - No risk-assessment or SL-T derivation method (TS 50701 / IEC 62443-3-2 alignment).",
  "The paper correctly states that the SL-T per zone calibrates monitoring intensity, but it provides "
  "no method for deriving SL-T: no consequence scale, no likelihood scale, no zone/conduit "
  "risk-assessment worked example. For a planning paper addressed to RDSO and the Railway Board "
  "this is a substantive gap. A consequence-by-likelihood method consistent with IEC 62443-3-2 and "
  "TS 50701, with a worked Kavach-corridor example, must be added (outline provided in Section H)."),
 ("M7 - Internal cross-reference errors and overloaded 'Tier' terminology (clarity / correctness).",
  "Concrete errors: Section 5 cites the residual-risk register as 'Section 16' (it is Section 15); "
  "Section 14 states 'Table 8 defines the governance model' (it is Table 11); Section 15 states 'Table 9 "
  "is an extract of the residual-risk register' (it is Table 12). Separately, the word 'Tier' is used for "
  "four distinct concepts - the geographic SOC echelons (Tier 1/2/3 Station/Zonal/National), the "
  "analyst seniority grades (T1-T4), the KPI reporting audiences (Board/SOC-Lead/Analyst) and the "
  "'three-tier KPI framework' - so that 'Tier-3' denotes both the National SOC and the senior threat "
  "hunter. This is genuinely confusing and must be disambiguated (proposed scheme in Section H)."),
 ("M8 - Regulatory obligations asserted, not substantiated (governance quality).",
  "The paper invokes a CERT-In 'statutory window' and NCIIPC 'designation of signalling as Critical "
  "Information Infrastructure' as if settled. The CERT-In Directions of 28 April 2022 impose a 6-hour "
  "incident-reporting obligation; whether, and how, railway signalling OT incidents fall under those "
  "Directions and under the NCIIPC CII regime (IT Act 2000, s.70) - and the precedence between the "
  "two reporting channels and the internal RDSO/Railway-Board chain - must be stated with specific "
  "citations, not assumed. The escalation timelines in the severity model (Table 8) should be "
  "reconciled explicitly against the Kavach Cyber Security SOP rather than presented as the authors' "
  "own figures."),
]
for hd, bd in mc:
    p(hd, bold=True); p(bd)

# ============================ C. MINOR COMMENTS ============================
h1("C. Minor Comments")
minor = [
 "m1 - SPAN versus TAP are conflated. A SPAN/mirror port can drop frames under oversubscription "
 "and is part of the switch's own resource budget; a passive optical/copper TAP cannot affect the "
 "monitored link. For SIL-4 vital segments the paper should mandate passive TAP and reserve SPAN "
 "for non-vital segments.",
 "m2 - The data-flow model (logs upward by data diode, control downward and authenticated) appears "
 "to contradict the unidirectional isolation it relies on. Clarify: either dual-diode with a separate "
 "out-of-band management path, or no online downward path to the Station echelon at all.",
 "m3 - Quantitative figures (14-18 FTE; 12-24-month talent pipeline; >99.5% sensor uptime; MTTD < 1 "
 "hour) are inherited from cross-sector practice and stated as railway targets without justification. "
 "Attribute them or justify them for the Indian Railways operating context.",
 "m4 - KPI targets expressed as 'approaching 100%' or 'minimise' are not measurable. Give concrete, "
 "time-bound values (e.g. 100% SIL-3/4 conduit coverage within 18 months of go-live).",
 "m5 - 'IDMZ' (industrial DMZ) is used before it is defined; define on first use and add a nomenclature "
 "table for the mixed signalling/cyber readership.",
 "m6 - The MTTD KPI is anchored to 'EN 50129 / FR6'; EN 50129 sets no detection-time target. Anchor "
 "MTTD to IEC 62443-3-3 SR 6.2 and to the operational hazard-exposure budget instead.",
 "m7 - The EN 50159 Category 2 description ('not under full control but threats partly mitigated') is "
 "loose; use the normative sense - a transmission system whose unauthorised access can be "
 "considered negligible/excluded, hence cryptographic defences are not mandated (these are required "
 "for Category 3).",
 "m8 - No indicative cost or cost-benefit model is offered, although the audience includes the Railway "
 "Board. A high-level capital/operating envelope for the pilot would strengthen the planning value.",
 "m9 - The byline carries a placeholder; add author names, affiliations and the IRSE paper-template "
 "front matter (and ORCID where applicable).",
 "m10 - References [10] (Kavach Cyber Security SOP) and [11] (IR signalling cyber security policy) lack "
 "document numbers and dates; tighten to specific issues.",
 "m11 - 'The majority of OT incidents begin with human error or supply-chain compromise' is stated as "
 "fact; add a citation (e.g. the cited cross-sector source and ENISA/Dragos annual ICS reports).",
 "m12 - Define MCOMM, RaSTA, MVB, CAN, BTS and SL-T in a nomenclature table; the audience spans "
 "signalling engineers and cyber practitioners who will not share all acronyms.",
]
for m in minor:
    bullet(m)

# ============================ D. MISSING REFERENCES ============================
h1("D. Missing References")
for r in [
 "MITRE ATT&CK for ICS - adversary tactics and techniques for industrial control systems (detection "
 "coverage mapping).",
 "NIST SP 800-82 Rev. 3, Guide to Operational Technology (OT) Security (cited generically as [14]; use "
 "the current revision and cite specific control families).",
 "IEC 62443-3-2, Security risk assessment for system design - cited as [1] but its risk-assessment "
 "method is not applied; cite and use it for SL-T derivation.",
 "IEC 62443-2-1 / IEC 62443-2-4, CSMS and service-provider requirements - to anchor the maturity "
 "and MSSP-governance discussion.",
 "ENISA, Railway Cybersecurity / Transport Threat Landscape reports - for the threat landscape and "
 "the human-error / supply-chain incidence claim.",
 "CERT-In Directions under sub-section (6) of section 70B of the IT Act, 28 April 2022 - for the "
 "6-hour incident-reporting obligation.",
 "NCIIPC / IT Act 2000 section 70 and the Critical Information Infrastructure Rules - for the CII "
 "designation basis.",
 "EU railway-cyber research outputs (CYRAIL, X2Rail, SAFETY4RAILS) and prior IRSE technical "
 "papers on signalling cyber security - for related-work positioning.",
 "RaSTA reference (IEC 61375 / DIN VDE V 0831-200) - if RaSTA is genuinely claimed for the EI/Kavach "
 "interface; otherwise remove the claim.",
]:
    bullet(r)

# ============================ E. REQUIRED REVISIONS ============================
h1("E. Required Revisions (checklist)")
for r in [
 "R1. Bound the evidence base as model-derived in the abstract and Section 5; add a Limitations "
 "section and field-validation future work (addresses M1).",
 "R2. Add the four required figures and reference them in the text (addresses M2).",
 "R3. Add a Related Work section and corroborate the adapted frameworks with primary literature; "
 "foreground the novel validated-architecture thesis (addresses M3).",
 "R4. Correct the Eurobalise/RFID conflation; qualify MCOMM/RaSTA/Kavach-EI-interface as RDSO-spec "
 "or model terms; normalise or justify the Purdue labelling (addresses M4).",
 "R5. Add the EN 50159 threat-defence-detection mapping table and the MITRE ATT&CK-for-ICS mapping; "
 "complete the use-case specification fields (addresses M5).",
 "R6. Add a TS 50701 / IEC 62443-3-2 risk-assessment and SL-T-derivation method with a worked "
 "Kavach example (addresses M6).",
 "R7. Fix the three cross-reference errors and disambiguate 'Tier' terminology throughout "
 "(addresses M7).",
 "R8. Substantiate CERT-In and NCIIPC obligations with citations and reconcile severity SLAs with the "
 "Kavach SOP (addresses M8).",
 "R9. Address minor comments m1-m12.",
]:
    bullet(r)

# ============================ F. RECOMMENDATION ============================
h1("F. Publication Recommendation")
p("MAJOR REVISION.", bold=True)
p("The manuscript's core thesis and structure are sound and the contribution is worthwhile, but the "
  "validity bounding of the evidence base (M1), the absence of figures (M2), the technical "
  "inaccuracies (M4) and the methodological gaps (M5, M6) are beyond cosmetic and must be "
  "remedied before acceptance. None of the issues is fatal; a revised manuscript addressing the "
  "checklist in Section E would, in the reviewers' joint judgement, be acceptable for publication.")

# ============================ G. REWRITTEN SECTIONS ============================
h1("G. Rewritten Sections")
p("The following are the revised texts for the sections requiring change. Unchanged sections of the "
  "manuscript are retained as submitted. New tables introduced by this review are numbered R-1 "
  "onward to avoid collision with the manuscript's own numbering.", italic=True)

h2("Abstract (revised - bounds the evidence base)")
p("[...unchanged opening...] The evidence base is an independent cyber security assessment of a "
  "reference railway OT architecture MODEL - a standards-traceable representation in which the "
  "semantic, governance and enforcement authorities were formally separated and validated; the "
  "findings reported here are therefore properties of that model and of the governance logic it "
  "encodes, and are presented as illustrative monitoring requirements rather than as measured "
  "deficiencies of any deployed Indian Railways installation. Field validation on representative Kavach "
  "and electronic-interlocking installations is identified as essential future work. With that scope "
  "understood, the assessment surfaced durable engineering findings (absence of cryptographic "
  "authentication on vital Category-1 buses, absence of multi-factor authentication on privileged "
  "engineering access, absence of timeliness monitoring on open radio bearers, incompletely "
  "instrumented trust boundaries and an unprotected telecom backhaul) and standards-compliance "
  "findings (the EN 50159 / IEC 62443 authentication-crediting tension and the need for TS 50701 "
  "lifecycle evidence). [...remainder of abstract unchanged...]")

h2("Section 4.2 (revised) - EN 50159 transmission categories, with threat-defence mapping")
p("EN 50159 classifies safety-related transmission systems into three categories defined by the "
  "controllability of the medium and the consequent need for cryptographic protection. In a Category 1 "
  "(closed) system the participants are fixed and known, the transmission characteristics are fixed, "
  "and unauthorised access to the medium is excluded by design; the safety defences are non-"
  "cryptographic (safety code, sequence number, time-out). In a Category 2 (controlled) system the "
  "transmission system is not under the full control of the designer, but the residual risk of "
  "unauthorised access is assessed as negligible and can be excluded by design or by procedural "
  "means, so cryptographic protection is not mandated. In a Category 3 (open) system the possibility of "
  "unauthorised access must be assumed; cryptographic procedures (message authentication, "
  "cryptographic identification) are therefore required in addition to the non-cryptographic defences. "
  "The Kavach radio interface and the balise air-gap are Category 3; the managed MPLS/IP backhaul is "
  "Category 2; the vital station buses are Category 1. Because the category fixes the threat model, it "
  "fixes the detector set the OT-SOC must field; Table R-1 makes the EN 50159 threat-defence-"
  "detection relationship explicit.")
table("EN 50159 threats mapped to safety defences and to OT-SOC detections",
      ["EN 50159 threat", "EN 50159 defence(s)", "OT-SOC detection (this paper)"],
      [["Repetition", "Sequence number; time stamp; time-out", "Sequence regression; duplicate-frame and freshness analysis (UC-03)"],
       ["Deletion", "Sequence number", "Sequence-gap detection; loss-of-message alert (UC-04)"],
       ["Insertion", "Sequence number; source/destination identifier; feedback; identification procedure", "Source-ID whitelist; rate/insertion anomaly (UC-01, UC-13)"],
       ["Re-sequencing", "Sequence number; time stamp", "Out-of-order sequence detection (UC-03)"],
       ["Corruption", "Safety code; cryptographic techniques", "Safety-code/CRC failure-rate and integrity-flag monitoring (UC-02, UC-10)"],
       ["Delay", "Time stamp; time-out", "End-to-end latency vs SIL-bounded threshold (UC-03)"],
       ["Masquerade", "Source/destination identifier; cryptographic identification; identification procedure", "Identity/closure monitoring; unexpected-participant detection (UC-01, UC-13)"]])
p("Two consequences for the OT-SOC follow and resolve the manuscript's central evidence finding. "
  "First, for the Category-1 vital buses EN 50159 does not mandate cryptographic source "
  "authentication - the closed-medium assumption and the safety code carry the assurance - which is "
  "why the reference model's vital protocols provide integrity and replay but not authentication; this is "
  "a standards-consistent design, not a defect, PROVIDED the closed-medium assumption holds, and the "
  "OT-SOC's role is to police that assumption continuously. Second, for the Category-3 open bearers "
  "cryptographic protection is mandatory and is necessary but not sufficient: timeliness, replay and "
  "RF-anomaly monitoring remain obligatory because delay and masquerade are first-class EN 50159 "
  "threats that cryptography alone does not detect in operation.")

h2("Section 5 (revised closing) - corrected cross-reference and Limitations")
p("[...unchanged body...] The residual findings of that assessment - referenced throughout this paper "
  "as evidence and consolidated in the residual-risk register of Section 15 - become the OT-SOC's "
  "standing monitoring requirements rather than items to be quietly closed.")
p("Limitations of the evidence base. The findings cited in this paper were obtained by assessing a "
  "reference architecture model, not by instrumenting deployed Kavach or electronic-interlocking "
  "equipment. They are therefore properties of the model and of the governance rules it encodes; their "
  "value here is to illustrate, concretely and traceably, the classes of monitoring requirement an "
  "OT-SOC must satisfy. Whether any specific finding (for example, the absence of cryptographic "
  "authentication on a particular vital interface) holds for a given fielded installation is a question of "
  "fact that must be established by a field assessment of that installation under the TS 50701 "
  "lifecycle. Such field validation - on a representative Kavach corridor and a representative electronic "
  "interlocking - is the immediate next step and is recorded as future work.")

h2("Section 7 (revised extract) - corrected balise nomenclature and TAP/SPAN precision")
p("The Process/Vital-state telemetry for the balise air-gap concerns the Kavach trackside RFID tag and "
  "the onboard RFID reader. (These are Kavach RFID balises; they are distinct from the ETCS "
  "Eurobalise and should not be conflated with it.) The reader decodes the passive, safety-coded "
  "telegram and derives position; the OT-SOC monitors telegram integrity and positional plausibility "
  "but does not expect cryptographic encryption or replay protection on this leg (Section 10). For all "
  "vital segments, collection must use a passive optical or copper TAP rather than a switch SPAN/mirror "
  "port: a TAP cannot, by construction, affect the monitored link, whereas a SPAN port consumes "
  "switch resources and can silently drop frames under load - an unacceptable risk on a SIL-4 bus. "
  "SPAN may be used only on non-vital segments where the consequence of frame loss is tolerable.")

h2("Section 8 (revised extract) - TAP/SPAN and the data-flow reconciliation")
p("The passivity requirement has a precise engineering meaning. On vital (Category-1) and onboard "
  "segments the collection device must be a passive TAP that is electrically and logically incapable of "
  "transmitting onto the monitored link; SPAN mirroring, which is a function of the production switch "
  "and can perturb it, is excluded from vital segments. The Monitoring Paradox is then enforced "
  "architecturally: telemetry leaves the Station echelon strictly upward through a data diode, and there "
  "is no online downward control path into the vital tier. Sensor configuration and detection-content "
  "updates are applied to Station-echelon collectors through a separate, scheduled, out-of-band "
  "management channel under change control - not through the monitoring path - so that the upward "
  "log diode remains genuinely unidirectional and the vital network cannot be reached from the SOC "
  "during normal operation.")

h2("Section 9 (revised opening) - disambiguated terminology")
p("Terminology. To avoid ambiguity this paper distinguishes three uses of the word that were "
  "previously all rendered as 'tier'. The geographic SOC structure is described in ECHELONS - E1 "
  "Station/Wayside, E2 Zonal/Divisional, E3 National; the analyst workforce is described in GRADES - "
  "G1 Monitoring Operator, G2 Investigator, G3 Senior Analyst/Hunter, G4 Lead; and the KPI framework "
  "is described by reporting AUDIENCE - Board, SOC-Lead, Analyst. The federated hybrid recommended "
  "in Section 6 is realised as a three-echelon architecture (E1 Station, E2 Zonal, E3 National) staffed "
  "across grades G1-G4. [...the remainder of Section 9 is retained, with 'Tier 1/2/3' replaced by "
  "'E1/E2/E3' throughout...]")

h2("Section 10 (revised extract) - protocol qualification and ATT&CK-for-ICS mapping")
p("The vital Category-1 buses - the electronic interlocking to Stationary-Kavach interface, the "
  "interlocking to object-controller and speed-sensor to Loco-Kavach links, and the Loco-Kavach to "
  "driver-machine-interface and brake-interface-unit links - carry, in the reference model, protocols "
  "denoted Kavach-EI-interface, CAN and MVB. (These identifiers are those of the reference model and "
  "of common rolling-stock buses; the authoritative protocol definitions for the Kavach radio and "
  "EI-interface are given in the RDSO Kavach specification, reference [9], and the manuscript's claims "
  "should be read against that specification rather than as independent protocol assertions.) To allow "
  "detection coverage to be assessed against a recognised adversary model, each use case is mapped "
  "to MITRE ATT&CK for ICS in Table R-2; this also exposes coverage gaps for future use-case "
  "development.")
table("Railway OT-SOC use cases mapped to MITRE ATT&CK for ICS",
      ["Use case", "ATT&CK for ICS tactic", "Representative technique"],
      [["UC-01 Spoofed vital field message", "Impair Process Control / Inhibit Response", "Spoof Reporting Message; Modify Parameter"],
       ["UC-03 Kavach RF replay / delay", "Inhibit Response Function / Impair Process Control", "Adversary-in-the-Middle; Manipulation of Control"],
       ["UC-04 RF jamming / loss of comms", "Inhibit Response Function", "Denial of Service; Block Reporting Message"],
       ["UC-05 Unauthorised logic download", "Lateral Movement / Execution", "Program Download; Modify Controller Tasking"],
       ["UC-07 Trust-boundary bypass", "Lateral Movement", "Exploitation of Remote Services; Default Credentials"],
       ["UC-08 Forbidden topology edge", "Initial Access / Persistence", "Hardware Additions; Rogue Master"],
       ["UC-13 New device on vital segment", "Initial Access", "Rogue Master; Hardware Additions"],
       ["UC-14 Removable media on EWS", "Initial Access", "Replication Through Removable Media"]])
p("Each use case in Table 5 should additionally record, in the deployed detection-content repository, "
  "its primary data source, the precise detection logic, the validated response playbook reference and "
  "a documented false-positive baseline; these operational fields are omitted from the paper for "
  "brevity but are mandatory in the engineering artefact.")

h2("New Section - Related Work (to be inserted after Section 2)")
p("Railway cyber security has an established literature that this paper builds upon. The CENELEC "
  "lifecycle standards (EN 50126, EN 50129) and the dedicated cyber standard CLC/TS 50701 provide "
  "the safety-security integration framework; IEC 62443 supplies the zone/conduit and security-level "
  "model adopted here; and EN 50159 supplies the communication-threat taxonomy. EU research "
  "programmes (CYRAIL, X2Rail, SAFETY4RAILS) have addressed railway-specific threat assessment and "
  "resilience, and ENISA has published sectoral transport and railway threat-landscape analyses. In "
  "the OT-SOC operational domain, NIST SP 800-82 and the ISA/IEC 62443 series define monitoring and "
  "management-system practice, and MITRE ATT&CK for ICS provides the adversary-technique reference "
  "now standard in detection engineering. The distinctive contribution of the present paper relative to "
  "this body of work is twofold: the argument that a machine-validated, authority-separated "
  "architecture model must precede and configure the OT-SOC, and the end-to-end railway "
  "contextualisation - from operating-model selection to a 90-day programme - anchored to Kavach and "
  "to Indian Railways' institutional structure. The structural OT-SOC frameworks adapted in Sections "
  "6-17 are drawn from cross-sector practice and are cited as such; they are not claimed as original.")

h2("New Section - Risk Assessment and SL-T Derivation (to be inserted before Section 6)")
p("TS 50701 and IEC 62443-3-2 require that target Security Levels (SL-T) be derived for each zone and "
  "conduit from a documented risk assessment, and the OT-SOC uses SL-T to calibrate monitoring "
  "intensity and response targets. This paper recommends a consequence-by-likelihood method. "
  "Consequence is scored primarily by the safety significance of the affected function - its SIL and its "
  "potential to cause a wrong-side failure - and secondarily by operational and reputational impact; a "
  "SIL-4 vital function (interlocking, Kavach) takes the highest consequence band by default. Likelihood "
  "is scored from exposure (Internet/enterprise reachability, open RF, vendor/engineering access), the "
  "EN 50159 transmission category of the conduit, and the strength of existing controls. The SL-T for a "
  "zone is the level required to reduce the residual risk of its highest-consequence credible scenario "
  "to acceptable; conduits inherit the higher SL-T of the zones they join. As a worked example, the "
  "Interlocking zone (SIL-4 consequence; Category-1 internal but reachable via engineering access and, "
  "through Kavach, exposed to an open Category-3 bearer) is assigned SL-T 3, driving the highest "
  "monitoring intensity, the shortest detection-to-advisory target and mandatory privileged-access and "
  "closure monitoring; the Telecom Core (Category-2, no direct vital consequence) is assigned SL-T 2; "
  "and the Enterprise zone SL-T 1-2. The full zone/conduit risk register is an artefact of the TS 50701 "
  "cyber security case and is the prerequisite (Phase 0) for any sensor deployment.")

h2("Section 13 (revised) - concrete KPI targets")
p("KPI targets must be measurable and time-bound. The revised framework replaces the qualitative "
  "targets with explicit values: vital-asset (SIL-3/4 conduit) monitoring coverage to reach 100% within "
  "18 months of pilot go-live and be sustained thereafter; trust-boundary instrumentation 100% of "
  "governance-defined boundaries within 12 months; EN 50159 detector completeness 100% (every "
  "conduit fielding the detector set its transmission category requires); sensor/TAP uptime not less "
  "than 99.5% monthly; mean time to detect a known vital-affecting technique under one hour and to "
  "first safe advisory under one hour for SEV1; architecture-conformance drift and exempt-flow false "
  "positives both zero; and privileged-session audit coverage 100% within 24 hours of session "
  "completion. The mean-time-to-detect target is anchored to IEC 62443-3-3 SR 6.2 (continuous "
  "monitoring) and to the operational hazard-exposure budget, not to EN 50129.")

h2("Section 14 (revised extract) - corrected reference and grade/echelon clarity")
p("Signalling and train-protection systems are Critical Information Infrastructure; their protection "
  "therefore engages NCIIPC for designation and national protection oversight under section 70 of the "
  "Information Technology Act 2000, and CERT-In for statutory incident reporting under the CERT-In "
  "Directions of 28 April 2022 (which impose reporting within six hours of noticing a reportable "
  "incident), in addition to the internal railway hierarchy. The interplay must be defined in the SOC "
  "charter: the OT-SOC notifies the internal RDSO/Railway-Board chain and, for incidents meeting the "
  "statutory thresholds, CERT-In within the mandated window, with NCIIPC engaged for CII-level "
  "events. Table 11 (not Table 8) defines the governance model and the principal responsibility of each "
  "entity. The staffing grades G1-G4 of Table 10 are distinct from the geographic echelons E1-E3 of "
  "Section 9: the IRSSE, Kavach and functional-safety subject-matter experts are grade-G2/G3 roles "
  "that may serve across echelons. The establishment figures cited (of the order of 14-18 full-time "
  "equivalents for 24x7 cover, and a 12-24-month competency pipeline) are adopted from cross-sector "
  "staffing experience [17] and should be validated against an Indian Railways workload study during "
  "the pilot.")

h2("Section 15 (revised extract) - corrected reference")
p("[...unchanged body...] Table 12 (not Table 9) is an extract of the residual-risk register populated "
  "from the assessment evidence, with each item routed to its acceptance authority; risk affecting any "
  "SIL-4 vital function is routed to the Railway Board and cannot be accepted at SOC or divisional "
  "level.")

h2("H. Required Figures (specification for the revision)")
p("The revision must add at least the following figures, each with a descriptive caption:", )
bullet("Figure 1 - IEC 62443 zone-and-conduit reference model for railway signalling: the zone "
       "hierarchy (Enterprise, IDMZ, Security-Management, Operations, Telecom Core, Radio Access, "
       "Interlocking, Field, Onboard) with the typed conduits and their EN 50159 transmission "
       "categories annotated.")
bullet("Figure 2 - Three-echelon OT-SOC topology (E1 Station, E2 Zonal, E3 National) showing passive "
       "TAP/SPAN collection, the unidirectional data-diode log path upward, the out-of-band "
       "management path, and the IT/OT boundary.")
bullet("Figure 3 - Kavach communication architecture distinguishing the legitimate end-to-end "
       "Loco-to-Stationary safety association over the radio bearer from the FORBIDDEN direct "
       "vital-to-base-station edge, with EN 50159 categories marked.")
bullet("Figure 4 - Four-gate incident-response flowchart (G1 Safety, G2 Operational impact, G3 "
       "Containment, G4 Recovery authorisation) showing decision authorities and the "
       "observe/advise/recommend escalation path under safe-working rules.")

# ---- save ----
doc.save("Reviewer_Assessment.docx")
print("SAVED Reviewer_Assessment.docx")
